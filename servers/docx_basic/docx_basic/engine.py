"""DOCX Basic engine — pure python-docx logic, no MCP imports."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from shared.address_resolver import (
    AddressError,
    build_docx_index,
    fetch_section_content,
)
from shared.file_utils import hint_for_error, resolve_path
from shared.handover import make_context, make_handover
from shared.live_edit import notify_reload
from shared.platform_utils import get_max_paragraphs, get_max_search_results, open_file
from shared.progress import fail, index_range, info, ok, warn
from shared.version_control import snapshot

from .helpers import (
    _error,
    _log_receipt,
    _not_found,
    _wrong_type,
    delete_at,
    diff_versions,
    get_history_tool,
    insert_at,
    read_receipt_tool,
    replace_at,
    restore_version,
)

# Re-export so existing imports from engine still work
__all__ = [
    "delete_at",
    "diff_versions",
    "get_history_tool",
    "insert_at",
    "read_receipt_tool",
    "replace_at",
    "restore_version",
]

logger = logging.getLogger(__name__)


def get_document_outline(file_path: str) -> dict[str, Any]:
    """Return only headings and paragraph indices — the structural skeleton."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        paragraphs = doc.paragraphs
        total = len(paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        outline: list[dict[str, Any]] = []
        for i, p in enumerate(paragraphs):
            name = p.style.name  # type: ignore[reportOptionalMemberAccess]
            if name and name.startswith("Heading "):  # type: ignore[reportOptionalMemberAccess]
                try:
                    level = int(name.split()[-1])  # type: ignore[reportOptionalMemberAccess]
                    if 1 <= level <= 6:
                        outline.append({"index": i, "level": level, "text": p.text})
                except (ValueError, IndexError):
                    pass

        progress.append(ok("Built outline", f"{len(outline)} headings"))

        return {
            "success": True,
            "file": str(path),
            "total_paragraphs": total,
            "outline": outline,
            "hint": "Pass an outline heading index to read_paragraph() or use get_document_index() for full section addresses.",
            "handover": make_handover(
                workflow_step="COLLECT",
                suggested_next=[
                    {
                        "tool": "get_document_index",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "get addressable section tree for surgical edits",
                    },
                    {
                        "tool": "search_paragraphs",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "find specific text within the document",
                    },
                ],
                carry_forward={"file_path": str(path)},
            ),
            "progress": progress,
            "token_estimate": len(str(outline)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check that file_path is a valid .docx file.", progress)


def get_document_index(file_path: str) -> dict[str, Any]:
    """Return section tree index. Zero paragraph content returned."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        index_data = build_docx_index(doc)
        total = index_data["total_paragraphs"]
        progress.append(ok(f"Indexed {path.name}", f"{total} paragraphs"))

        return {
            "success": True,
            "file": str(path),
            "total_paragraphs": total,
            "address_scheme": index_data["address_scheme"],
            "sections": index_data["sections"],
            "hint": "Use an address like '§2' or '§2.p4' with fetch_section() to read targeted content.",
            "handover": make_handover(
                workflow_step="INSPECT",
                suggested_next=[
                    {
                        "tool": "fetch_section",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "read content of a specific section by address",
                    },
                    {
                        "tool": "search_paragraphs",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "find paragraphs matching a query",
                    },
                ],
                carry_forward={"file_path": str(path)},
            ),
            "progress": progress,
            "token_estimate": len(str(index_data["sections"])) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check that file_path is a valid .docx file.", progress)


def fetch_section(file_path: str, address: str) -> dict[str, Any]:
    """Return paragraphs of one section. Only addressed content returned."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        doc = Document(str(path))
        progress.append(ok(f"Opened {path.name}"))

        content = fetch_section_content(doc, address)
        para_count = len(content["paragraphs"])
        progress.append(ok(f"Fetched {address}", f"{para_count} paragraphs"))

        return {
            "success": True,
            "file": str(path),
            **content,
            "hint": "Call replace_text() or apply_patch() to edit the paragraphs you found.",
            "handover": make_handover(
                workflow_step="INSPECT",
                suggested_next=[
                    {
                        "tool": "replace_text",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "replace text in this section",
                    },
                    {
                        "tool": "search_paragraphs",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "narrow down to a specific paragraph",
                    },
                ],
                carry_forward={"file_path": str(path), "address": address},
            ),
            "progress": progress,
            "token_estimate": len(str(content)) // 4,
        }
    except AddressError as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "hint": "Use get_document_index to see valid section addresses.",
            "progress": progress,
            "token_estimate": 30,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path and address.", progress)


def read_document(file_path: str) -> dict[str, Any]:
    """Read full doc. Use get_document_index for large files (>10 pages)."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        all_paras = doc.paragraphs
        total = len(all_paras)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        max_p = get_max_paragraphs() * 3  # read_document gets 3x the normal limit
        truncated = total > max_p
        para_list = all_paras[:max_p]

        paragraphs = [{"index": i, "text": p.text, "style": p.style.name} for i, p in enumerate(para_list)]  # type: ignore[reportOptionalMemberAccess]

        result: dict[str, Any] = {
            "success": True,
            "file": str(path),
            "paragraph_count": total,
            "paragraphs": paragraphs,
            "hint": "Call search_paragraphs() or get_document_index() to locate text before editing.",
            "context": make_context(
                op="read_document",
                summary=f"Opened {path.name} — {total} paragraphs",
                artifacts=[{"type": "docx", "path": str(path), "role": "input"}],
            ),
            "handover": make_handover(
                workflow_step="INSPECT",
                suggested_next=[
                    {
                        "tool": "search_paragraphs",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "locate specific text before editing",
                    },
                    {
                        "tool": "get_document_index",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "get section tree for large documents",
                    },
                    {
                        "tool": "replace_text",
                        "server": "docx_basic",
                        "domain": "office",
                        "reason": "make a targeted text replacement",
                    },
                ],
                carry_forward={"file_path": str(path)},
            ),
            "progress": progress,
            "token_estimate": len(str(paragraphs)) // 4,
        }

        if truncated:
            result["truncated"] = True
            result["truncated_at"] = max_p
            result["warning"] = "Large document. Use get_document_index + fetch_section for targeted access."
            progress.append(
                warn(
                    f"Document truncated at {max_p} paragraphs",
                    "Use get_document_index for large files",
                )
            )

        return result

    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check that file_path is a valid .docx file.", progress)


def read_paragraph(file_path: str, index: int) -> dict[str, Any]:
    """Return one paragraph with full run detail."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        doc = Document(str(path))
        paras = doc.paragraphs
        total = len(paras)

        if index < 0 or index >= total:
            progress.append(fail(f"Index {index} out of range"))
            return {
                "success": False,
                "error": f"paragraph_index {index} out of range {index_range(total, 'paragraphs')}",
                "hint": "Use get_document_outline to see valid paragraph indices.",
                "progress": progress,
                "token_estimate": 20,
            }

        p = paras[index]
        runs = [
            {
                "text": r.text,
                "bold": r.bold,
                "italic": r.italic,
                "font": r.font.name,
                "size": r.font.size.pt if r.font.size else None,
            }
            for r in p.runs
        ]
        progress.append(ok(f"Read paragraph {index}", f"{len(runs)} runs"))

        return {
            "success": True,
            "index": index,
            "text": p.text,
            "style": p.style.name,  # type: ignore[reportOptionalMemberAccess]
            "runs": runs,
            "progress": progress,
            "token_estimate": len(str(runs)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path and index.", progress)


def read_paragraph_range(file_path: str, start_index: int, end_index: int) -> dict[str, Any]:
    """Return a bounded range of paragraphs. Max 50 (20 in 8GB mode)."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        max_range = get_max_paragraphs()
        if end_index - start_index > max_range:
            return {
                "success": False,
                "error": f"Range exceeds {max_range} paragraphs.",
                "hint": f"Use smaller ranges. Max range is {max_range} paragraphs.",
                "progress": [fail(f"Range too large ({end_index - start_index})")],
                "token_estimate": 15,
            }

        doc = Document(str(path))
        paras = doc.paragraphs
        total = len(paras)

        start = max(0, start_index)
        end = min(end_index, total - 1)

        result_paras = [
            {"index": i, "text": paras[i].text, "style": paras[i].style.name}  # type: ignore[reportOptionalMemberAccess]
            for i in range(start, end + 1)
        ]
        progress.append(ok(f"Read paragraphs {start}-{end}", f"{len(result_paras)} paragraphs"))

        return {
            "success": True,
            "start_index": start,
            "end_index": end,
            "paragraphs": result_paras,
            "truncated": False,
            "progress": progress,
            "token_estimate": len(str(result_paras)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path and indices.", progress)


def search_paragraphs(file_path: str, query: str, max_results: int = 10) -> dict[str, Any]:
    """Scan paragraphs for matching text. Returns only matches."""
    progress: list[dict[str, Any]] = []
    if not query:
        return {
            "success": False,
            "error": "query cannot be empty",
            "hint": "Provide a non-empty search string.",
            "progress": [fail("Empty query")],
            "token_estimate": 10,
        }

    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        max_results = min(max_results, get_max_search_results())
        doc = Document(str(path))
        paras = doc.paragraphs
        total = len(paras)

        matches: list[dict[str, Any]] = []
        q_lower = query.lower()
        for i, p in enumerate(paras):
            if q_lower in p.text.lower():
                matches.append({"index": i, "text": p.text, "style": p.style.name})  # type: ignore[reportOptionalMemberAccess]
                if len(matches) >= max_results:
                    break

        progress.append(
            ok(
                f"Found {len(matches)} matches for '{query}'",
                f"scanned {total} paragraphs",
            )
        )

        return {
            "success": True,
            "query": query,
            "matches": matches,
            "total_paragraphs_scanned": total,
            "truncated": len(matches) >= max_results,
            "progress": progress,
            "token_estimate": len(str(matches)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path and query.", progress)


def replace_text(
    file_path: str,
    match_text: str,
    new_text: str,
    preserve_style: bool = True,
    dry_run: bool = False,
    open_after: bool = False,
) -> dict[str, Any]:
    """Find text and replace in-place, preserving run formatting."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    # preserve_style was declared on the tool, forwarded by the wrapper, and
    # read nowhere: the replacement goes through docxedit.replace_string, which
    # preserves run formatting and offers no way not to. So preserve_style=False
    # did exactly what preserve_style=True does, and said it succeeded.
    if not preserve_style:
        return {
            "success": False,
            "op": "replace_text",
            "error": "replace_text cannot drop the existing run formatting",
            "hint": (
                "It replaces text inside the runs, which keeps their formatting. Leave "
                "preserve_style=True, then use set_font() or set_paragraph_style() to restyle."
            ),
            "progress": [fail("Unsupported argument", "preserve_style=False")],
            "token_estimate": 40,
        }
    try:
        import docxedit  # type: ignore[import-untyped]
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        para_count = len(doc.paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{para_count} paragraphs"))

        # Count occurrences first
        matches: list[int] = []
        for i, p in enumerate(doc.paragraphs):
            if match_text in p.text:
                matches.append(i)

        if not matches:
            progress.append(fail(f"'{match_text}' not found"))
            _log_receipt(
                file_path,
                "replace_text",
                {"match_text": match_text, "new_text": new_text},
                "✘ match_text not found",
                None,
                False,
            )
            return {
                "success": False,
                "error": "match_text not found in document",
                "hint": ("Use search_paragraphs to verify the exact text before replacing."),
                "progress": progress,
                "token_estimate": 20,
            }

        if dry_run:
            progress.append(
                info(
                    f"[DRY RUN] Would replace '{match_text}' → '{new_text}'",
                    f"{len(matches)} occurrence{'s' if len(matches) != 1 else ''}",
                )
            )
            progress.append(info("Dry run complete — no changes made"))
            return {
                "success": True,
                "op": "replace_text",
                "dry_run": True,
                "would_replace": len(matches),
                "match_locations": matches,
                "progress": progress,
                "token_estimate": 30,
            }

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        # docxedit.replace_string() has no return statement (always None) —
        # use the paragraph match count computed above instead.
        docxedit.replace_string(doc, match_text, new_text)
        replaced = len(matches)
        progress.append(
            ok(
                f"Replaced '{match_text}' → '{new_text}'",
                f"{replaced} occurrence{'s' if replaced != 1 else ''}",
            )
        )

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(ok(f"Saved {path.name}", f"{para_count} paragraphs"))
        progress.append(notify_reload(str(path), "docx"))

        _log_receipt(
            file_path,
            "replace_text",
            {"match_text": match_text, "new_text": new_text},
            f"✔ Replaced {replaced} occurrences",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "replace_text",
            "match": match_text,
            "new_text": new_text,
            "replaced_count": replaced,
            "replaced_in_paragraphs": matches,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        _log_receipt(
            file_path,
            "replace_text",
            {"match_text": match_text, "new_text": new_text},
            f"✘ {e}",
            backup,
            False,
        )
        return {
            "success": False,
            "error": str(e),
            "backup": backup,
            "hint": hint_for_error(e, path),
            "progress": progress,
            "token_estimate": 20,
        }


def insert_paragraph(
    file_path: str,
    after_index: int,
    text: str,
    style: str = "Body Text",
    open_after: bool = False,
) -> dict[str, Any]:
    """Insert a paragraph after index N with the given style."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        doc = Document(str(path))
        paras = doc.paragraphs
        total = len(paras)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        # A brand-new document has no paragraphs, so there was no anchor to
        # insert after and every index failed. -1 -- the documented "insert at
        # the beginning" value, and the one a caller reaches for on an empty
        # file -- failed worst: it passed this guard, reached doc.paragraphs[0]
        # below and surfaced IndexError to the caller as "list index out of
        # range", with a hint offering to undo a change that never happened.
        # With no paragraphs there is exactly one position, so any index means
        # it and there is nothing to reject.
        if total and (after_index < -1 or after_index >= total):
            progress.append(fail(f"Index {after_index} out of range"))
            return {
                "success": False,
                "error": f"paragraph_index {after_index} out of range {index_range(total, 'paragraphs')}",
                "hint": "Use read_document to get current paragraph count.",
                "progress": progress,
                "token_estimate": 15,
            }

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        # Insert after the target paragraph using XML manipulation
        from docx.oxml import OxmlElement  # type: ignore[import-untyped]

        new_para = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_r.append(new_t)
        new_para.append(new_r)

        if total == 0:
            # The body ends with a w:sectPr; appending the raw element would
            # land the paragraph after it, which is not valid. The public API
            # places it correctly.
            doc.add_paragraph(text)
        elif after_index == -1:
            # Insert at beginning
            doc.paragraphs[0]._element.addprevious(new_para)
        else:
            paras[after_index]._element.addnext(new_para)

        inserted_at = 0 if total == 0 else after_index + 1

        # Apply style by setting it via python-docx after inserting
        # Re-read to find the new paragraph and apply style
        doc.save(str(path))

        # Re-open and find the inserted paragraph to apply style
        doc2 = Document(str(path))
        target_idx = inserted_at
        if target_idx < len(doc2.paragraphs):
            try:
                doc2.paragraphs[target_idx].style = doc2.styles[style]  # type: ignore[reportAttributeAccessIssue]
            except KeyError:
                pass  # Style doesn't exist — use default
            doc2.save(str(path))

        if open_after:
            open_file(path)
        progress.append(ok(f"Inserted paragraph at index {inserted_at}", style))
        progress.append(notify_reload(str(path), "docx"))

        _log_receipt(
            file_path,
            "insert_paragraph",
            {"after_index": after_index, "text": text, "style": style},
            f"✔ Inserted paragraph at {inserted_at}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "insert_paragraph",
            "inserted_at_index": inserted_at,
            "text": text,
            "style": style,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "backup": backup,
            "hint": hint_for_error(e, path),
            "progress": progress,
            "token_estimate": 20,
        }


def delete_paragraph(
    file_path: str,
    paragraph_index: int = -1,
    match_text: str = "",
    open_after: bool = False,
) -> dict[str, Any]:
    """Delete a paragraph by index or matching text."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        doc = Document(str(path))
        paras = doc.paragraphs
        total = len(paras)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        # Both selectors are optional in the schema, so the call the schema
        # documents -- file_path and nothing else -- reached the range check
        # below carrying the unset default and answered "paragraph_index -1 out
        # of range (0-6)": an index the caller never wrote, under a hint
        # offering to count the paragraphs. Nothing said the tool needs to be
        # told which paragraph to delete.
        if paragraph_index < 0 and not match_text:
            progress.append(fail("No paragraph chosen"))
            return {
                "success": False,
                "error": "Nothing to delete: give paragraph_index or match_text",
                "hint": (
                    "delete_paragraph(file_path, paragraph_index=3) deletes by position; "
                    "delete_paragraph(file_path, match_text='...') deletes the first paragraph "
                    "containing that text. Use get_document_outline to see the paragraphs."
                ),
                "progress": progress,
                "token_estimate": 30,
            }

        # Find target index
        target_idx = paragraph_index
        if match_text:
            for i, p in enumerate(paras):
                if match_text in p.text:
                    target_idx = i
                    break
            else:
                progress.append(fail(f"'{match_text}' not found"))
                return {
                    "success": False,
                    "error": "match_text not found in document",
                    "hint": "Use search_paragraphs to verify exact text.",
                    "progress": progress,
                    "token_estimate": 15,
                }

        if target_idx < 0 or target_idx >= total:
            progress.append(fail(f"Index {target_idx} out of range"))
            return {
                "success": False,
                "error": f"paragraph_index {target_idx} out of range {index_range(total, 'paragraphs')}",
                "hint": "Use read_document to get current paragraph count.",
                "progress": progress,
                "token_estimate": 15,
            }

        deleted_text = paras[target_idx].text[:80]
        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        # Remove the paragraph element from XML
        para_elem = paras[target_idx]._element
        para_elem.getparent().remove(para_elem)

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(ok(f"Deleted paragraph {target_idx}", deleted_text[:40]))
        progress.append(notify_reload(str(path), "docx"))

        _log_receipt(
            file_path,
            "delete_paragraph",
            {"paragraph_index": target_idx, "match_text": match_text},
            f"✔ Deleted paragraph {target_idx}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "delete_paragraph",
            "deleted_index": target_idx,
            "deleted_text": deleted_text,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "backup": backup,
            "hint": hint_for_error(e, path),
            "progress": progress,
            "token_estimate": 20,
        }


def append_text(file_path: str, text: str, style: str = "Body Text", open_after: bool = False) -> dict[str, Any]:
    """Append a new paragraph at document end."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        para_count = len(doc.paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{para_count} paragraphs"))

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        new_para = doc.add_paragraph(text)
        try:
            new_para.style = doc.styles[style]  # type: ignore[reportAttributeAccessIssue]
        except KeyError:
            pass  # Default style

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(ok(f"Appended paragraph at index {para_count}", style))
        progress.append(notify_reload(str(path), "docx"))

        # insert_paragraph, delete_paragraph and replace_text all log; this one
        # did not, so a document's receipt held every edit except the appends --
        # and get_history and read_receipt both read that log.
        _log_receipt(
            file_path,
            "append_text",
            {"text": text, "style": style},
            f"✔ Appended paragraph at index {para_count}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "append_text",
            "appended_at_index": para_count,
            "text": text,
            "style": style,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "backup": backup,
            "hint": hint_for_error(e, path),
            "progress": progress,
            "token_estimate": 20,
        }
