"""DOCX Basic helpers — surgical address tools, version/receipt tools, private helpers."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from shared.address_resolver import AddressError, resolve_docx_address
from shared.doc_diff import diff_docx
from shared.file_utils import hint_for_message, resolve_path
from shared.live_edit import notify_reload
from shared.progress import describe_error, fail, ok, undo, warn
from shared.receipt import append_receipt, read_receipt_log
from shared.version_control import get_history, restore, snapshot

logger = logging.getLogger(__name__)


# ─── private helpers ──────────────────────────────────────────────────────────


def _not_found(path: Path, progress: list) -> dict[str, Any]:
    progress.append(fail("File not found", str(path.name)))
    return {
        "success": False,
        "error": f"File not found: {path}",
        "hint": (
            "Check that file_path is absolute and the file exists. "
            "On macOS/Linux: /Users/you/doc.docx or ~/Documents/doc.docx"
        ),
        "progress": progress,
        "token_estimate": 20,
    }


def _wrong_type(path: Path, expected: str, progress: list) -> dict[str, Any]:
    progress.append(fail(f"Wrong file type: {path.suffix}", f"Expected {expected}"))
    return {
        "success": False,
        "error": f"Expected {expected} file, got {path.suffix}",
        "hint": f"Use the correct server for {path.suffix} files.",
        "progress": progress,
        "token_estimate": 15,
    }


def _error(msg: str, hint: str, progress: list) -> dict[str, Any]:
    return {
        "success": False,
        "error": describe_error(msg),
        "hint": hint_for_message(msg, hint),
        "progress": progress,
        "token_estimate": 15,
    }


def _log_receipt(
    file_path: str,
    tool: str,
    args: dict,
    result: str,
    backup: str | None,
    success: bool,
) -> None:
    try:
        append_receipt(file_path, tool, "docx-basic", args, result, backup, success)
    except Exception:
        pass


# ─── surgical address-based write tools ───────────────────────────────────────


def replace_at(file_path: str, address: str, new_text: str) -> dict[str, Any]:
    """Replace content at exact address. Preserves all run formatting."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        doc = Document(str(path))
        node = resolve_docx_address(doc, address)
        old_text = node.paragraph.text if node.paragraph else ""

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        if node.paragraph:
            # Run-level replacement preserving formatting
            for run in node.paragraph.runs:
                if run.text:
                    run.text = ""
            if node.paragraph.runs:
                node.paragraph.runs[0].text = new_text
            else:
                node.paragraph.add_run(new_text)

        doc.save(str(path))
        progress.append(ok(f"Replaced at {address}"))
        progress.append(notify_reload(str(path), "docx"))

        return {
            "success": True,
            "op": "replace_at",
            "address": address,
            "old_text": old_text,
            "new_text": new_text,
            "backup": backup,
            "tokens_read": 0,
            "tokens_written": len(new_text) // 4,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except AddressError as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "hint": "Use get_document_index to find valid addresses.",
            "progress": progress,
            "token_estimate": 20,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "backup": backup,
            "hint": "Use restore_version to undo.",
            "progress": progress,
            "token_estimate": 20,
        }


def insert_at(file_path: str, address: str, new_text: str, style: str = "Body Text") -> dict[str, Any]:
    """Insert new paragraph after address. Snapshot taken before write."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.oxml import OxmlElement  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        doc = Document(str(path))
        node = resolve_docx_address(doc, address)
        progress.append(ok(f"Resolved address {address}"))

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        new_para = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_r.append(new_t)
        new_para.append(new_r)

        if node.paragraph:
            node.paragraph._element.addnext(new_para)

        doc.save(str(path))
        progress.append(ok(f"Inserted paragraph after {address}", style))
        progress.append(notify_reload(str(path), "docx"))

        return {
            "success": True,
            "op": "insert_at",
            "after_address": address,
            "new_text": new_text,
            "style": style,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except AddressError as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "hint": "Use get_document_index to find valid addresses.",
            "progress": progress,
            "token_estimate": 20,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "backup": backup,
            "hint": "Use restore_version to undo.",
            "progress": progress,
            "token_estimate": 20,
        }


def delete_at(file_path: str, address: str) -> dict[str, Any]:
    """Delete paragraph or section at address."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)

        doc = Document(str(path))
        node = resolve_docx_address(doc, address)
        deleted_text = node.paragraph.text[:80] if node.paragraph else ""
        progress.append(ok(f"Resolved address {address}"))

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        if node.paragraph:
            elem = node.paragraph._element
            elem.getparent().remove(elem)

        doc.save(str(path))
        progress.append(ok(f"Deleted paragraph at {address}"))
        progress.append(notify_reload(str(path), "docx"))

        return {
            "success": True,
            "op": "delete_at",
            "address": address,
            "deleted_text": deleted_text,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except AddressError as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "hint": "Use get_document_index to find valid addresses.",
            "progress": progress,
            "token_estimate": 20,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return {
            "success": False,
            "error": str(e),
            "backup": backup,
            "hint": "Use restore_version to undo.",
            "progress": progress,
            "token_estimate": 20,
        }


# ─── version / history / receipt tools ────────────────────────────────────────


def get_history_tool(file_path: str) -> dict[str, Any]:
    """Return version snapshot history for a document."""
    progress: list[dict[str, Any]] = []
    try:
        path = resolve_path(file_path)
        history = get_history(str(path))
        progress.append(ok(f"Read history for {path.name}", f"{len(history)} snapshots"))
        return {
            "success": True,
            "file": str(path),
            "history": history,
            "progress": progress,
            "token_estimate": len(str(history)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check that file_path is correct.", progress)


def restore_version(file_path: str, timestamp: str, create_branch: str = "") -> dict[str, Any]:
    """Restore a document to a previous snapshot version."""
    progress: list[dict[str, Any]] = []
    try:
        path = resolve_path(file_path)

        if create_branch:
            from shared import gitops

            repo_dir = str(path.parent)
            if gitops.is_git_repo(str(path)):
                branch_ok = gitops.create_branch(repo_dir, create_branch)
                if branch_ok:
                    progress.append(ok(f"Created branch '{create_branch}'"))
                else:
                    progress.append(warn(f"Could not create branch '{create_branch}'"))

        restored = restore(str(path), timestamp)
        if not restored:
            progress.append(fail(f"Snapshot '{timestamp}' not found"))
            return {
                "success": False,
                "error": f"Snapshot with timestamp '{timestamp}' not found.",
                "hint": "Use get_history to see available timestamps.",
                "progress": progress,
                "token_estimate": 15,
            }

        progress.append(undo("Restored from snapshot", timestamp))
        progress.append(notify_reload(str(path), "docx"))

        # A restore replaces the document's entire contents, which makes it the
        # single event most worth being able to look up later -- and it was the
        # one write on this server that recorded nothing. A reader asking
        # read_receipt "what happened to this file?" saw the edits and no sign
        # that any of them had since been rolled back.
        _log_receipt(
            file_path,
            "restore_version",
            {"timestamp": timestamp},
            f"✔ Restored from snapshot {timestamp}",
            None,
            True,
        )

        return {
            "success": True,
            "op": "restore_version",
            "file": str(path),
            "restored_timestamp": timestamp,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path and timestamp.", progress)


def diff_versions(file_path: str, timestamp_a: str, timestamp_b: str = "current") -> dict[str, Any]:
    """Compare two versions of a document."""
    progress: list[dict[str, Any]] = []
    try:
        path = resolve_path(file_path)
        history = get_history(str(path))
        timestamps = {e["timestamp"]: e["backup_path"] for e in history}

        if timestamp_a not in timestamps:
            return {
                "success": False,
                "error": f"Snapshot '{timestamp_a}' not found.",
                "hint": "Use get_history to see available timestamps.",
                "progress": [fail(f"Snapshot '{timestamp_a}' not found")],
                "token_estimate": 15,
            }

        path_a = timestamps[timestamp_a]
        path_b = str(path) if timestamp_b == "current" else timestamps.get(timestamp_b, "")

        if not path_b:
            return {
                "success": False,
                "error": f"Snapshot '{timestamp_b}' not found.",
                "hint": "Use get_history or pass 'current'.",
                "progress": [fail(f"Snapshot '{timestamp_b}' not found")],
                "token_estimate": 15,
            }

        progress.append(ok("Comparing versions"))
        result = diff_docx(path_a, path_b)
        if result.get("success"):
            progress.append(ok("Diff complete", f"{result.get('change_count', 0)} changes"))

        result["progress"] = progress
        result["token_estimate"] = len(str(result)) // 4
        return result

    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path and timestamps.", progress)


def read_receipt_tool(file_path: str, last_n: int = 10) -> dict[str, Any]:
    """Show recent tool operations on this file."""
    progress: list[dict[str, Any]] = []
    try:
        path = resolve_path(file_path)
        entries = read_receipt_log(str(path), last_n)
        progress.append(ok(f"Read receipt log for {path.name}", f"{len(entries)} entries"))
        return {
            "success": True,
            "file": path.name,
            "entries": entries,
            "progress": progress,
            "token_estimate": len(str(entries)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check that file_path is correct.", progress)
