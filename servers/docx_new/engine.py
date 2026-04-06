"""DOCX New engine — create Word documents from scratch, no MCP imports."""
from __future__ import annotations

import logging
import sys
from pathlib import Path
from typing import Any

# Allow imports from repo root (shared/)
_ROOT = Path(__file__).resolve().parents[2]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from shared.platform_utils import open_file
from shared.progress import fail, info, ok, warn

logging.basicConfig(stream=sys.stderr, level=logging.WARNING)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _token_estimate(obj: Any) -> int:
    return len(str(obj)) // 4


def _ensure_parent(path: Path) -> None:
    """Create parent directories if they do not exist."""
    path.parent.mkdir(parents=True, exist_ok=True)


def _open_if_requested(path: Path, open_after: bool, progress: list[dict[str, Any]]) -> None:
    if open_after:
        open_file(path)
        progress.append(ok(f"Opened {path.name} in default app"))


def _err(progress: list[dict[str, Any]], error: str, hint: str) -> dict[str, Any]:
    return {
        "success": False,
        "error": error,
        "hint": hint,
        "progress": progress,
        "token_estimate": _token_estimate(progress),
    }


# ---------------------------------------------------------------------------
# Public engine functions
# ---------------------------------------------------------------------------

def create_document(output_path: str, open_after: bool = True) -> dict[str, Any]:
    """Create a blank Word document and save to output_path."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = Path(output_path).resolve()
        _ensure_parent(path)
        progress.append(info(f"Creating blank document", path.name))

        doc = Document()
        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", str(path.parent)))

        _open_if_requested(path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "create_document",
            "output": str(path),
            "output_name": path.name,
            "progress": progress,
            "token_estimate": _token_estimate(progress),
        }
        return result
    except Exception as exc:
        logger.warning("create_document failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check that output_path is a valid writable path ending in .docx.",
        )


def create_from_text(
    output_path: str,
    paragraphs: list[dict[str, Any]],
    open_after: bool = True,
) -> dict[str, Any]:
    """Create a Word document from a list of {text, style} paragraph dicts."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        if not isinstance(paragraphs, list):
            progress.append(fail("paragraphs must be a list of dicts"))
            return _err(
                progress,
                "paragraphs must be a list",
                'Pass a list like [{"text": "Hello", "style": "Normal"}].',
            )

        path = Path(output_path).resolve()
        _ensure_parent(path)
        progress.append(info(f"Creating document from {len(paragraphs)} paragraphs", path.name))

        doc = Document()
        added = 0
        for item in paragraphs:
            text = item.get("text", "") if isinstance(item, dict) else str(item)
            style = item.get("style", "Normal") if isinstance(item, dict) else "Normal"
            if not style:
                style = "Normal"
            try:
                doc.add_paragraph(text, style=style)
            except KeyError:
                # Style not found — fall back to Normal
                progress.append(warn(f"Style '{style}' not found, using Normal"))
                doc.add_paragraph(text, style="Normal")
            added += 1

        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", f"{added} paragraphs written"))

        _open_if_requested(path, open_after, progress)

        return {
            "success": True,
            "op": "create_from_text",
            "output": str(path),
            "output_name": path.name,
            "paragraph_count": added,
            "progress": progress,
            "token_estimate": _token_estimate(progress),
        }
    except Exception as exc:
        logger.warning("create_from_text failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Ensure paragraphs is a list of dicts with 'text' and optional 'style' keys.",
        )


def create_from_sections(
    output_path: str,
    title: str,
    sections: list[dict[str, Any]],
    open_after: bool = True,
) -> dict[str, Any]:
    """Create a structured document from a title and list of {heading, body} sections."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        if not isinstance(sections, list):
            progress.append(fail("sections must be a list of dicts"))
            return _err(
                progress,
                "sections must be a list",
                'Pass a list like [{"heading": "Intro", "body": "Text here"}].',
            )

        path = Path(output_path).resolve()
        _ensure_parent(path)
        progress.append(info(f"Creating structured document", path.name))

        doc = Document()

        # Title as Heading 1
        doc.add_paragraph(title, style="Heading 1")
        progress.append(ok(f"Added title: {title[:40]}"))

        section_count = 0
        for sec in sections:
            heading = sec.get("heading", "") if isinstance(sec, dict) else ""
            body = sec.get("body", "") if isinstance(sec, dict) else str(sec)
            if heading:
                doc.add_paragraph(heading, style="Heading 2")
            if body:
                doc.add_paragraph(body, style="Normal")
            section_count += 1

        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", f"{section_count} sections written"))

        _open_if_requested(path, open_after, progress)

        return {
            "success": True,
            "op": "create_from_sections",
            "output": str(path),
            "output_name": path.name,
            "section_count": section_count,
            "progress": progress,
            "token_estimate": _token_estimate(progress),
        }
    except Exception as exc:
        logger.warning("create_from_sections failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Ensure sections is a list of dicts with 'heading' and 'body' keys.",
        )


def create_from_template(
    template_path: str,
    output_path: str,
    substitutions: dict[str, str],
    open_after: bool = True,
) -> dict[str, Any]:
    """Copy a .docx template, apply text substitutions, save to output_path."""
    progress: list[dict[str, Any]] = []
    try:
        import shutil

        import docxedit  # type: ignore[import-untyped]
        from docx import Document  # type: ignore[import-untyped]

        tpl_path = Path(template_path).resolve()
        if not tpl_path.exists():
            progress.append(fail(f"Template not found", str(tpl_path)))
            return _err(
                progress,
                f"File not found: {template_path}",
                "Check that template_path is an absolute path to an existing .docx file.",
            )
        if tpl_path.suffix.lower() != ".docx":
            progress.append(fail(f"Template is not a .docx file", tpl_path.suffix))
            return _err(
                progress,
                f"Expected .docx file, got {tpl_path.suffix}",
                "Use the correct server for this file type.",
            )

        out_path = Path(output_path).resolve()
        _ensure_parent(out_path)

        progress.append(ok(f"Opened template {tpl_path.name}"))

        # Copy template to output location first
        shutil.copy2(str(tpl_path), str(out_path))

        # Open the copy and apply substitutions using run-level editing
        doc = Document(str(out_path))

        if not isinstance(substitutions, dict):
            progress.append(fail("substitutions must be a dict"))
            return _err(
                progress,
                "substitutions must be a dict",
                'Pass a dict like {"PLACEHOLDER": "replacement value"}.',
            )

        applied = 0
        for old_str, new_str in substitutions.items():
            old_s = str(old_str)
            new_s = str(new_str)
            # docxedit.replace_string operates on runs — never assigns .text directly
            count = docxedit.replace_string(doc, old_s, new_s)
            if count:
                progress.append(
                    ok(f"Replaced '{old_s}' → '{new_s}'", f"{count} occurrence{'s' if count != 1 else ''}")
                )
                applied += count
            else:
                progress.append(warn(f"Placeholder '{old_s}' not found in template"))

        doc.save(str(out_path))
        progress.append(ok(f"Saved {out_path.name}", f"{applied} total substitutions"))

        _open_if_requested(out_path, open_after, progress)

        return {
            "success": True,
            "op": "create_from_template",
            "output": str(out_path),
            "output_name": out_path.name,
            "substitutions_applied": applied,
            "progress": progress,
            "token_estimate": _token_estimate(progress),
        }
    except Exception as exc:
        logger.warning("create_from_template failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check that template_path points to a valid .docx file and substitutions is a dict.",
        )


def create_letter(
    output_path: str,
    from_name: str,
    to_name: str,
    subject: str,
    body: str,
    open_after: bool = True,
) -> dict[str, Any]:
    """Create a formatted business letter .docx."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt  # type: ignore[import-untyped]

        path = Path(output_path).resolve()
        _ensure_parent(path)
        progress.append(info(f"Creating business letter", path.name))

        doc = Document()

        # Sender name
        sender_para = doc.add_paragraph()
        sender_run = sender_para.add_run(from_name)
        sender_run.bold = True

        # Blank line
        doc.add_paragraph()

        # Recipient
        doc.add_paragraph(f"To: {to_name}")

        # Blank line
        doc.add_paragraph()

        # Subject line — bold
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run(f"Subject: {subject}")
        subject_run.bold = True

        # Blank line
        doc.add_paragraph()

        # Body paragraphs — split on newlines
        body_paras = body.split("\n")
        for line in body_paras:
            doc.add_paragraph(line)

        # Blank line before closing
        doc.add_paragraph()

        # Closing with sender name
        closing_para = doc.add_paragraph()
        closing_para.add_run("Sincerely,")
        doc.add_paragraph()
        sign_para = doc.add_paragraph()
        sign_run = sign_para.add_run(from_name)
        sign_run.bold = True

        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", f"From: {from_name} | To: {to_name}"))

        _open_if_requested(path, open_after, progress)

        return {
            "success": True,
            "op": "create_letter",
            "output": str(path),
            "output_name": path.name,
            "from_name": from_name,
            "to_name": to_name,
            "subject": subject,
            "progress": progress,
            "token_estimate": _token_estimate(progress),
        }
    except Exception as exc:
        logger.warning("create_letter failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check that output_path is a valid writable path ending in .docx.",
        )
