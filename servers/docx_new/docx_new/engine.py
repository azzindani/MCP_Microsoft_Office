"""DOCX New engine — create Word documents from scratch, no MCP imports."""

from __future__ import annotations

import logging
import sys
from pathlib import Path
from typing import Any

# Allow imports from repo root (shared/)
_ROOT = Path(__file__).resolve().parents[2]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from shared.arg_alias import (  # noqa: E402
    ENTRY_HEADING_KEYS,
    ENTRY_TEXT_KEYS,
    entry_value,
    unnamed_entry_note,
)
from shared.file_utils import embed_content  # noqa: E402
from shared.platform_utils import open_file, resolve_output_path  # noqa: E402
from shared.progress import fail, info, ok, warn  # noqa: E402
from shared.template_fill import ordered_pairs, resolve_targets, sentinel_for  # noqa: E402

logging.basicConfig(stream=sys.stderr, level=logging.WARNING)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _token_estimate(obj: Any) -> int:
    return len(str(obj)) // 4


def _count_occurrences(doc: Any, needle: str) -> int:
    """Count paragraphs (incl. table cells) containing needle.

    docxedit.replace_string() has no return statement (always None), so
    callers must count occurrences themselves before/independent of calling it.
    """
    n = sum(1 for p in doc.paragraphs if needle in p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                n += sum(1 for p in cell.paragraphs if needle in p.text)
    return n


def _all_text(doc: Any) -> str:
    """Every paragraph the substitution pass can reach, as one string.

    Only used to decide what a key should match, so joining is enough -- a
    placeholder never spans two paragraphs.
    """
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _ensure_parent(path: Path) -> None:
    """Create parent directories if they do not exist."""
    path.parent.mkdir(parents=True, exist_ok=True)


def _open_if_requested(path: Path, open_after: bool, progress: list[dict[str, Any]]) -> None:
    if open_after:
        open_file(path)
        progress.append(ok(f"Opened {path.name} in default app"))


def _err(progress: list[dict[str, Any]], error: str, hint: str) -> dict[str, Any]:
    return {
        "success": False,
        "error": error,
        "hint": hint,
        "progress": progress,
        "token_estimate": _token_estimate(progress),
    }


# ---------------------------------------------------------------------------
# Public engine functions
# ---------------------------------------------------------------------------


def create_document(output_path: str, open_after: bool = True, return_content: bool = False) -> dict[str, Any]:
    """Create a blank Word document and save to output_path."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_output_path(output_path, "document.docx")
        _ensure_parent(path)
        progress.append(info("Creating blank document", path.name))

        doc = Document()
        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", str(path.parent)))

        _open_if_requested(path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "create_document",
            "output": str(path),
            "output_name": path.name,
            "progress": progress,
        }
        embed_content(result, path, return_content)
        result["token_estimate"] = _token_estimate(result)
        return result
    except Exception as exc:
        logger.warning("create_document failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check that output_path is a valid writable path ending in .docx.",
        )


def create_from_text(
    output_path: str,
    paragraphs: list[dict[str, Any]],
    open_after: bool = True,
    return_content: bool = False,
) -> dict[str, Any]:
    """Create a Word document from a list of {text, style} paragraph dicts."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        if not isinstance(paragraphs, list):
            progress.append(fail("paragraphs must be a list of dicts"))
            return _err(
                progress,
                "paragraphs must be a list",
                'Pass a list like [{"text": "Hello", "style": "Normal"}].',
            )

        path = resolve_output_path(output_path, "document.docx")
        _ensure_parent(path)
        progress.append(info(f"Creating document from {len(paragraphs)} paragraphs", path.name))

        # Every entry is inspected before the document is built: a paragraph
        # written under an unrecognised key used to become an empty paragraph
        # and still be counted, so [{"content": "hello"}] saved a .docx with no
        # text in it and reported "1 paragraphs written".
        unnamed = [
            note
            for i, item in enumerate(paragraphs)
            if (note := unnamed_entry_note(i, item, ENTRY_TEXT_KEYS, "Paragraph"))
        ]
        if unnamed and len(unnamed) == len(paragraphs):
            progress.append(fail("No paragraph carried any text"))
            return _err(
                progress,
                "; ".join(unnamed),
                'Pass a list like [{"text": "Hello", "style": "Normal"}]. Nothing was written.',
            )

        doc = Document()
        added = 0
        for i, item in enumerate(paragraphs):
            text = entry_value(item, ENTRY_TEXT_KEYS)
            style = item.get("style", "Normal") if isinstance(item, dict) else "Normal"
            if not style:
                style = "Normal"
            try:
                doc.add_paragraph(text, style=style)
            except KeyError:
                # Style not found — fall back to Normal
                progress.append(warn(f"Style '{style}' not found, using Normal"))
                doc.add_paragraph(text, style="Normal")
            added += 1

        for note in unnamed:
            progress.append(warn("Paragraph written empty", note))

        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", f"{added} paragraphs written"))

        _open_if_requested(path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "create_from_text",
            "output": str(path),
            "output_name": path.name,
            "paragraph_count": added,
            "progress": progress,
        }
        embed_content(result, path, return_content)
        result["token_estimate"] = _token_estimate(result)
        return result
    except Exception as exc:
        logger.warning("create_from_text failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Ensure paragraphs is a list of dicts with 'text' and optional 'style' keys.",
        )


def create_from_sections(
    output_path: str,
    title: str,
    sections: list[dict[str, Any]],
    open_after: bool = True,
    return_content: bool = False,
) -> dict[str, Any]:
    """Create a structured document from a title and list of {heading, body} sections."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        if not isinstance(sections, list):
            progress.append(fail("sections must be a list of dicts"))
            return _err(
                progress,
                "sections must be a list",
                'Pass a list like [{"heading": "Intro", "body": "Text here"}].',
            )

        path = resolve_output_path(output_path, "document.docx")
        _ensure_parent(path)
        progress.append(info("Creating structured document", path.name))

        doc = Document()

        # Title as Heading 1
        doc.add_paragraph(title, style="Heading 1")
        progress.append(ok(f"Added title: {title[:40]}"))

        # `header` for `heading` used to drop the heading and keep the body, so
        # the document came back a level flatter than it was asked for with
        # nothing said about it. A section carrying neither is a typo, not a
        # request for a blank one.
        section_count = 0
        for i, sec in enumerate(sections):
            heading = entry_value(sec, ENTRY_HEADING_KEYS)
            body = entry_value(sec, ENTRY_TEXT_KEYS) if isinstance(sec, dict) else str(sec)
            if not heading and not body:
                progress.append(
                    warn(
                        f"Section {i} written empty",
                        unnamed_entry_note(i, sec, ENTRY_HEADING_KEYS, "Section"),
                    )
                )
            if heading:
                doc.add_paragraph(heading, style="Heading 2")
            if body:
                doc.add_paragraph(body, style="Normal")
            section_count += 1

        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", f"{section_count} sections written"))

        _open_if_requested(path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "create_from_sections",
            "output": str(path),
            "output_name": path.name,
            "section_count": section_count,
            "progress": progress,
        }
        embed_content(result, path, return_content)
        result["token_estimate"] = _token_estimate(result)
        return result
    except Exception as exc:
        logger.warning("create_from_sections failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Ensure sections is a list of dicts with 'heading' and 'body' keys.",
        )


# The accent a document gets when the caller names none. Word's own "Dark
# Blue, Text 2, Darker 25%" -- neutral enough for any subject and dark enough
# to carry white text in a table header.
_DEFAULT_ACCENT = "1F3864"

# How far toward white the derived fills sit. Banded rows have to be barely
# there; a callout can be a shade stronger without competing with body text.
_BAND_TINT = 0.92
_CALLOUT_TINT = 0.86

# The severity vocabulary a risks table is scored on. Three levels, the same
# three the data side uses in `insights` and the quality alerts, so a document
# assembled from those findings does not introduce a fourth scale.
RISK_LEVELS: tuple[str, ...] = ("high", "medium", "low")

# The one list. `BLOCK_KINDS` is derived from it below rather than written
# beside it, because a second list is a list that can disagree -- and the whole
# point of the census in `test_a_kind_the_tool_lists_and_cannot_draw.py` is that
# the lists describing a capability keep drifting apart.
#
# It also exists because twelve kinds do not fit in an 80-character tool
# description, and two of this review's asks met head-on: "no dead ops" wants
# every kind named where an agent can read it, and the docstring cap wants the
# description short. The review answered that itself, in its scaling section:
# *"Never list 1000: expose `search/get_capabilities/get_state`, sparse outline
# then detail"*. So the description points at `list_block_kinds`, and this is
# what that returns.
BLOCK_KIND_SPECS: dict[str, dict[str, Any]] = {
    "heading": {
        "does": "A section heading, coloured with the accent.",
        "keys": {"text": "required", "level": "1-6, default 2"},
        "example": {"kind": "heading", "text": "Findings", "level": 2},
    },
    "text": {
        "does": "One body paragraph.",
        "keys": {"text": "required"},
        "example": {"kind": "text", "text": "Charge-off rate is 13.82%."},
    },
    "bullets": {
        "does": "A bulleted or numbered list.",
        "keys": {"items": "required, list of strings", "numbered": "bool, default false"},
        "example": {"kind": "bullets", "items": ["First", "Second"]},
    },
    "table": {
        "does": "A table with a shaded header and banded rows.",
        "keys": {"header": "list", "rows": "list of lists", "widths": "optional cm per column"},
        "example": {"kind": "table", "header": ["Metric", "Value"], "rows": [["Rows", "38,576"]]},
    },
    "kpi": {
        "does": "Figures across the page, value large, label beneath.",
        "keys": {"items": "required, [{value, label}]"},
        "example": {"kind": "kpi", "items": [{"value": "13.82%", "label": "charge-off rate"}]},
    },
    "callout": {
        "does": "A tinted box for the one message that must land.",
        "keys": {"text": "required", "title": "optional"},
        "example": {"kind": "callout", "title": "Risk", "text": "The model may be leaking."},
    },
    "image": {
        "does": "A picture, from a local path or an http(s) URL.",
        "keys": {"path": "or url", "width_in": "inches, default 6.0", "caption": "optional"},
        "example": {"kind": "image", "path": "/data/chart.png", "width_in": 6.0},
    },
    "links": {
        "does": "Clickable links out to charts or other artifacts.",
        "keys": {"items": "required, [{label, url, note}]", "title": "optional"},
        "example": {"kind": "links", "items": [{"label": "Dashboard", "url": "https://…/dash.html"}]},
    },
    "risks": {
        "does": f"A risk table, severity coloured. Levels: {', '.join(RISK_LEVELS)}.",
        "keys": {"items": "required, [{risk, level, impact, mitigation, owner}]"},
        "example": {"kind": "risks", "items": [{"risk": "Leakage", "level": "high"}]},
    },
    "checklist": {
        "does": "Ticked and unticked actions.",
        "keys": {"items": "required, [{text, done}] or plain strings"},
        "example": {"kind": "checklist", "items": [{"text": "Drop id", "done": True}]},
    },
    "rule": {
        "does": "A horizontal rule.",
        "keys": {"width_pt": "default 1.0"},
        "example": {"kind": "rule"},
    },
    "pagebreak": {
        "does": "A page break.",
        "keys": {},
        "example": {"kind": "pagebreak"},
    },
}

BLOCK_KINDS: tuple[str, ...] = tuple(BLOCK_KIND_SPECS)

# Red / amber / green, dark enough to carry white text.
_RISK_FILLS: dict[str, str] = {"high": "C0392B", "medium": "D68910", "low": "1E8449"}

# Word's own check-box glyphs render in every default font; a font-dependent
# glyph would come out as a box on a machine without it, which is the one
# outcome a checklist cannot survive.
_CHECK_DONE = "☒"
_CHECK_OPEN = "☐"

# What python-docx can actually place. A user review put the gap plainly:
# "create_from_sections (docx) cannot embed images/charts. The board paper
# references charts that live in separate HTML files... Room for improvement:
# an `image` block type accepting a `data/` path or URL."
_IMAGE_EXTS: tuple[str, ...] = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".emf", ".wmf")

# The one a caller most often has and most often tries: every chart this fleet
# writes is a self-contained .html. Refusing it by name, with the tool that
# converts it, is worth more than "unsupported image format".
_HTML_EXTS: tuple[str, ...] = (".html", ".htm")

# Wide enough to read a chart, narrow enough to sit inside default margins.
_DEFAULT_IMAGE_WIDTH_IN = 6.0


def _add_image(doc: Any, block: dict[str, Any], docx_style: Any, hue: str) -> str:
    """Place one image block. Returns "" on success, or why it wrote nothing.

    Accepts a local path or an http(s) URL, because the review's case was a
    board paper referencing charts that this fleet had already written to the
    shared output directory -- and the same directory is served over HTTP, so a
    caller holds one or the other depending on where it is running.

    Three refusals are worth more than a generic failure, and each names the way
    out:

    * **A .html chart.** Every chart these servers produce is a self-contained
      HTML page, so it is the file a caller has in hand and the first thing they
      will pass. python-docx cannot place it. Naming the converter is the whole
      value of the message.
    * **A missing file**, which is usually a path from the wrong side of the
      container boundary.
    * **A URL that did not fetch.** Reported with the status, because "the
      document has no image" and "the file server refused" are different
      problems.
    """
    from io import BytesIO

    from docx.shared import Inches  # type: ignore[import-untyped]

    source = ""
    for key in ("path", "src", "source", "url", "file", "file_path", "image"):
        value = block.get(key)
        if isinstance(value, str) and value.strip():
            source = value.strip()
            break
    if not source:
        return "an image block with no path or url"

    suffix = Path(source.split("?")[0]).suffix.lower()
    if suffix in _HTML_EXTS:
        return (
            f"{source} is an HTML page, which cannot be placed in a .docx. "
            "Charts from this fleet are self-contained HTML; render one to PNG first "
            "(or ask the chart tool for an image) and pass that path"
        )

    stream: Any
    if source.lower().startswith(("http://", "https://")):
        try:
            import urllib.request

            with urllib.request.urlopen(source, timeout=20) as resp:  # noqa: S310
                if getattr(resp, "status", 200) >= 400:
                    return f"{source} returned HTTP {resp.status}"
                stream = BytesIO(resp.read())
        except Exception as exc:
            return f"{source} could not be fetched ({type(exc).__name__}: {exc})"
    else:
        img_path = Path(source)
        if not img_path.is_absolute():
            img_path = (Path.cwd() / img_path).resolve()
        if not img_path.exists():
            return f"{source} does not exist (paths are resolved on the server, not the caller)"
        if suffix and suffix not in _IMAGE_EXTS:
            return f"{source} is a {suffix} file; python-docx places {', '.join(_IMAGE_EXTS)}"
        stream = str(img_path)

    try:
        width_in = float(block.get("width_in") or block.get("width") or _DEFAULT_IMAGE_WIDTH_IN)
    except TypeError, ValueError:
        width_in = _DEFAULT_IMAGE_WIDTH_IN
    try:
        doc.add_picture(stream, width=Inches(max(0.5, min(width_in, 9.0))))
    except Exception as exc:
        return f"{source} is not a readable image ({type(exc).__name__}: {exc})"

    caption = entry_value(block, ENTRY_TEXT_KEYS) or str(block.get("caption") or "")
    if caption:
        para = doc.add_paragraph(caption, style="Normal")
        docx_style.style_runs(para, color=hue, italic="true", font_size=9.0)
    return ""


def _block_kind(block: dict[str, Any]) -> str:
    for key in ("kind", "type", "block", "block_type"):
        value = block.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip().lower()
    return ""


def create_from_blocks(
    output_path: str,
    title: str,
    blocks: list[dict[str, Any]],
    accent: str = "",
    open_after: bool = True,
    return_content: bool = False,
    font: str = "",
    heading_font: str = "",
) -> dict[str, Any]:
    """Create a .docx from typed blocks — headings, bullets, tables, KPIs.

    `create_from_sections` takes {heading, body} strings and can only ever
    produce one Heading and one Normal paragraph per section. That is not a
    limitation you notice until someone reads the output: handed a data
    summary it wrote ten headings over ten 180-word paragraphs, no bullets and
    no tables, and the reply that came back was "so many text and numbers in
    single paragraphs, not readable for my director".

    The primitives to fix that by hand all exist, but only one paragraph at a
    time and addressed by index, so a four-page brief is ~150 sequential calls
    against indices that shift under every insert. A model that wants a
    readable document reaches for python-docx instead, and did.

    A block is {"kind": ..., ...}. Every kind below is one call:

        heading   {kind, text, level}          level 1-6, accent coloured
        text      {kind, text}
        bullets   {kind, items[], numbered}    bulleted unless numbered
        table     {kind, header[], rows[][]}   shaded header, banded body
        kpi       {kind, items[{value,label}]} figures across the page
        callout   {kind, text, title}          tinted box for the one message
        image     {kind, path|url, width_in}   a picture, from a file or a URL
        links     {kind, items[{label,url}]}   clickable links to the charts
        risks     {kind, items[{risk,level}]}  risk table, severity coloured
        checklist {kind, items[{text,done}]}   ticked and unticked actions
        rule      {kind}                       horizontal rule
        pagebreak {kind}

    `accent`, `font` and `heading_font` are the brand tokens: one colour and
    two typefaces, applied to every block rather than to each one separately.
    A review asked for "style/brand tokens" and this is the whole of them --
    anything more would be a template, which `create_from_template` already is.
    """
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.enum.text import WD_BREAK  # type: ignore[import-untyped]
        from docx.shared import Pt  # type: ignore[import-untyped]

        from shared import docx_style

        if not isinstance(blocks, list):
            progress.append(fail("blocks must be a list of dicts"))
            return _err(
                progress,
                "blocks must be a list",
                'Pass a list like [{"kind": "bullets", "items": ["First", "Second"]}].',
            )

        try:
            hue = docx_style.parse_hex(accent, "accent") if accent else _DEFAULT_ACCENT
        except docx_style.DocxStyleError as exc:
            progress.append(fail(str(exc)))
            return _err(progress, str(exc), "Accent is a 6-digit hex colour, e.g. '0B1D3A'.")
        band = docx_style.tint(hue, _BAND_TINT)
        callout_fill = docx_style.tint(hue, _CALLOUT_TINT)

        path = resolve_output_path(output_path, "document.docx")
        _ensure_parent(path)
        progress.append(info("Creating block document", path.name))

        doc = Document()
        if title:
            heading = doc.add_paragraph(title, style="Heading 1")
            docx_style.style_runs(heading, color=hue)
            docx_style.set_paragraph_rule(heading, color=hue, width_pt=1.0)
            docx_style.set_spacing(heading, space_after=10)

        counts: dict[str, int] = {}
        unknown: list[str] = []
        tables_made = 0
        images_made = 0
        links_made = 0

        for index, block in enumerate(blocks):
            if not isinstance(block, dict):
                unknown.append(f"block {index} is {type(block).__name__}, not a dict")
                continue
            kind = _block_kind(block)
            if kind not in BLOCK_KINDS:
                unknown.append(f"block {index} has kind={kind or 'missing'!r}")
                continue

            if kind == "heading":
                level = max(1, min(int(block.get("level", 2) or 2), 6))
                text = entry_value(block, ENTRY_TEXT_KEYS) or entry_value(block, ENTRY_HEADING_KEYS)
                para = doc.add_paragraph(text, style=f"Heading {level}")
                docx_style.style_runs(para, color=hue)

            elif kind == "text":
                doc.add_paragraph(entry_value(block, ENTRY_TEXT_KEYS), style="Normal")

            elif kind == "bullets":
                items = block.get("items") or block.get("bullets") or []
                # "List Bullet"/"List Number" are built-in Word styles; a
                # caller's template can lack them, and a missing style must not
                # cost the content.
                style = "List Number" if block.get("numbered") else "List Bullet"
                for item in items:
                    text = entry_value(item, ENTRY_TEXT_KEYS) if isinstance(item, dict) else str(item)
                    try:
                        doc.add_paragraph(text, style=style)
                    except KeyError:
                        doc.add_paragraph(f"• {text}", style="Normal")

            elif kind == "table":
                header = block.get("header") or block.get("headers") or []
                rows = block.get("rows") or []
                if not header and not rows:
                    unknown.append(f"block {index} is a table with no header and no rows")
                    continue
                _add_block_table(doc, header, rows, hue, band, block.get("widths") or [], docx_style)
                _spacer(doc, docx_style)
                tables_made += 1

            elif kind == "kpi":
                items = block.get("items") or []
                if not items:
                    unknown.append(f"block {index} is a kpi row with no items")
                    continue
                _add_kpi_row(doc, items, hue, docx_style)
                _spacer(doc, docx_style)
                tables_made += 1

            elif kind == "callout":
                _add_callout(doc, block, hue, callout_fill, docx_style)
                _spacer(doc, docx_style)
                tables_made += 1

            elif kind == "image":
                note = _add_image(doc, block, docx_style, hue)
                if note:
                    unknown.append(f"block {index}: {note}")
                    continue
                images_made += 1

            elif kind == "links":
                items = block.get("items") or block.get("links") or []
                note = _add_links(doc, block, items, hue, docx_style)
                if note:
                    unknown.append(f"block {index}: {note}")
                    continue
                links_made += sum(1 for i in items if isinstance(i, dict) and i.get("url"))

            elif kind == "risks":
                items = block.get("items") or block.get("risks") or []
                if not items:
                    unknown.append(f"block {index} is a risks table with no items")
                    continue
                bad_level = _add_risks(doc, items, hue, band, docx_style)
                _spacer(doc, docx_style)
                tables_made += 1
                if bad_level:
                    unknown.append(f"block {index}: {bad_level}")

            elif kind == "checklist":
                items = block.get("items") or block.get("checklist") or []
                if not items:
                    unknown.append(f"block {index} is a checklist with no items")
                    continue
                _add_checklist(doc, items, hue, docx_style)

            elif kind == "rule":
                para = doc.add_paragraph("")
                docx_style.set_paragraph_rule(para, color=hue, width_pt=float(block.get("width_pt", 1.0) or 1.0))

            elif kind == "pagebreak":
                doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)

            counts[kind] = counts.get(kind, 0) + 1

        # Body text at 11pt with a little air, so the default is a document
        # somebody can read rather than one that merely opens.
        normal = doc.styles["Normal"]
        normal.font.size = Pt(11)  # type: ignore[reportAttributeAccessIssue]
        normal.paragraph_format.space_after = Pt(6)  # type: ignore[reportAttributeAccessIssue]
        fonts_applied = _apply_fonts(doc, font, heading_font)

        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", f"{sum(counts.values())} block(s), {tables_made} table(s)"))
        if unknown:
            progress.append(
                warn(
                    f"{len(unknown)} block(s) written nothing",
                    "; ".join(unknown) + f". Valid kinds: {', '.join(BLOCK_KINDS)}.",
                )
            )

        _open_if_requested(path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "create_from_blocks",
            "output": str(path),
            "output_name": path.name,
            "block_count": sum(counts.values()),
            "blocks_by_kind": counts,
            "images_embedded": images_made,
            "links_embedded": links_made,
            "skipped": unknown,
            "accent": f"#{hue}",
            "fonts": fonts_applied,
            "progress": progress,
        }
        embed_content(result, path, return_content)
        result["token_estimate"] = _token_estimate(result)
        return result
    except Exception as exc:
        logger.warning("create_from_blocks failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            f"Each block is a dict with a 'kind' of: {', '.join(BLOCK_KINDS)}.",
        )


def _spacer(doc: Any, docx_style: Any) -> None:
    """A short empty paragraph after every table-shaped block.

    Not only for air. Word merges two tables that sit next to each other in the
    body with nothing between them, so a kpi row followed by a table would come
    out as one seven-row grid -- the layout defect that is hardest to see in a
    success response and obvious the moment anyone opens the file.
    """
    docx_style.set_spacing(doc.add_paragraph(""), space_after=6)


def _add_block_table(
    doc: Any,
    header: list[Any],
    rows: list[list[Any]],
    hue: str,
    band: str,
    widths: list[Any],
    docx_style: Any,
) -> None:
    """A header-shaded, row-banded table. Borders come from 'Table Grid'."""
    body = [[str(cell) for cell in row] for row in rows]
    n_cols = max([len(header)] + [len(r) for r in body]) or 1
    n_rows = len(body) + (1 if header else 0)
    try:
        table = doc.add_table(rows=n_rows, cols=n_cols, style="Table Grid")
    except KeyError:
        table = doc.add_table(rows=n_rows, cols=n_cols)

    offset = 0
    if header:
        for col in range(n_cols):
            cell = table.rows[0].cells[col]
            cell.text = str(header[col]) if col < len(header) else ""
            docx_style.set_cell_fill(cell, hue)
            for para in cell.paragraphs:
                docx_style.style_runs(para, bold="true", color="FFFFFF")
        offset = 1

    for r, row in enumerate(body):
        for col in range(n_cols):
            cell = table.rows[r + offset].cells[col]
            cell.text = row[col] if col < len(row) else ""
        if r % 2 == 1:
            for col in range(n_cols):
                docx_style.set_cell_fill(table.rows[r + offset].cells[col], band)

    if widths:
        docx_style.set_column_widths(table, [float(w) for w in widths])


def _add_kpi_row(doc: Any, items: list[Any], hue: str, docx_style: Any) -> None:
    """Figures across the page: value large and coloured, label small beneath.

    A borderless two-row table, because Word has no other way to put four
    figures side by side that survives being edited afterwards.
    """
    from docx.shared import Pt  # type: ignore[import-untyped]

    table = doc.add_table(rows=2, cols=len(items))
    for column, item in enumerate(items):
        value = str(item.get("value", "")) if isinstance(item, dict) else str(item)
        label = str(item.get("label", "")) if isinstance(item, dict) else ""
        value_cell = table.rows[0].cells[column]
        value_cell.text = value
        for para in value_cell.paragraphs:
            docx_style.set_alignment(para, "center")
            docx_style.style_runs(para, font_size=20, bold="true", color=hue)
        label_cell = table.rows[1].cells[column]
        label_cell.text = label
        for para in label_cell.paragraphs:
            docx_style.set_alignment(para, "center")
            docx_style.style_runs(para, font_size=8, color="595959", all_caps="true")
            para.paragraph_format.space_after = Pt(10)


def list_block_kinds() -> dict[str, Any]:
    """Every block kind `create_from_blocks` accepts, with its keys and an example.

    The tool description cannot hold twelve kinds and stay inside the 80
    character cap the docstring census enforces, and an undiscoverable kind is a
    capability that does not exist. The review's own answer to that tension is
    the sparse-then-detail shape: a short description that names this, and this
    for the detail.

    Reads `BLOCK_KIND_SPECS`, which is also what `BLOCK_KINDS` is derived from,
    so this cannot fall out of step with what the validator accepts.
    """
    return {
        "success": True,
        "op": "list_block_kinds",
        "kinds": [
            {"kind": name, "does": spec["does"], "keys": spec["keys"], "example": spec["example"]}
            for name, spec in BLOCK_KIND_SPECS.items()
        ],
        "kind_count": len(BLOCK_KIND_SPECS),
        "risk_levels": list(RISK_LEVELS),
        "brand": {
            "accent": "6-digit hex, e.g. '0B1D3A'",
            "font": "body typeface, e.g. 'Georgia'",
            "heading_font": "heading typeface; defaults to font",
        },
        "token_estimate": 0,
    }


def _add_links(doc: Any, block: dict[str, Any], items: list[Any], hue: str, docx_style: Any) -> str:
    """Clickable links to the artifacts the paper talks about.

    A user review's board paper "references charts that live in separate HTML
    files", and the fix it asked for was `public_url` links to them. Real
    `w:hyperlink` elements, not blue text: text that looks like a link and does
    nothing when clicked is a worse answer than the bare URL it replaced.

    Returns "" or why the block wrote nothing.
    """
    rows = [i for i in items if isinstance(i, dict) and str(i.get("url", "")).strip()]
    if not rows:
        return "links block has no items with a url"

    heading = entry_value(block, ENTRY_HEADING_KEYS)
    if heading:
        para = doc.add_paragraph(heading, style="Normal")
        docx_style.style_runs(para, bold="true", color=hue)

    for item in rows:
        url = str(item["url"]).strip()
        label = str(item.get("label") or item.get("text") or url)
        note = str(item.get("note") or "")
        para = doc.add_paragraph("", style="Normal")
        para.add_run("• ")
        docx_style.add_hyperlink(para, url, label)
        if note:
            from docx.shared import Pt  # type: ignore[import-untyped]

            run = para.add_run(f" — {note}")
            run.font.size = Pt(9)
    return ""


def _add_risks(doc: Any, items: list[Any], hue: str, band: str, docx_style: Any) -> str:
    """The risks table, scored on the same three levels everything else uses.

    A generic `table` block could hold this. It would also let every caller
    invent its own columns and its own severity words, and the review asked for
    a risks table beside a KPI table and a findings table precisely so a reader
    meets the same shape each time.

    Returns "" or a note about levels that were not recognised. An unrecognised
    level is written through uncoloured rather than dropped -- losing a risk to
    a spelling is the one failure a risk register must not have.
    """
    header = ["Risk", "Level", "Impact", "Mitigation", "Owner"]
    present = [
        column
        for column in header
        if column == "Risk"
        or column == "Level"
        or any(str(i.get(column.lower(), "")).strip() for i in items if isinstance(i, dict))
    ]

    rows: list[list[str]] = []
    levels: list[str] = []
    unrecognised: list[str] = []
    for item in items:
        if not isinstance(item, dict):
            item = {"risk": str(item)}
        level = str(item.get("level") or item.get("severity") or "").strip().lower()
        if level and level not in RISK_LEVELS:
            unrecognised.append(level)
        levels.append(level)
        rows.append(
            [
                str(item.get("risk") or item.get("text") or "") if c == "Risk" else str(item.get(c.lower(), ""))
                for c in present
            ]
        )
        rows[-1][present.index("Level")] = level.title() if level else ""

    _add_block_table(doc, present, rows, hue, band, [], docx_style)

    # Colour the Level cell after the table is built, so the banding underneath
    # is already in place and this wins on top of it.
    table = doc.tables[-1]
    column = present.index("Level")
    for row_index, level in enumerate(levels, start=1):
        fill = _RISK_FILLS.get(level)
        if not fill:
            continue
        cell = table.rows[row_index].cells[column]
        docx_style.set_cell_fill(cell, fill)
        for para in cell.paragraphs:
            docx_style.style_runs(para, bold="true", color="FFFFFF")

    if unrecognised:
        return (
            f"level(s) {sorted(set(unrecognised))} are not one of {', '.join(RISK_LEVELS)}; "
            "those rows were written uncoloured"
        )
    return ""


def _add_checklist(doc: Any, items: list[Any], hue: str, docx_style: Any) -> None:
    """Ticked and unticked actions, as text rather than as form controls.

    Word's real check-box content control needs a `w:sdt` and only behaves
    inside a form-protected document; opened anywhere else it is an empty grey
    rectangle. A glyph is what a reader can actually see, print and tick.
    """
    for item in items:
        if isinstance(item, dict):
            text = entry_value(item, ENTRY_TEXT_KEYS)
            done = bool(item.get("done") or item.get("checked") or item.get("complete"))
        else:
            text, done = str(item), False
        para = doc.add_paragraph("", style="Normal")
        mark = para.add_run(f"{_CHECK_DONE if done else _CHECK_OPEN}  ")
        mark.bold = True
        if done:
            from docx.shared import RGBColor  # type: ignore[import-untyped]

            mark.font.color.rgb = RGBColor.from_string(_RISK_FILLS["low"])
        para.add_run(text)


def _apply_fonts(doc: Any, font: str, heading_font: str) -> dict[str, str]:
    """Brand typefaces, applied to the styles rather than to every run.

    Setting a style's font is what makes text a caller adds *afterwards* match
    too. Word also needs the East-Asian name set, or it silently substitutes on
    a machine with a different default -- the kind of difference that only shows
    up on somebody else's screen.
    """
    applied: dict[str, str] = {}
    body = str(font or "").strip()
    heads = str(heading_font or font or "").strip()

    def _set(style_name: str, family: str) -> bool:
        try:
            style = doc.styles[style_name]
        except KeyError:
            return False
        style.font.name = family
        element = style.element.rPr
        if element is not None and element.rFonts is not None:
            from docx.oxml.ns import qn  # type: ignore[import-untyped]

            element.rFonts.set(qn("w:eastAsia"), family)
        return True

    if body and _set("Normal", body):
        applied["body"] = body
    if heads:
        for level in range(1, 7):
            _set(f"Heading {level}", heads)
        applied["headings"] = heads
    return applied


def _add_callout(doc: Any, block: dict[str, Any], hue: str, fill: str, docx_style: Any) -> None:
    """One tinted, borderless cell for the single message that must land."""
    heading = entry_value(block, ENTRY_HEADING_KEYS)
    text = entry_value(block, ENTRY_TEXT_KEYS)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    docx_style.set_cell_fill(cell, fill)
    first = cell.paragraphs[0]
    if heading:
        first.text = heading
        docx_style.style_runs(first, bold="true", color=hue)
        if text:
            body = cell.add_paragraph(text)
            docx_style.style_runs(body, color="1A1A1A")
    else:
        first.text = text
        docx_style.style_runs(first, color="1A1A1A")


def create_from_template(
    template_path: str,
    output_path: str,
    substitutions: dict[str, str],
    open_after: bool = True,
    return_content: bool = False,
) -> dict[str, Any]:
    """Copy a .docx template, apply text substitutions, save to output_path."""
    progress: list[dict[str, Any]] = []
    try:
        import shutil

        import docxedit  # type: ignore[import-untyped]
        from docx import Document  # type: ignore[import-untyped]

        tpl_path = Path(template_path).resolve()
        if not tpl_path.exists():
            progress.append(fail("Template not found", str(tpl_path)))
            return _err(
                progress,
                f"File not found: {template_path}",
                "Check that template_path is an absolute path to an existing .docx file.",
            )
        if tpl_path.suffix.lower() != ".docx":
            progress.append(fail("Template is not a .docx file", tpl_path.suffix))
            return _err(
                progress,
                f"Expected .docx file, got {tpl_path.suffix}",
                "Use the correct server for this file type.",
            )

        out_path = resolve_output_path(output_path, "document.docx")
        _ensure_parent(out_path)

        progress.append(ok(f"Opened template {tpl_path.name}"))

        # Copy template to output location first
        shutil.copy2(str(tpl_path), str(out_path))

        # Open the copy and apply substitutions using run-level editing
        doc = Document(str(out_path))

        if not isinstance(substitutions, dict):
            progress.append(fail("substitutions must be a dict"))
            return _err(
                progress,
                "substitutions must be a dict",
                'Pass a dict like {"PLACEHOLDER": "replacement value"}.',
            )

        # Decide what to search for before touching anything: a key that only
        # appears delimited is matched in that form, and a key that matches
        # nothing is known to match nothing rather than having been eaten by an
        # earlier replacement.
        targets, notes = resolve_targets(_all_text(doc), substitutions)
        for note in notes:
            progress.append(info(note))

        pairs = ordered_pairs(targets, substitutions)
        counts = {target: _count_occurrences(doc, target) for target, _ in pairs}

        # Two phases, because one sequential pass lets a key consume another key
        # or a value an earlier replacement produced. Every target becomes a
        # sentinel first; only then do sentinels become values.
        # docxedit.replace_string operates on runs — never assigns .text directly.
        for index, (target, _) in enumerate(pairs):
            docxedit.replace_string(doc, target, sentinel_for(index))
        for index, (_, value) in enumerate(pairs):
            docxedit.replace_string(doc, sentinel_for(index), value)

        applied = 0
        for target, value in pairs:
            count = counts.get(target, 0)
            progress.append(
                ok(
                    f"Replaced '{target}' → '{value}'",
                    f"{count} occurrence{'s' if count != 1 else ''}",
                )
            )
            applied += count
        for key in substitutions:
            if str(key) not in targets:
                progress.append(warn(f"Placeholder '{key}' not found in template"))

        doc.save(str(out_path))
        progress.append(ok(f"Saved {out_path.name}", f"{applied} total substitutions"))

        _open_if_requested(out_path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "create_from_template",
            "output": str(out_path),
            "output_name": out_path.name,
            "substitutions_applied": applied,
            "progress": progress,
        }
        embed_content(result, out_path, return_content)
        result["token_estimate"] = _token_estimate(result)
        return result
    except Exception as exc:
        logger.warning("create_from_template failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check that template_path points to a valid .docx file and substitutions is a dict.",
        )


def create_letter(
    output_path: str,
    from_name: str,
    to_name: str,
    subject: str,
    body: str,
    open_after: bool = True,
    return_content: bool = False,
) -> dict[str, Any]:
    """Create a formatted business letter .docx."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_output_path(output_path, "document.docx")
        _ensure_parent(path)
        progress.append(info("Creating business letter", path.name))

        doc = Document()

        # Sender name
        sender_para = doc.add_paragraph()
        sender_run = sender_para.add_run(from_name)
        sender_run.bold = True

        # Blank line
        doc.add_paragraph()

        # Recipient
        doc.add_paragraph(f"To: {to_name}")

        # Blank line
        doc.add_paragraph()

        # Subject line — bold
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run(f"Subject: {subject}")
        subject_run.bold = True

        # Blank line
        doc.add_paragraph()

        # Body paragraphs — split on newlines
        body_paras = body.split("\n")
        for line in body_paras:
            doc.add_paragraph(line)

        # Blank line before closing
        doc.add_paragraph()

        # Closing with sender name
        closing_para = doc.add_paragraph()
        closing_para.add_run("Sincerely,")
        doc.add_paragraph()
        sign_para = doc.add_paragraph()
        sign_run = sign_para.add_run(from_name)
        sign_run.bold = True

        doc.save(str(path))
        progress.append(ok(f"Saved {path.name}", f"From: {from_name} | To: {to_name}"))

        _open_if_requested(path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "create_letter",
            "output": str(path),
            "output_name": path.name,
            "from_name": from_name,
            "to_name": to_name,
            "subject": subject,
            "progress": progress,
        }
        embed_content(result, path, return_content)
        result["token_estimate"] = _token_estimate(result)
        return result
    except Exception as exc:
        logger.warning("create_letter failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check that output_path is a valid writable path ending in .docx.",
        )


def merge_documents(
    file_paths: list,
    output_path: str,
    add_page_break: bool = True,
    open_after: bool = True,
    return_content: bool = False,
) -> dict[str, Any]:
    """Merge multiple .docx files into one document."""
    progress: list[dict[str, Any]] = []
    try:
        import copy

        from docx import Document  # type: ignore[import-untyped]
        from docx.oxml import OxmlElement  # type: ignore[import-untyped]
        from docx.oxml.ns import qn  # type: ignore[import-untyped]

        if not isinstance(file_paths, list) or len(file_paths) == 0:
            progress.append(fail("file_paths must be a non-empty list"))
            return _err(
                progress,
                "file_paths must be a non-empty list",
                "Pass a list of absolute paths to .docx files.",
            )

        out_path = resolve_output_path(output_path, "document.docx")
        if out_path.suffix.lower() != ".docx":
            progress.append(fail("output_path must end in .docx", out_path.suffix))
            return _err(
                progress,
                f"Expected .docx output_path, got {out_path.suffix}",
                "Set output_path to a path ending in .docx.",
            )

        # Validate all source files up front
        resolved: list[Path] = []
        for fp in file_paths:
            p = Path(str(fp)).resolve()
            if not p.exists():
                progress.append(fail("File not found", str(p)))
                return _err(
                    progress,
                    f"File not found: {fp}",
                    "Check that all file_paths exist and are accessible.",
                )
            if p.suffix.lower() != ".docx":
                progress.append(fail("Not a .docx file", p.suffix))
                return _err(
                    progress,
                    f"Expected .docx file, got {p.suffix}",
                    "All file_paths must point to .docx files.",
                )
            resolved.append(p)

        _ensure_parent(out_path)
        progress.append(info(f"Merging {len(resolved)} documents", out_path.name))

        base_doc = Document()
        # Remove the default empty paragraph python-docx adds to a new document
        for para in base_doc.paragraphs:
            p_elem = para._element
            p_elem.getparent().remove(p_elem)

        for idx, src_path in enumerate(resolved):
            src_doc = Document(str(src_path))

            if add_page_break and idx > 0:
                page_para = base_doc.add_paragraph()
                run = page_para.add_run()
                break_elem = OxmlElement("w:br")
                break_elem.set(qn("w:type"), "page")
                run._r.append(break_elem)
                progress.append(info(f"Added page break before {src_path.name}"))

            for para in src_doc.paragraphs:
                new_para = copy.deepcopy(para._element)
                base_doc.element.body.append(new_para)

            progress.append(ok(f"Merged {src_path.name}"))

        base_doc.save(str(out_path))
        progress.append(ok(f"Saved {out_path.name}", f"{len(resolved)} documents merged"))

        _open_if_requested(out_path, open_after, progress)

        result: dict[str, Any] = {
            "success": True,
            "op": "merge_documents",
            "output": str(out_path),
            "output_name": out_path.name,
            "merged_count": len(resolved),
            "progress": progress,
        }
        embed_content(result, out_path, return_content)
        result["token_estimate"] = _token_estimate(result)
        return result
    except Exception as exc:
        logger.warning("merge_documents failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check that all file_paths are valid .docx files and output_path is writable.",
        )


def batch_create_from_template(
    template_path: str,
    data_list: list,
    output_dir: str,
    filename_key: str = "",
    open_after: bool = False,
) -> dict[str, Any]:
    """Generate N .docx files from a template + list of {key:value} dicts."""
    progress: list[dict[str, Any]] = []
    try:
        import shutil

        import docxedit  # type: ignore[import-untyped]
        from docx import Document  # type: ignore[import-untyped]

        tpl_path = Path(template_path).resolve()
        if not tpl_path.exists():
            progress.append(fail("Template not found", str(tpl_path)))
            return _err(
                progress,
                f"File not found: {template_path}",
                "Check that template_path is an absolute path to an existing .docx file.",
            )
        if tpl_path.suffix.lower() != ".docx":
            progress.append(fail("Template is not a .docx file", tpl_path.suffix))
            return _err(
                progress,
                f"Expected .docx file, got {tpl_path.suffix}",
                "Use the correct server for this file type.",
            )

        if not isinstance(data_list, list) or len(data_list) == 0:
            progress.append(fail("data_list must be a non-empty list"))
            return _err(
                progress,
                "data_list must be a non-empty list of dicts",
                'Pass a list like [{"NAME": "Alice", "DATE": "April 1"}].',
            )

        out_dir = resolve_output_path(output_dir, "documents")
        out_dir.mkdir(parents=True, exist_ok=True)
        progress.append(info(f"Generating {len(data_list)} documents", str(out_dir)))

        created_files: list[str] = []

        for idx, data_dict in enumerate(data_list):
            if not isinstance(data_dict, dict):
                progress.append(warn(f"Item {idx} is not a dict, skipping"))
                continue

            # Determine output filename. The stem is a stem, so a caller who
            # names the row's file the obvious way -- "alice.docx" -- must not
            # get alice.docx.docx. Only the extension this tool would append is
            # stripped, and only from the end: "alice.docx.bak" keeps its name.
            if filename_key and filename_key in data_dict:
                stem = str(data_dict[filename_key])
                if stem.lower().endswith(".docx"):
                    stem = stem[: -len(".docx")]
            else:
                stem = f"document_{idx + 1:03d}"

            out_file = out_dir / f"{stem}.docx"

            # Copy template to output location
            shutil.copy2(str(tpl_path), str(out_file))

            # Open copy and apply substitutions using run-level editing.
            # This used to hardcode "{{key}}" while create_from_template beside
            # it matched the raw key, so one template could not serve both: a
            # sweep pointed this at a {key} template and got two documents that
            # were byte-identical to it, every placeholder still in place, under
            # success: true. Both now resolve targets the same way.
            doc = Document(str(out_file))
            targets, _notes = resolve_targets(_all_text(doc), data_dict)
            pairs = ordered_pairs(targets, data_dict)
            unfilled = [str(k) for k in data_dict if str(k) not in targets]
            if unfilled:
                progress.append(warn(f"{stem}: no placeholder for {', '.join(unfilled)}"))
            for index, (target, _value) in enumerate(pairs):
                docxedit.replace_string(doc, target, sentinel_for(index))
            for index, (_target, value) in enumerate(pairs):
                docxedit.replace_string(doc, sentinel_for(index), str(value))

            doc.save(str(out_file))
            created_files.append(out_file.name)
            progress.append(ok(f"Created {out_file.name}"))

            if open_after:
                open_file(out_file)

        progress.append(ok("Batch complete", f"{len(created_files)} of {len(data_list)} created"))

        return {
            "success": True,
            "op": "batch_create",
            "output_dir": str(out_dir),
            "created_count": len(created_files),
            "files": created_files,
            "progress": progress,
            "token_estimate": _token_estimate(progress),
        }
    except Exception as exc:
        logger.warning("batch_create_from_template failed: %s", exc)
        progress.append(fail(str(exc)))
        return _err(
            progress,
            str(exc),
            "Check template_path, output_dir permissions, and that data_list is a list of dicts.",
        )
