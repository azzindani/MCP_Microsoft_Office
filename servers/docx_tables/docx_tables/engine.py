"""DOCX Tables engine — pure python-docx logic, no MCP imports."""

from pathlib import Path
from typing import Any

from shared.file_utils import hint_for_error, hint_for_message, resolve_path
from shared.live_edit import notify_reload
from shared.platform_utils import open_file
from shared.progress import describe_error, fail, index_range, ok, warn
from shared.receipt import append_receipt
from shared.version_control import snapshot

# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _not_found(path: Path, progress: list[dict[str, Any]]) -> dict[str, Any]:
    progress.append(fail("Cannot open file", f"Path not found: {path}"))
    return {
        "success": False,
        "error": f"File not found: {path}",
        "hint": (
            "Check that file_path is absolute and the file exists. "
            "On Windows use forward slashes: C:/Users/you/doc.docx"
        ),
        "progress": progress,
        "token_estimate": 40,
    }


def _wrong_type(path: Path, expected: str, progress: list[dict[str, Any]]) -> dict[str, Any]:
    progress.append(fail(f"Wrong file type: {path.suffix}", f"Expected: {expected}"))
    return {
        "success": False,
        "error": f"Expected {expected} file, got {path.suffix}",
        "hint": "Use the correct server for this file type.",
        "progress": progress,
        "token_estimate": 30,
    }


def _error(msg: str, hint: str, progress: list[dict[str, Any]], backup: str | None = None) -> dict[str, Any]:
    result: dict[str, Any] = {
        "success": False,
        "error": describe_error(msg),
        "hint": hint_for_message(msg, hint),
        "progress": progress,
        "token_estimate": len(str(progress)) // 4,
    }
    if backup is not None:
        result["backup"] = backup
    return result


def _table_dims(table: Any) -> tuple[int, int]:
    """Return (row_count, col_count) for a table."""
    rows = len(table.rows)
    cols = len(table.columns) if rows else 0
    return rows, cols


# A column narrower than this is unreadable whatever is in it -- a date or a
# figure still needs room not to break across lines.
_MIN_COL_EMU = 457200  # 0.5 inch

# Roughly the width of one character of body text, plus the cell's own padding.
# Used to ask "how much does this column actually need", not to typeset -- the
# real answer depends on the font, and a column an eighth of an inch too wide
# costs nothing while one an eighth too narrow breaks a figure in half.
_CHAR_EMU = 77000  # ~0.084 inch
_CELL_PADDING_EMU = 137000  # ~0.15 inch

# Leftover width is shared on the SQUARE ROOT of a column's longest cell rather
# than on the length itself, so one enormous column does not take everything.
_WEIGHT_EXPONENT = 0.5


def _fit_columns(doc: Any, table: Any, data: list[list[str]], cols: int) -> None:
    """Size columns from their content instead of splitting the page evenly.

    `table.autofit = True` is the obvious answer and it does not work: it sets
    `tblLayout`, and both Word and LibreOffice still lay the table out on the
    equal `gridCol` widths python-docx wrote. Rendered and looked at, a column
    holding `DecreaseIncreaseInPlacementsWithOtherBanksAndBankIndonesia` got the
    same third of the page as one holding `4.019`, and the identifier broke
    mid-word across three lines -- `DecreaseIncreaseInPla / cementsWithOtherBan /
    ksAndBankIndonesia`. Explicit widths are honoured by both.

    Width has to be set on every CELL, not only on the column: Word reads
    `tcW` per cell and treats `gridCol` as a hint, so setting one and not the
    other leaves the layout unchanged in Word while looking right in LibreOffice.

    **Columns that fit are served first, and the rest share what is left.** A
    purely proportional split -- even damped by a square root -- still starves
    the small column when another is enormous. Measured on the reconciliation
    appendix, whose context column runs to 120 characters: the value column came
    out at 1.13in and broke `-17.906.497` into `-17.906.4 / 97`. A wrapped
    identifier is a cosmetic problem; a figure split across two lines invites a
    misread, so a column asking for less than an even share is simply given what
    it asks for, and only the greedy columns compete for the remainder.
    """
    from docx.shared import Emu  # type: ignore[import-untyped]

    section = doc.sections[0]
    usable = int(section.page_width - section.left_margin - section.right_margin)
    if usable <= 0 or cols <= 0:
        return

    longest = [1] * cols
    for row in data:
        for c_idx in range(min(cols, len(row))):
            longest[c_idx] = max(longest[c_idx], len(str(row[c_idx])))

    needed = [n * _CHAR_EMU + _CELL_PADDING_EMU for n in longest]
    even = usable // cols

    # A column that wants no more than an even share gets exactly what it wants.
    modest = [c for c in range(cols) if needed[c] <= even]
    widths = [0] * cols
    for c_idx in modest:
        widths[c_idx] = max(_MIN_COL_EMU, needed[c_idx])

    greedy = [c for c in range(cols) if c not in modest]
    remaining = usable - sum(widths)
    if greedy:
        weights = {c: longest[c] ** _WEIGHT_EXPONENT for c in greedy}
        total = sum(weights.values()) or 1.0
        for c_idx in greedy:
            widths[c_idx] = max(_MIN_COL_EMU, int(remaining * weights[c_idx] / total))

    # Rounding and the floor can push the row past the printable width; scale
    # back proportionally rather than letting the last column run off the page.
    over = sum(widths) - usable
    if over > 0:
        widths = [max(_MIN_COL_EMU, w - int(over * w / sum(widths))) for w in widths]

    table.autofit = False
    for c_idx, width in enumerate(widths):
        table.columns[c_idx].width = Emu(width)
        for cell in table.columns[c_idx].cells:
            cell.width = Emu(width)


def _set_cell_text(cell: Any, text: str) -> None:
    """
    Set cell text safely — removes extra paragraphs, preserves the first.

    Does NOT use cell.text = value which collapses formatting in multi-para
    cells. Instead clears the first paragraph's runs and sets new text.
    """
    # Remove extra paragraphs (keep only the first)
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)  # type: ignore[attr-defined]

    # Work with the first (and now only) paragraph
    first_para = cell.paragraphs[0]
    # Remove existing runs
    for run in first_para.runs:
        run._element.getparent().remove(run._element)  # type: ignore[attr-defined]
    # Add a single run with the new text
    first_para.add_run(text)


# ---------------------------------------------------------------------------
# Read tools
# ---------------------------------------------------------------------------


def list_tables(file_path: str) -> dict[str, Any]:
    """Return count and row/col dimensions for every table in the document."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        table_list = []
        for i, tbl in enumerate(tables):
            rows, cols = _table_dims(tbl)
            table_list.append({"index": i, "rows": rows, "cols": cols})

        return {
            "success": True,
            "file": str(path),
            "table_count": len(table_list),
            "tables": table_list,
            "progress": progress,
            "token_estimate": len(str(table_list)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check that file_path points to a valid .docx file.", progress)


def read_table(file_path: str, table_index: int) -> dict[str, Any]:
    """Return full 2-D cell array for one table. Merged cells marked."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        if table_index < 0 or table_index >= len(tables):
            progress.append(
                fail(
                    f"Table index {table_index} out of range",
                    f"Available: 0-{len(tables) - 1}" if tables else "No tables found",
                )
            )
            return _error(
                f"Table index {table_index} out of range {index_range(len(tables), 'tables')}",
                "Add one with add_table() first." if not tables else "Use list_tables to see available tables.",
                progress,
            )

        tbl = tables[table_index]
        rows, cols = _table_dims(tbl)
        progress.append(ok(f"Reading table {table_index}", f"{rows} rows × {cols} cols"))

        # Track merged cells by their XML element — cells sharing the same
        # element are part of a horizontal merge.
        data: list[list[dict[str, Any]]] = []
        for r_idx, row in enumerate(tbl.rows):
            row_data: list[dict[str, Any]] = []
            seen_elements: set[int] = set()
            for c_idx, cell in enumerate(row.cells):
                elem_id = id(cell._element)  # type: ignore[attr-defined]
                if elem_id in seen_elements:
                    row_data.append({"col": c_idx, "text": "merged"})
                else:
                    seen_elements.add(elem_id)
                    row_data.append({"col": c_idx, "text": cell.text})
            data.append(row_data)

        return {
            "success": True,
            "table_index": table_index,
            "rows": rows,
            "cols": cols,
            "data": data,
            "progress": progress,
            "token_estimate": len(str(data)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path and table_index.", progress)


def search_table_cells(file_path: str, query: str, max_results: int = 10) -> dict[str, Any]:
    """Scan all table cells for text match. Return cell coordinates only."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)
        if not query:
            return _error("query cannot be empty", "Provide a non-empty search string.", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        query_lower = query.lower()
        matches: list[dict[str, Any]] = []
        cells_scanned = 0

        for t_idx, tbl in enumerate(tables):
            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    cells_scanned += 1
                    if query_lower in cell.text.lower():
                        matches.append(
                            {
                                "table_index": t_idx,
                                "row": r_idx,
                                "col": c_idx,
                                "text": cell.text,
                            }
                        )
                        if len(matches) >= max_results:
                            break
                if len(matches) >= max_results:
                    break
            if len(matches) >= max_results:
                break

        progress.append(
            ok(
                f"Found {len(matches)} match{'es' if len(matches) != 1 else ''} for '{query}'",
                f"{cells_scanned} cells scanned",
            )
        )

        return {
            "success": True,
            "query": query,
            "matches": matches,
            "total_cells_scanned": cells_scanned,
            "truncated": len(matches) >= max_results,
            "progress": progress,
            "token_estimate": len(str(matches)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check that file_path points to a valid .docx file.", progress)


def read_table_row(file_path: str, table_index: int, row: int) -> dict[str, Any]:
    """Return all cells in one table row."""
    progress: list[dict[str, Any]] = []
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        if table_index < 0 or table_index >= len(tables):
            progress.append(fail(f"Table index {table_index} out of range"))
            return _error(
                f"Table index {table_index} out of range {index_range(len(tables), 'tables')}",
                "Add one with add_table() first." if not tables else "Use list_tables to see available tables.",
                progress,
            )

        tbl = tables[table_index]
        rows, cols = _table_dims(tbl)

        if row < 0 or row >= rows:
            progress.append(fail(f"Row {row} out of range"))
            return _error(
                f"Row {row} out of range {index_range(rows, 'rows')}",
                "Use read_table to see the full table structure.",
                progress,
            )

        cells = [{"col": c_idx, "text": cell.text} for c_idx, cell in enumerate(tbl.rows[row].cells)]
        progress.append(ok(f"Read row {row} of table {table_index}", f"{len(cells)} cells"))

        return {
            "success": True,
            "table_index": table_index,
            "row": row,
            "cells": cells,
            "progress": progress,
            "token_estimate": len(str(cells)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(str(e), "Check file_path, table_index, and row.", progress)


# ---------------------------------------------------------------------------
# Write tools
# ---------------------------------------------------------------------------


def set_cell(
    file_path: str, table_index: int, row: int, col: int, text: str, open_after: bool = False
) -> dict[str, Any]:
    """Write text to a specific table cell. Snapshot taken before write."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        if table_index < 0 or table_index >= len(tables):
            progress.append(fail(f"Table index {table_index} out of range"))
            return _error(
                f"Table index {table_index} out of range {index_range(len(tables), 'tables')}",
                "Add one with add_table() first." if not tables else "Use list_tables to see available tables.",
                progress,
            )

        tbl = tables[table_index]
        rows, cols = _table_dims(tbl)

        if row < 0 or row >= rows:
            progress.append(fail(f"Row {row} out of range {index_range(rows, 'rows')}"))
            return _error(
                f"Row index {row} out of range {index_range(rows, 'rows')}",
                "Use read_table to see the full table structure.",
                progress,
            )
        if col < 0 or col >= cols:
            progress.append(fail(f"Col {col} out of range {index_range(cols, 'columns')}"))
            return _error(
                f"Col index {col} out of range {index_range(cols, 'columns')}",
                "Use read_table to see the full table structure.",
                progress,
            )

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        cell = tbl.rows[row].cells[col]
        old_text = cell.text
        _set_cell_text(cell, text)

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            tool="set_cell",
            server="docx-tables",
            args={"table_index": table_index, "row": row, "col": col, "text": text},
            result=f"✔ Set table[{table_index}][{row}][{col}] = '{text}'",
            backup=backup,
            success=True,
        )

        return {
            "success": True,
            "op": "set_cell",
            "table_index": table_index,
            "row": row,
            "col": col,
            "old_text": old_text,
            "new_text": text,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            tool="set_cell",
            server="docx-tables",
            args={"table_index": table_index, "row": row, "col": col, "text": text},
            result=f"✘ {e}",
            backup=backup,
            success=False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def add_row(file_path: str, table_index: int, data: list[str], open_after: bool = False) -> dict[str, Any]:
    """Append a row to table N. data is a list of cell strings."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        if table_index < 0 or table_index >= len(tables):
            progress.append(fail(f"Table index {table_index} out of range"))
            return _error(
                f"Table index {table_index} out of range {index_range(len(tables), 'tables')}",
                "Add one with add_table() first." if not tables else "Use list_tables to see available tables.",
                progress,
            )

        tbl = tables[table_index]
        rows, cols = _table_dims(tbl)

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        # Add a new row by copying the structure of the last row
        new_row = tbl.add_row()
        for c_idx, cell in enumerate(new_row.cells):
            cell_text = data[c_idx] if c_idx < len(data) else ""
            _set_cell_text(cell, cell_text)

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(
            ok(
                f"Added row {rows} to table {table_index}",
                f"{len(data)} cell values provided",
            )
        )
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            tool="add_row",
            server="docx-tables",
            args={"table_index": table_index, "data": data},
            result=f"✔ Added row to table[{table_index}]",
            backup=backup,
            success=True,
        )

        return {
            "success": True,
            "op": "add_row",
            "table_index": table_index,
            "new_row_index": rows,
            "data": data,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            tool="add_row",
            server="docx-tables",
            args={"table_index": table_index, "data": data},
            result=f"✘ {e}",
            backup=backup,
            success=False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def delete_row(file_path: str, table_index: int, row: int, open_after: bool = False) -> dict[str, Any]:
    """Remove a row from table N. Rows below shift up."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        if table_index < 0 or table_index >= len(tables):
            progress.append(fail(f"Table index {table_index} out of range"))
            return _error(
                f"Table index {table_index} out of range {index_range(len(tables), 'tables')}",
                "Add one with add_table() first." if not tables else "Use list_tables to see available tables.",
                progress,
            )

        tbl = tables[table_index]
        rows, cols = _table_dims(tbl)

        if row < 0 or row >= rows:
            progress.append(fail(f"Row {row} out of range {index_range(rows, 'rows')}"))
            return _error(
                f"Row index {row} out of range {index_range(rows, 'rows')}",
                "Use read_table to see the full table structure.",
                progress,
            )

        # Capture deleted row text for the receipt
        deleted_text = [tbl.rows[row].cells[c].text for c in range(cols)]

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        # Remove the row element from the table's XML
        row_element = tbl.rows[row]._element  # type: ignore[attr-defined]
        row_element.getparent().remove(row_element)

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(ok(f"Deleted row {row} from table {table_index}"))
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            tool="delete_row",
            server="docx-tables",
            args={"table_index": table_index, "row": row},
            result=f"✔ Deleted row {row} from table[{table_index}]",
            backup=backup,
            success=True,
        )

        return {
            "success": True,
            "op": "delete_row",
            "table_index": table_index,
            "deleted_row": row,
            "deleted_text": deleted_text,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            tool="delete_row",
            server="docx-tables",
            args={"table_index": table_index, "row": row},
            result=f"✘ {e}",
            backup=backup,
            success=False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def add_table(
    file_path: str,
    after_paragraph_index: int,
    rows: int,
    cols: int,
    data: list[list[str]] | None = None,
    open_after: bool = False,
) -> dict[str, Any]:
    """Insert table after paragraph N. -1 = before the first. data is optional."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        if rows < 1 or cols < 1:
            return _error(
                "rows and cols must be at least 1",
                "Provide positive integer values for rows and cols.",
                progress,
            )

        doc = Document(str(path))
        paragraphs = doc.paragraphs
        total_paras = len(paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{total_paras} paragraphs"))

        # A brand-new document has no paragraphs at all, so every anchor index --
        # including 0 and -1 -- used to be "out of range" and the only tool that
        # can create a table was unreachable on the one file that most needs it.
        # create_document() then add_table() is the obvious two-step, and it
        # failed; the hint sent the caller to read_document, which returns an
        # empty list. With no paragraphs there is exactly one place a table can
        # go, so there is nothing to disambiguate: it goes in the empty body,
        # which is where python-docx has already put it below.
        #
        # -1 means "before the first paragraph", which is what insert_paragraph
        # in docx-basic already documents for its own after_index. The two tools
        # take the same kind of anchor into the same document, so a caller who
        # learns the convention from one reaches for it in the other: a sweep
        # called add_table(-1) and got "Paragraph index -1 out of range (0-6)"
        # from the tool right next to the one that accepts it.
        if total_paras and (after_paragraph_index < -1 or after_paragraph_index >= total_paras):
            progress.append(fail(f"Paragraph index {after_paragraph_index} out of range"))
            return _error(
                f"Paragraph index {after_paragraph_index} out of range {index_range(total_paras, 'paragraphs')}",
                "Use read_document to see paragraph indices.",
                progress,
            )

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        # "No paragraphs" is not the same as "empty". A document can hold tables
        # and nothing else, and then there is more than one place a table can go
        # after all -- so -1 has a front to be placed at, and calling the file
        # empty in the reply describes a document that is not.
        body_before = [child for child in doc.element.body if not child.tag.endswith("}sectPr")]

        # Create a new table (appended to body first by python-docx).
        #
        # WITH BORDERS, and that is not decoration. python-docx's default style
        # is "Normal Table", which draws no rules at all, so what Word shows is
        # columns of text floating in whitespace -- and a table without rules is
        # read as a table only where the columns happen to line up. Rendered and
        # looked at, a three-column reconciliation came out as unreadable
        # gutters; the same data under "Table Grid" is a table.
        #
        # `autofit` sizes columns to their content instead of splitting the page
        # into equal fractions. Without it a column holding
        # `DecreaseIncreaseInPlacementsWithOtherBanksAndBankIndonesia` gets the
        # same third of the page as one holding `4.019`, and the long name is
        # broken mid-word across three lines.
        #
        # "Table Grid" is a built-in Word style, present in the default template
        # python-docx ships and in every real .docx, but a caller's template can
        # omit it -- so a missing style falls back to the unstyled table rather
        # than failing a write the caller asked for.
        try:
            tbl = doc.add_table(rows=rows, cols=cols, style="Table Grid")
        except KeyError:
            tbl = doc.add_table(rows=rows, cols=cols)
            progress.append(
                warn("This document's template has no 'Table Grid' style", "table inserted without borders")
            )

        # Populate with data if provided
        if data:
            for r_idx in range(min(rows, len(data))):
                for c_idx in range(min(cols, len(data[r_idx]))):
                    _set_cell_text(tbl.rows[r_idx].cells[c_idx], data[r_idx][c_idx])
            _fit_columns(doc, tbl, data, cols)

        # Move the table element to appear after the target paragraph. With no
        # paragraphs there is no anchor to move to, and python-docx has already
        # appended it to the body, which is the only position available.
        if total_paras:
            anchor_para = paragraphs[max(after_paragraph_index, 0)]._element  # type: ignore[attr-defined]
            tbl_element = tbl._element  # type: ignore[attr-defined]
            # Remove from current position (end of body)
            tbl_element.getparent().remove(tbl_element)
            if after_paragraph_index == -1:
                anchor_para.addprevious(tbl_element)
            else:
                anchor_para.addnext(tbl_element)
        elif after_paragraph_index == -1 and body_before:
            # No paragraphs, but the body is not empty -- honour "first" against
            # whatever content is there rather than leaving the table where
            # python-docx appended it, which is the other end entirely.
            tbl_element = tbl._element  # type: ignore[attr-defined]
            tbl_element.getparent().remove(tbl_element)
            body_before[0].addprevious(tbl_element)

        doc.save(str(path))
        if open_after:
            open_file(path)
        if not total_paras and not body_before:
            placed = "in the empty document"
        elif not total_paras:
            placed = (
                "before the existing content"
                if after_paragraph_index == -1
                else "at the end (the document has no paragraphs to anchor to)"
            )
        elif after_paragraph_index == -1:
            placed = "before paragraph 0"
        else:
            placed = f"after paragraph {after_paragraph_index}"
        progress.append(
            ok(
                f"Inserted {rows}×{cols} table {placed}",
                f"{'with data' if data else 'empty'}",
            )
        )
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            tool="add_table",
            server="docx-tables",
            args={"after_paragraph_index": after_paragraph_index, "rows": rows, "cols": cols},
            result=f"✔ Inserted {rows}×{cols} table {placed}",
            backup=backup,
            success=True,
        )

        return {
            "success": True,
            "op": "add_table",
            "after_paragraph_index": after_paragraph_index,
            "rows": rows,
            "cols": cols,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            tool="add_table",
            server="docx-tables",
            args={"after_paragraph_index": after_paragraph_index, "rows": rows, "cols": cols},
            result=f"✘ {e}",
            backup=backup,
            success=False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def delete_table(file_path: str, table_index: int, open_after: bool = False) -> dict[str, Any]:
    """Remove table N from the document entirely."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    path: Path | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        tables = doc.tables
        progress.append(ok(f"Opened {path.name}", f"{len(tables)} tables"))

        if table_index < 0 or table_index >= len(tables):
            progress.append(fail(f"Table index {table_index} out of range"))
            return _error(
                f"Table index {table_index} out of range {index_range(len(tables), 'tables')}",
                "Add one with add_table() first." if not tables else "Use list_tables to see available tables.",
                progress,
            )

        tbl = tables[table_index]
        rows, cols = _table_dims(tbl)

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        # Remove table element from the document body
        tbl_element = tbl._element  # type: ignore[attr-defined]
        tbl_element.getparent().remove(tbl_element)

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(ok(f"Deleted table {table_index}", f"was {rows} rows × {cols} cols"))
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            tool="delete_table",
            server="docx-tables",
            args={"table_index": table_index},
            result=f"✔ Deleted table[{table_index}] ({rows}×{cols})",
            backup=backup,
            success=True,
        )

        return {
            "success": True,
            "op": "delete_table",
            "table_index": table_index,
            "deleted_rows": rows,
            "deleted_cols": cols,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            tool="delete_table",
            server="docx-tables",
            args={"table_index": table_index},
            result=f"✘ {e}",
            backup=backup,
            success=False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)
