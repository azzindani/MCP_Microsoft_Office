"""DOCX Layout engine — pure python-docx logic, no MCP imports."""

import subprocess
from pathlib import Path
from typing import Any

from shared.file_utils import embed_content, hint_for_error, resolve_path
from shared.live_edit import notify_reload
from shared.platform_utils import get_pdf_converter, open_file, resolve_output_path
from shared.progress import describe_error, fail, index_range, info, ok
from shared.receipt import append_receipt
from shared.version_control import snapshot

# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _not_found(path: Path, progress: list[dict[str, Any]]) -> dict[str, Any]:
    progress.append(fail("File not found", str(path.name)))
    return {
        "success": False,
        "error": f"File not found: {path}",
        "hint": "Check that file_path is absolute and the file exists.",
        "progress": progress,
        "token_estimate": len(str(progress)) // 4,
    }


def _wrong_type(path: Path, expected: str, progress: list[dict[str, Any]]) -> dict[str, Any]:
    progress.append(fail(f"Wrong file type: {path.suffix}", f"Expected {expected}"))
    return {
        "success": False,
        "error": f"Expected {expected} file, got {path.suffix}",
        "hint": "Use the correct server for this file type.",
        "progress": progress,
        "token_estimate": len(str(progress)) // 4,
    }


def _error(msg: str, hint: str, progress: list[dict[str, Any]], backup: str | None = None) -> dict[str, Any]:
    result: dict[str, Any] = {
        "success": False,
        "error": describe_error(msg),
        "hint": hint,
        "progress": progress,
        "token_estimate": len(str(progress)) // 4,
    }
    if backup is not None:
        result["backup"] = backup
    return result


def _out_of_range(idx: int, total: int, progress: list[dict[str, Any]], backup: str | None = None) -> dict[str, Any]:
    progress.append(fail(f"Paragraph index {idx} out of range", f"0-{total - 1}"))
    return _error(
        f"paragraph_index {idx} out of range {index_range(total, 'paragraphs')}",
        "Use read_document to get current paragraph count.",
        progress,
        backup,
    )


# ---------------------------------------------------------------------------
# Tool implementations
# ---------------------------------------------------------------------------


def set_heading(file_path: str, paragraph_index: int, level: int, open_after: bool = False) -> dict[str, Any]:
    """Apply Heading 1-6 style to paragraph N."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)
        if level < 1 or level > 6:
            progress.append(fail(f"Invalid heading level: {level}", "Must be 1-6"))
            return _error(
                f"level {level} is invalid. Must be 1-6.",
                "Provide a heading level between 1 and 6.",
                progress,
            )

        doc = Document(str(path))
        total = len(doc.paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        if paragraph_index < 0 or paragraph_index >= total:
            return _out_of_range(paragraph_index, total, progress)

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        para = doc.paragraphs[paragraph_index]
        style_name = f"Heading {level}"
        para.style = doc.styles[style_name]  # type: ignore[reportAttributeAccessIssue]
        progress.append(ok(f"Applied {style_name} to paragraph {paragraph_index}"))

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            "set_heading",
            "docx_layout",
            {"paragraph_index": paragraph_index, "level": level},
            f"✔ Applied Heading {level} to paragraph {paragraph_index}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "set_heading",
            "paragraph_index": paragraph_index,
            "style": style_name,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            "set_heading",
            "docx_layout",
            {"paragraph_index": paragraph_index, "level": level},
            f"✘ {e}",
            backup,
            False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def set_font(
    file_path: str,
    paragraph_index: int,
    font_name: str = "",
    font_size: float = 0,
    bold: bool = False,
    italic: bool = False,
    open_after: bool = False,
) -> dict[str, Any]:
    """Set font attributes on all runs in paragraph N."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        total = len(doc.paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        if paragraph_index < 0 or paragraph_index >= total:
            return _out_of_range(paragraph_index, total, progress)

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        para = doc.paragraphs[paragraph_index]
        changes: list[str] = []

        for run in para.runs:
            if font_name:
                run.font.name = font_name
            if font_size > 0:
                run.font.size = Pt(font_size)
            if bold:
                run.bold = True
            if italic:
                run.italic = True

        if font_name:
            changes.append(f"name={font_name}")
        if font_size > 0:
            changes.append(f"size={font_size}pt")
        if bold:
            changes.append("bold=True")
        if italic:
            changes.append("italic=True")

        detail = ", ".join(changes) if changes else "no changes"
        progress.append(ok(f"Updated font on paragraph {paragraph_index}", detail))

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            "set_font",
            "docx_layout",
            {
                "paragraph_index": paragraph_index,
                "font_name": font_name,
                "font_size": font_size,
                "bold": bold,
                "italic": italic,
            },
            f"✔ Font updated: {detail}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "set_font",
            "paragraph_index": paragraph_index,
            "changes": changes,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            "set_font",
            "docx_layout",
            {"paragraph_index": paragraph_index},
            f"✘ {e}",
            backup,
            False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def set_paragraph_style(
    file_path: str, paragraph_index: int, style_name: str, open_after: bool = False
) -> dict[str, Any]:
    """Apply a named style from the document style gallery to paragraph N."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        total = len(doc.paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        # Validate style exists in document
        available_styles = [s.name for s in doc.styles if s.type is not None]
        if style_name not in available_styles:
            progress.append(fail(f"Style not found: {style_name}"))
            return _error(
                f"Style '{style_name}' not found in document.",
                "Available styles include: Normal, Body Text, Heading 1-6.",
                progress,
            )

        if paragraph_index < 0 or paragraph_index >= total:
            return _out_of_range(paragraph_index, total, progress)

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        para = doc.paragraphs[paragraph_index]
        para.style = doc.styles[style_name]  # type: ignore[reportAttributeAccessIssue]
        progress.append(ok(f"Applied style '{style_name}' to paragraph {paragraph_index}"))

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            "set_paragraph_style",
            "docx_layout",
            {"paragraph_index": paragraph_index, "style_name": style_name},
            f"✔ Applied style '{style_name}' to paragraph {paragraph_index}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "set_paragraph_style",
            "paragraph_index": paragraph_index,
            "style": style_name,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            "set_paragraph_style",
            "docx_layout",
            {"paragraph_index": paragraph_index, "style_name": style_name},
            f"✘ {e}",
            backup,
            False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def add_image(
    file_path: str,
    paragraph_index: int,
    image_path: str,
    width_inches: float = 4.0,
    open_after: bool = False,
) -> dict[str, Any]:
    """Insert an image into paragraph N. Supported: PNG, JPG, GIF, BMP, TIFF."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    _SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif"}
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Inches  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        img_path = Path(image_path).resolve()
        if not img_path.exists():
            progress.append(fail("Image file not found", str(img_path.name)))
            return _error(
                f"Image file not found: {image_path}",
                "Check that image_path is an absolute path to an existing image file.",
                progress,
            )
        if img_path.suffix.lower() not in _SUPPORTED_IMAGE_EXTS:
            progress.append(fail(f"Unsupported image format: {img_path.suffix}"))
            return _error(
                f"Unsupported image format: {img_path.suffix}",
                "Supported formats: PNG, JPG, GIF, BMP, TIFF.",
                progress,
            )

        doc = Document(str(path))
        total = len(doc.paragraphs)
        progress.append(ok(f"Opened {path.name}", f"{total} paragraphs"))

        # An image needs a paragraph to live in, and a brand-new document has
        # none -- so every index was out of range and nothing could put the
        # first image into a freshly created file. add_image is the only tool
        # that adds an image at all, so there was no way round it. Give the
        # image a paragraph of its own instead of refusing.
        if total and (paragraph_index < 0 or paragraph_index >= total):
            return _out_of_range(paragraph_index, total, progress)

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        para = doc.paragraphs[paragraph_index] if total else doc.add_paragraph()
        run = para.add_run()
        # python-docx raises UnrecognizedImageError bare, with no message at
        # all, so str(exc) is "" and the generic handler below returned a
        # failure whose `error` was an empty string -- under a hint offering to
        # undo a change that had not happened. The extension check above only
        # reads the suffix, so a text file named .png reaches here: a sweep
        # wrote base64 to q2_dot.png without decoding it and got exactly that
        # blank answer back.
        try:
            from docx.image.exceptions import (  # type: ignore[import-untyped]
                InvalidImageStreamError,
                UnexpectedEndOfFileError,
                UnrecognizedImageError,
            )

            image_errors: tuple[type[BaseException], ...] = (
                InvalidImageStreamError,
                UnexpectedEndOfFileError,
                UnrecognizedImageError,
            )
        except ImportError:  # pragma: no cover - defensive
            image_errors = ()

        try:
            run.add_picture(str(img_path), width=Inches(width_inches))
        except image_errors:
            progress.append(fail("Not a readable image", img_path.name))
            return _error(
                f"Not a readable image: {img_path.name}",
                f"The name ends in {img_path.suffix} but the contents are not a valid image — "
                "check it is the real file and not a placeholder or undecoded base64 text.",
                progress,
                backup,
            )
        progress.append(
            ok(
                f"Inserted image into paragraph {paragraph_index}",
                f"{img_path.name} at {width_inches}in",
            )
        )

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            "add_image",
            "docx_layout",
            {
                "paragraph_index": paragraph_index,
                "image_path": image_path,
                "width_inches": width_inches,
            },
            f"✔ Inserted {img_path.name} at paragraph {paragraph_index}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "add_image",
            "paragraph_index": paragraph_index,
            "image": img_path.name,
            "width_inches": width_inches,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            "add_image",
            "docx_layout",
            {"paragraph_index": paragraph_index, "image_path": image_path},
            f"✘ {e}",
            backup,
            False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def set_page_margins(
    file_path: str,
    top: float = 2.54,
    bottom: float = 2.54,
    left: float = 2.54,
    right: float = 2.54,
    open_after: bool = False,
) -> dict[str, Any]:
    """Set all page margins in cm for every section in the document."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Cm  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        doc = Document(str(path))
        section_count = len(doc.sections)
        progress.append(ok(f"Opened {path.name}", f"{section_count} section(s)"))

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        for section in doc.sections:
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)

        detail = f"top={top}cm bottom={bottom}cm left={left}cm right={right}cm"
        progress.append(ok(f"Set margins on {section_count} section(s)", detail))

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            "set_page_margins",
            "docx_layout",
            {"top": top, "bottom": bottom, "left": left, "right": right},
            f"✔ Margins updated: {detail}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "set_page_margins",
            "top_cm": top,
            "bottom_cm": bottom,
            "left_cm": left,
            "right_cm": right,
            "sections_updated": section_count,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            "set_page_margins",
            "docx_layout",
            {"top": top, "bottom": bottom, "left": left, "right": right},
            f"✘ {e}",
            backup,
            False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def add_header_footer(
    file_path: str,
    text: str,
    location: str = "header",
    open_after: bool = False,
) -> dict[str, Any]:
    """Set header or footer text for all sections. location: header or footer."""
    progress: list[dict[str, Any]] = []
    backup: str | None = None
    try:
        from docx import Document  # type: ignore[import-untyped]

        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)
        if location not in ("header", "footer"):
            progress.append(fail(f"Invalid location: {location}"))
            return _error(
                f"location must be 'header' or 'footer', got '{location}'.",
                "Use location='header' or location='footer'.",
                progress,
            )

        doc = Document(str(path))
        section_count = len(doc.sections)
        progress.append(ok(f"Opened {path.name}", f"{section_count} section(s)"))

        backup = snapshot(str(path))
        progress.append(ok("Snapshot saved", Path(backup).name))

        for section in doc.sections:
            if location == "header":
                section.header.paragraphs[0].text = text
            else:
                section.footer.paragraphs[0].text = text

        progress.append(ok(f"Set {location} text on {section_count} section(s)", f'"{text[:60]}"'))

        doc.save(str(path))
        if open_after:
            open_file(path)
        progress.append(notify_reload(str(path), "docx"))

        append_receipt(
            str(path),
            "add_header_footer",
            "docx_layout",
            {"text": text, "location": location},
            f"✔ {location.capitalize()} set: {text[:60]}",
            backup,
            True,
        )

        return {
            "success": True,
            "op": "add_header_footer",
            "location": location,
            "text": text,
            "sections_updated": section_count,
            "backup": backup,
            "progress": progress,
            "token_estimate": len(str(progress)) // 4,
        }
    except Exception as e:
        progress.append(fail(str(e)))
        append_receipt(
            file_path,
            "add_header_footer",
            "docx_layout",
            {"text": text, "location": location},
            f"✘ {e}",
            backup,
            False,
        )
        return _error(str(e), hint_for_error(e, path), progress, backup)


def export_pdf(
    file_path: str,
    output_path: str = "",
    open_after: bool = True,
    return_content: bool = False,
) -> dict[str, Any]:
    """Export .docx to PDF. Requires Word (Win/Mac) or LibreOffice (Linux)."""
    progress: list[dict[str, Any]] = []
    try:
        path = resolve_path(file_path)
        if not path.exists():
            return _not_found(path, progress)
        if path.suffix.lower() != ".docx":
            return _wrong_type(path, ".docx", progress)

        progress.append(ok(f"Opened {path.name}"))

        converter = get_pdf_converter()
        if converter is None:
            progress.append(fail("No PDF converter found"))
            return _error(
                "No PDF converter available.",
                (
                    "Install LibreOffice: sudo apt install libreoffice (Ubuntu) "
                    "or sudo dnf install libreoffice (Fedora). "
                    "On Windows/macOS, install Microsoft Word."
                ),
                progress,
            )

        # Determine output path — default to Downloads directory
        out_path = resolve_output_path(output_path or path.stem + ".pdf", path.stem + ".pdf")

        out_dir = out_path.parent
        out_dir.mkdir(parents=True, exist_ok=True)

        progress.append(info(f"Using converter: {converter}"))

        if converter == "libreoffice":
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(path),
            ]
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            # LibreOffice may exit 0 but still fail to produce a file
            lo_output = out_dir / (path.stem + ".pdf")
            if proc.returncode != 0 or not lo_output.exists():
                err = proc.stderr.strip() or proc.stdout.strip() or "No output produced"
                progress.append(fail("LibreOffice conversion failed", err[:100]))
                return _error(
                    f"LibreOffice conversion failed: {err}",
                    "Ensure LibreOffice is installed and the file is a valid .docx.",
                    progress,
                )
            if lo_output != out_path:
                lo_output.rename(out_path)

        elif converter == "word":
            try:
                import docx2pdf  # type: ignore[import-untyped]

                docx2pdf.convert(str(path), str(out_path))
            except Exception as e:
                progress.append(fail(f"Word conversion failed: {e}"))
                return _error(
                    f"Word conversion failed: {e}",
                    "Ensure Microsoft Word is installed and the file is valid.",
                    progress,
                )

        progress.append(ok("Exported to PDF", out_path.name))

        if open_after:
            from shared.platform_utils import open_file

            open_file(out_path)
            progress.append(ok("Opened PDF in default viewer"))

        result: dict[str, Any] = {
            "success": True,
            "op": "export_pdf",
            "input": str(path),
            "output": str(out_path),
            "output_name": out_path.name,
            "converter": converter,
            "progress": progress,
        }
        embed_content(result, out_path, return_content)
        result["token_estimate"] = len(str(result)) // 4
        return result
    except Exception as e:
        progress.append(fail(str(e)))
        return _error(
            str(e),
            "Check that the file is a valid .docx and a PDF converter is available.",
            progress,
        )
