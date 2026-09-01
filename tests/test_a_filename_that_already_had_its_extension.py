"""batch_create_from_template glued .docx onto a filename that already had it.

    data_list=[{..., "filename": "alice.docx"}], filename_key="filename"
      -> files: ["alice.docx.docx"]

`success: true`, `created_count: 2`, both files valid and correctly filled. A
caller who names the row's output the obvious way just gets a directory of
`*.docx.docx` -- and it is the value in `files` that a later read_document is
pointed at, so the wrong name is not only cosmetic.
"""

from __future__ import annotations

import docx

from docx_new import engine


def _template(path):
    d = docx.Document()
    d.add_paragraph("Hello {{name}}")
    d.save(str(path))
    return str(path)


def test_a_filename_that_already_ends_in_docx_is_not_doubled(tmp_path):
    tpl = _template(tmp_path / "tpl.docx")
    out = tmp_path / "out"

    result = engine.batch_create_from_template(
        tpl,
        [{"name": "Alice", "filename": "alice.docx"}],
        str(out),
        filename_key="filename",
    )

    assert result["success"] is True
    assert result["files"] == ["alice.docx"]
    assert (out / "alice.docx").exists()
    assert not (out / "alice.docx.docx").exists()


def test_a_filename_without_the_extension_still_gets_one(tmp_path):
    tpl = _template(tmp_path / "tpl.docx")
    out = tmp_path / "out"

    result = engine.batch_create_from_template(
        tpl,
        [{"name": "Bob", "filename": "bob"}],
        str(out),
        filename_key="filename",
    )

    assert result["files"] == ["bob.docx"]
    assert (out / "bob.docx").exists()


def test_the_match_is_case_insensitive(tmp_path):
    tpl = _template(tmp_path / "tpl.docx")
    out = tmp_path / "out"

    result = engine.batch_create_from_template(
        tpl,
        [{"name": "Cara", "filename": "cara.DOCX"}],
        str(out),
        filename_key="filename",
    )

    assert result["files"] == ["cara.docx"]


def test_only_a_trailing_docx_is_stripped(tmp_path):
    # "alice.docx.bak" is a name, not a name plus our extension. Stripping
    # anywhere but the end would rewrite names the caller meant.
    tpl = _template(tmp_path / "tpl.docx")
    out = tmp_path / "out"

    result = engine.batch_create_from_template(
        tpl,
        [{"name": "Dee", "filename": "notes.docx.bak"}],
        str(out),
        filename_key="filename",
    )

    assert result["files"] == ["notes.docx.bak.docx"]


def test_the_generated_names_are_unaffected(tmp_path):
    # No filename_key: the stem is ours already and must not change shape.
    tpl = _template(tmp_path / "tpl.docx")
    out = tmp_path / "out"

    result = engine.batch_create_from_template(tpl, [{"name": "Eve"}, {"name": "Fay"}], str(out))

    assert result["files"] == ["document_001.docx", "document_002.docx"]


def test_the_substitutions_still_land(tmp_path):
    # The rename must not disturb what the batch is for.
    tpl = _template(tmp_path / "tpl.docx")
    out = tmp_path / "out"

    engine.batch_create_from_template(
        tpl,
        [{"name": "Alice", "filename": "alice.docx"}],
        str(out),
        filename_key="filename",
    )

    text = "\n".join(p.text for p in docx.Document(str(out / "alice.docx")).paragraphs)
    assert "Alice" in text and "{{name}}" not in text
