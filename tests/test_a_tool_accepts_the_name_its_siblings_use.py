"""Three tools spelled a shared concept their own way, and phases stalled on them.

A census of every `@mcp.tool()` signature in this repo, grouped by what the
argument means:

    the sheet     sheet_name       28    source_sheet 1    old_name 1
    the slide     slide_index      10
    the paragraph paragraph_index   5    index        1

`slide_index` is the control -- ten tools, one spelling, and no sweep has ever
mis-called one. Every outlier has cost an attempt.

Watched live in round 8, a sweep model called

    rename_sheet(old_name="Sweep15Data", new_name="Sweep15Copy")   # tool 12

and then, two tools later, wrote the same `new_name` for copy_sheet, which
takes `new_sheet_name`:

    Error executing tool copy_sheet: 1 validation error for copy_sheetArguments
    new_sheet_name
      Field required [type=missing, ...]

pydantic refuses that before any server code runs, so the tool cannot name what
it wanted, and the live schemas carry no property descriptions -- the parameter
name is the entire contract. The phase reported "14/14 PASS" with the failure
in a notes column.

read_paragraph was the same shape with a tell in the source: the engine's own
out-of-range error read "paragraph_index N out of range" while the parameter
was called `index`. The error named an argument the tool did not take.

Renaming outright would fix the guess and break every existing caller, so each
outlier now accepts both spellings, canonical being whatever the majority of
its siblings use. The census runs as a test, so a new outlier fails the build
rather than waiting for a sweep to trip over it.
"""

from __future__ import annotations

import ast
import collections
from pathlib import Path

import openpyxl
import pytest
from docx import Document

SERVERS = Path(__file__).parent.parent / "servers"


def tool_defs(path: Path) -> list[ast.FunctionDef]:
    tree = ast.parse(path.read_text(encoding="utf-8"))
    return [n for n in ast.walk(tree) if isinstance(n, ast.FunctionDef) and n.decorator_list]


def every_tool() -> list[tuple[str, ast.FunctionDef]]:
    out = []
    for path in sorted(SERVERS.glob("*/*/server.py")):
        out.extend((path.parts[-3], fn) for fn in tool_defs(path))
    return out


@pytest.fixture()
def book(tmp_path: Path) -> str:
    wb = openpyxl.Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Data"
    ws["A1"] = "x"
    dst = tmp_path / "b.xlsx"
    wb.save(str(dst))
    wb.close()
    return str(dst)


@pytest.fixture()
def doc(tmp_path: Path) -> str:
    d = Document()
    d.add_paragraph("first")
    d.add_paragraph("second")
    dst = tmp_path / "d.docx"
    d.save(str(dst))
    return str(dst)


def sheets(path: str) -> list[str]:
    wb = openpyxl.load_workbook(path)
    names = list(wb.sheetnames)
    wb.close()
    return names


class TestCopySheetTakesEitherSpelling:
    """The call that actually failed mid-sweep."""

    def test_the_majority_spelling_works(self, book: str):
        from xlsx_basic.server import copy_sheet

        r = copy_sheet(book, sheet_name="Data", new_name="Copy")
        assert r["success"] is True, r.get("error")
        assert "Copy" in sheets(book)

    def test_the_original_spelling_still_works(self, book: str):
        from xlsx_basic.server import copy_sheet

        r = copy_sheet(book, source_sheet="Data", new_sheet_name="Copy")
        assert r["success"] is True, r.get("error")
        assert "Copy" in sheets(book)

    def test_the_two_can_be_mixed(self, book: str):
        from xlsx_basic.server import copy_sheet

        r = copy_sheet(book, source_sheet="Data", new_name="Copy")
        assert r["success"] is True, r.get("error")

    def test_the_alias_is_recorded_in_progress(self, book: str):
        from xlsx_basic.server import copy_sheet

        r = copy_sheet(book, source_sheet="Data", new_sheet_name="Copy")
        msgs = " ".join(str(p.get("msg", "")) for p in r["progress"])
        assert "alias" in msgs.lower(), r["progress"]

    def test_the_canonical_spelling_is_not_announced(self, book: str):
        from xlsx_basic.server import copy_sheet

        r = copy_sheet(book, sheet_name="Data", new_name="Copy")
        msgs = " ".join(str(p.get("msg", "")) for p in r["progress"])
        assert "alias" not in msgs.lower()

    @pytest.mark.parametrize(
        "kwargs,missing_name",
        [
            ({"new_name": "Copy"}, "sheet_name"),
            ({"sheet_name": "Data"}, "new_name"),
            ({}, "sheet_name"),
        ],
    )
    def test_neither_spelling_names_what_is_missing(self, book: str, kwargs: dict, missing_name: str):
        from xlsx_basic.server import copy_sheet

        r = copy_sheet(book, **kwargs)
        assert r["success"] is False
        assert missing_name in r["error"], r


class TestRenameSheetTakesEitherSpelling:
    def test_the_majority_spelling_works(self, book: str):
        from xlsx_basic.server import rename_sheet

        r = rename_sheet(book, sheet_name="Data", new_name="Renamed")
        assert r["success"] is True, r.get("error")
        assert sheets(book) == ["Renamed"]

    def test_the_original_spelling_still_works(self, book: str):
        from xlsx_basic.server import rename_sheet

        r = rename_sheet(book, old_name="Data", new_name="Renamed")
        assert r["success"] is True, r.get("error")
        assert sheets(book) == ["Renamed"]

    def test_new_name_is_still_required(self, book: str):
        """It is the one spelling both tools already agreed on."""
        import inspect

        from xlsx_basic.server import rename_sheet

        sig = inspect.signature(rename_sheet)
        assert sig.parameters["new_name"].default is inspect.Parameter.empty

    def test_no_sheet_at_all_says_sheet_name(self, book: str):
        from xlsx_basic.server import rename_sheet

        r = rename_sheet(book, new_name="Renamed")
        assert r["success"] is False and "sheet_name" in r["error"]


class TestReadParagraphTakesEitherSpelling:
    def test_the_majority_spelling_works(self, doc: str):
        from docx_basic.server import read_paragraph

        r = read_paragraph(doc, paragraph_index=1)
        assert r["success"] is True, r.get("error")
        assert r["text"] == "second"

    def test_the_original_spelling_still_works(self, doc: str):
        from docx_basic.server import read_paragraph

        r = read_paragraph(doc, index=1)
        assert r["success"] is True, r.get("error")
        assert r["text"] == "second"

    def test_paragraph_zero_is_reachable_under_both(self, doc: str):
        """0 is falsy; a truthiness check here would drop the first paragraph."""
        from docx_basic.server import read_paragraph

        assert read_paragraph(doc, paragraph_index=0)["text"] == "first"
        assert read_paragraph(doc, index=0)["text"] == "first"

    def test_neither_given_says_which_to_pass(self, doc: str):
        from docx_basic.server import read_paragraph

        r = read_paragraph(doc)
        assert r["success"] is False
        assert "paragraph_index" in r["error"] and "paragraph_index=0" in r["hint"]

    def test_out_of_range_still_reports_the_range(self, doc: str):
        from docx_basic.server import read_paragraph

        r = read_paragraph(doc, paragraph_index=99)
        assert r["success"] is False and "out of range" in r["error"]

    def test_the_error_names_a_real_parameter(self, doc: str):
        """The engine said "paragraph_index" while the tool took "index"."""
        import inspect

        from docx_basic.server import read_paragraph

        r = read_paragraph(doc, paragraph_index=99)
        params = set(inspect.signature(read_paragraph).parameters)
        named = {w.strip(".,'\"") for w in r["error"].split() if w.strip(".,'\"") in {"index", "paragraph_index"}}
        assert named and named <= params, (named, params)


class TestTheCensusHasNoNewOutliers:
    """What made these three findable. A concept spelled one way by many tools
    and another way by one or two is where a caller's guess goes wrong."""

    CONCEPTS = {
        "the sheet": {"sheet_name", "source_sheet", "old_name"},
        "the slide": {"slide_index", "slide_number", "slide"},
        "the paragraph": {"paragraph_index", "para_index", "index"},
    }

    def census(self, names: set[str]) -> dict[str, list[str]]:
        counts: dict[str, list[str]] = collections.defaultdict(list)
        for server, fn in every_tool():
            for arg in fn.args.args:
                if arg.arg in names:
                    counts[arg.arg].append(f"{server}.{fn.name}")
        return counts

    @pytest.mark.parametrize("concept", list(CONCEPTS))
    def test_every_minority_spelling_is_an_accepted_alias(self, concept: str):
        """A tool may use the minority spelling only if it also takes the
        majority one -- which is what these fixes did."""
        counts = self.census(self.CONCEPTS[concept])
        if not counts:
            pytest.skip(f"no tools use {concept}")
        majority = max(counts, key=lambda k: len(counts[k]))
        offenders = []
        for spelling, tools in counts.items():
            if spelling == majority:
                continue
            for qualified in tools:
                server, tool = qualified.split(".", 1)
                fn = next(f for s, f in every_tool() if s == server and f.name == tool)
                if majority not in {a.arg for a in fn.args.args}:
                    offenders.append(f"{qualified} takes {spelling} but not {majority}")
        assert not offenders, offenders

    def test_the_slide_stayed_consistent(self):
        """The control: if this ever splits, the same fix is needed there."""
        counts = self.census(self.CONCEPTS["the slide"])
        assert set(counts) == {"slide_index"}, counts

    def test_the_three_known_aliases_are_still_wired(self):
        wired = {
            ("xlsx_basic", "copy_sheet"): {"sheet_name", "source_sheet", "new_name", "new_sheet_name"},
            ("xlsx_basic", "rename_sheet"): {"sheet_name", "old_name", "new_name"},
            ("docx_basic", "read_paragraph"): {"paragraph_index", "index"},
        }
        for (server, tool), expected in wired.items():
            fn = next(f for s, f in every_tool() if s == server and f.name == tool)
            assert expected <= {a.arg for a in fn.args.args}, (server, tool)

    def test_the_docstrings_still_fit(self):
        for _server, fn in every_tool():
            doc = ast.get_docstring(fn) or ""
            assert len(doc) <= 80, (fn.name, len(doc), doc)
