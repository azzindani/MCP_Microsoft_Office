"""preserve_style=False did exactly what preserve_style=True does.

An AST walk over every @mcp.tool in the fleet, asking which declared parameters
the code never reads, found this one. `preserve_style` is declared on the tool,
forwarded by the wrapper, and referenced nowhere in the engine: the replacement
goes through `docxedit.replace_string`, which edits the text inside the
existing runs and offers no way not to.

So a caller asking for the replacement to take default formatting instead got
the original formatting, and was told the call succeeded. That is the same
failure the round-11 patch-op work was about, one level up: a knob the schema
advertises that can only ever be in one position.

Nothing here can implement the False case, so it is refused with the pair of
tools that actually restyle text. True -- the default, and what every existing
caller passes -- is untouched.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
for p in (str(ROOT), str(ROOT / "shared"), str(ROOT / "servers" / "docx_basic")):
    if p not in sys.path:
        sys.path.insert(0, p)

from docx_basic import engine  # noqa: E402


@pytest.fixture
def doc(tmp_path) -> Path:
    from docx import Document

    d = Document()
    d.add_paragraph("keep one and two")
    path = tmp_path / "d.docx"
    d.save(str(path))
    return path


def text_of(path: Path) -> str:
    from docx import Document

    return "\n".join(p.text for p in Document(str(path)).paragraphs)


class TestTheDefaultStillWorks:
    def test_a_replacement_lands(self, doc):
        r = engine.replace_text(str(doc), "one", "ONE")
        assert r["success"] is True, r.get("error")
        assert "ONE" in text_of(doc)

    def test_it_reports_how_many_it_replaced(self, doc):
        r = engine.replace_text(str(doc), "one", "ONE")
        assert r["replaced_count"] == 1, r

    def test_passing_the_flag_explicitly_changes_nothing(self, doc):
        r = engine.replace_text(str(doc), "one", "ONE", preserve_style=True)
        assert r["success"] is True, r.get("error")
        assert "ONE" in text_of(doc)


class TestTheImpossibleCaseIsRefused:
    def test_it_says_no_rather_than_pretending(self, doc):
        r = engine.replace_text(str(doc), "one", "ONE", preserve_style=False)
        assert r["success"] is False, "preserve_style=False reported success while preserving style"

    def test_the_document_is_left_alone(self, doc):
        before = text_of(doc)
        engine.replace_text(str(doc), "one", "ONE", preserve_style=False)
        assert text_of(doc) == before

    def test_the_hint_names_what_does_restyle(self, doc):
        r = engine.replace_text(str(doc), "one", "ONE", preserve_style=False)
        hint = r.get("hint", "")
        assert "set_font" in hint, hint
        assert "preserve_style=True" in hint, hint

    def test_no_snapshot_is_taken_for_a_refused_call(self, doc, tmp_path):
        engine.replace_text(str(doc), "one", "ONE", preserve_style=False)
        versions = tmp_path / ".mcp_versions"
        assert not versions.exists() or not list(versions.glob("*")), "a refused call left a snapshot"


class TestTheWrapperStillOffersIt:
    def test_the_parameter_is_declared(self):
        import inspect

        from docx_basic import server

        fn = getattr(server, "replace_text")
        fn = getattr(fn, "fn", fn)
        assert "preserve_style" in inspect.signature(fn).parameters

    def test_it_still_defaults_to_true(self):
        import inspect

        from docx_basic import server

        fn = getattr(server, "replace_text")
        fn = getattr(fn, "fn", fn)
        assert inspect.signature(fn).parameters["preserve_style"].default is True
