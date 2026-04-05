"""Tests for the docx_layout server engine."""

import shutil
from pathlib import Path

import pytest

FIXTURES = Path(__file__).parent / "fixtures"


def _copy(name: str, tmp_path: Path) -> Path:
    src = FIXTURES / name
    if not src.exists():
        pytest.skip(f"Fixture {name} not found.")
    dst = tmp_path / name
    shutil.copy(src, dst)
    return dst


# ---------------------------------------------------------------------------
# set_heading
# ---------------------------------------------------------------------------

class TestSetHeading:
    def test_applies_heading_style(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        path = _copy("contract_simple.docx", tmp_path)
        result = set_heading(str(path), 0, 1)
        assert result["success"] is True
        assert result["style"] == "Heading 1"
        assert "backup" in result
        assert result["progress"]

    def test_heading_style_persisted(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        from docx import Document
        path = _copy("contract_simple.docx", tmp_path)
        set_heading(str(path), 0, 2)
        doc = Document(str(path))
        assert "Heading" in doc.paragraphs[0].style.name

    def test_heading_levels_1_through_6(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        path = _copy("contract_simple.docx", tmp_path)
        for level in range(1, 7):
            result = set_heading(str(path), 0, level)
            assert result["success"] is True, f"Level {level} failed: {result}"

    def test_invalid_level_returns_error(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        path = _copy("contract_simple.docx", tmp_path)
        result = set_heading(str(path), 0, 7)
        assert result["success"] is False
        assert "7" in result["error"]

    def test_invalid_level_zero_returns_error(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        path = _copy("contract_simple.docx", tmp_path)
        result = set_heading(str(path), 0, 0)
        assert result["success"] is False

    def test_out_of_range_index(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        path = _copy("contract_simple.docx", tmp_path)
        result = set_heading(str(path), 9999, 1)
        assert result["success"] is False
        assert "range" in result["error"].lower()

    def test_file_not_found(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        result = set_heading(str(tmp_path / "ghost.docx"), 0, 1)
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_wrong_file_type(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        f = tmp_path / "test.xlsx"
        f.write_bytes(b"fake")
        result = set_heading(str(f), 0, 1)
        assert result["success"] is False


# ---------------------------------------------------------------------------
# set_font
# ---------------------------------------------------------------------------

class TestSetFont:
    def test_sets_font_name(self, tmp_path):
        from servers.docx_layout.engine import set_font
        from docx import Document
        path = _copy("contract_complex.docx", tmp_path)
        result = set_font(str(path), 0, font_name="Arial")
        assert result["success"] is True
        doc = Document(str(path))
        for run in doc.paragraphs[0].runs:
            assert run.font.name == "Arial"

    def test_sets_font_size(self, tmp_path):
        from servers.docx_layout.engine import set_font
        from docx import Document
        from docx.shared import Pt
        path = _copy("contract_complex.docx", tmp_path)
        result = set_font(str(path), 0, font_size=14)
        assert result["success"] is True
        doc = Document(str(path))
        for run in doc.paragraphs[0].runs:
            assert run.font.size == Pt(14)

    def test_sets_bold(self, tmp_path):
        from servers.docx_layout.engine import set_font
        from docx import Document
        path = _copy("contract_simple.docx", tmp_path)
        result = set_font(str(path), 0, bold=True)
        assert result["success"] is True
        doc = Document(str(path))
        for run in doc.paragraphs[0].runs:
            assert run.bold is True

    def test_sets_italic(self, tmp_path):
        from servers.docx_layout.engine import set_font
        from docx import Document
        path = _copy("contract_simple.docx", tmp_path)
        result = set_font(str(path), 0, italic=True)
        assert result["success"] is True
        doc = Document(str(path))
        for run in doc.paragraphs[0].runs:
            assert run.italic is True

    def test_no_changes_still_succeeds(self, tmp_path):
        from servers.docx_layout.engine import set_font
        path = _copy("contract_simple.docx", tmp_path)
        result = set_font(str(path), 0)
        assert result["success"] is True

    def test_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import set_font
        path = _copy("contract_simple.docx", tmp_path)
        result = set_font(str(path), 0, font_name="Times New Roman")
        assert result["success"] is True
        assert "backup" in result
        assert Path(result["backup"]).exists()

    def test_out_of_range_index(self, tmp_path):
        from servers.docx_layout.engine import set_font
        path = _copy("contract_simple.docx", tmp_path)
        result = set_font(str(path), 99999, font_name="Arial")
        assert result["success"] is False

    def test_file_not_found(self, tmp_path):
        from servers.docx_layout.engine import set_font
        result = set_font(str(tmp_path / "ghost.docx"), 0, font_name="Arial")
        assert result["success"] is False


# ---------------------------------------------------------------------------
# set_paragraph_style
# ---------------------------------------------------------------------------

class TestSetParagraphStyle:
    def test_applies_normal_style(self, tmp_path):
        from servers.docx_layout.engine import set_paragraph_style
        path = _copy("contract_simple.docx", tmp_path)
        result = set_paragraph_style(str(path), 0, "Normal")
        assert result["success"] is True
        assert result["style"] == "Normal"

    def test_style_persisted(self, tmp_path):
        from servers.docx_layout.engine import set_paragraph_style
        from docx import Document
        path = _copy("contract_simple.docx", tmp_path)
        set_paragraph_style(str(path), 0, "Normal")
        doc = Document(str(path))
        assert doc.paragraphs[0].style.name == "Normal"

    def test_invalid_style_returns_error(self, tmp_path):
        from servers.docx_layout.engine import set_paragraph_style
        path = _copy("contract_simple.docx", tmp_path)
        result = set_paragraph_style(str(path), 0, "XYZZY_INVALID_STYLE")
        assert result["success"] is False
        assert "XYZZY_INVALID_STYLE" in result["error"]

    def test_out_of_range_index(self, tmp_path):
        from servers.docx_layout.engine import set_paragraph_style
        path = _copy("contract_simple.docx", tmp_path)
        result = set_paragraph_style(str(path), 9999, "Normal")
        assert result["success"] is False

    def test_file_not_found(self, tmp_path):
        from servers.docx_layout.engine import set_paragraph_style
        result = set_paragraph_style(str(tmp_path / "ghost.docx"), 0, "Normal")
        assert result["success"] is False

    def test_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import set_paragraph_style
        path = _copy("contract_simple.docx", tmp_path)
        result = set_paragraph_style(str(path), 0, "Normal")
        assert result["success"] is True
        assert "backup" in result
        assert Path(result["backup"]).exists()


# ---------------------------------------------------------------------------
# add_image
# ---------------------------------------------------------------------------

class TestAddImage:
    def test_invalid_image_path_returns_error(self, tmp_path):
        from servers.docx_layout.engine import add_image
        path = _copy("contract_simple.docx", tmp_path)
        result = add_image(str(path), 0, str(tmp_path / "nonexistent.png"))
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_unsupported_image_format(self, tmp_path):
        from servers.docx_layout.engine import add_image
        path = _copy("contract_simple.docx", tmp_path)
        fake_img = tmp_path / "image.xyz"
        fake_img.write_bytes(b"fake")
        result = add_image(str(path), 0, str(fake_img))
        assert result["success"] is False
        assert "unsupported" in result["error"].lower()

    def test_file_not_found(self, tmp_path):
        from servers.docx_layout.engine import add_image
        result = add_image(str(tmp_path / "ghost.docx"), 0, str(tmp_path / "img.png"))
        assert result["success"] is False

    def test_inserts_real_image(self, tmp_path):
        """If a valid PNG is available, verify insertion succeeds."""
        from servers.docx_layout.engine import add_image
        from docx import Document

        # Create a minimal 1x1 PNG
        import struct, zlib
        def _minimal_png() -> bytes:
            sig = b"\x89PNG\r\n\x1a\n"
            ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
            ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data)
            ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
            idat_data = zlib.compress(b"\x00\xff\xff\xff")
            idat_crc = zlib.crc32(b"IDAT" + idat_data)
            idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + struct.pack(">I", idat_crc)
            iend_crc = zlib.crc32(b"IEND")
            iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
            return sig + ihdr + idat + iend

        img_path = tmp_path / "test_image.png"
        img_path.write_bytes(_minimal_png())

        doc_path = _copy("contract_simple.docx", tmp_path)
        result = add_image(str(doc_path), 0, str(img_path), width_inches=2.0)
        assert result["success"] is True
        assert result["image"] == "test_image.png"
        assert "backup" in result


# ---------------------------------------------------------------------------
# set_page_margins
# ---------------------------------------------------------------------------

class TestSetPageMargins:
    def test_sets_margins(self, tmp_path):
        from servers.docx_layout.engine import set_page_margins
        from docx import Document
        from docx.shared import Cm
        path = _copy("contract_simple.docx", tmp_path)
        result = set_page_margins(str(path), top=3.0, bottom=2.0, left=2.5, right=2.5)
        assert result["success"] is True
        assert result["top_cm"] == 3.0
        doc = Document(str(path))
        # Allow ±500 EMU tolerance for floating-point rounding in unit conversion
        assert abs(doc.sections[0].top_margin - Cm(3.0)) < 500
        assert abs(doc.sections[0].bottom_margin - Cm(2.0)) < 500

    def test_default_margins(self, tmp_path):
        from servers.docx_layout.engine import set_page_margins
        path = _copy("contract_simple.docx", tmp_path)
        result = set_page_margins(str(path))
        assert result["success"] is True
        assert result["top_cm"] == 2.54

    def test_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import set_page_margins
        path = _copy("contract_simple.docx", tmp_path)
        result = set_page_margins(str(path), top=3.0)
        assert result["success"] is True
        assert "backup" in result
        assert Path(result["backup"]).exists()

    def test_file_not_found(self, tmp_path):
        from servers.docx_layout.engine import set_page_margins
        result = set_page_margins(str(tmp_path / "ghost.docx"))
        assert result["success"] is False


# ---------------------------------------------------------------------------
# add_header_footer
# ---------------------------------------------------------------------------

class TestAddHeaderFooter:
    def test_add_header_text(self, tmp_path):
        from servers.docx_layout.engine import add_header_footer
        from docx import Document
        path = _copy("contract_simple.docx", tmp_path)
        result = add_header_footer(str(path), "My Header", location="header")
        assert result["success"] is True
        assert result["location"] == "header"
        doc = Document(str(path))
        assert doc.sections[0].header.paragraphs[0].text == "My Header"

    def test_add_footer_text(self, tmp_path):
        from servers.docx_layout.engine import add_header_footer
        from docx import Document
        path = _copy("contract_simple.docx", tmp_path)
        result = add_header_footer(str(path), "Page Footer", location="footer")
        assert result["success"] is True
        assert result["location"] == "footer"
        doc = Document(str(path))
        assert doc.sections[0].footer.paragraphs[0].text == "Page Footer"

    def test_invalid_location_returns_error(self, tmp_path):
        from servers.docx_layout.engine import add_header_footer
        path = _copy("contract_simple.docx", tmp_path)
        result = add_header_footer(str(path), "text", location="sidebar")
        assert result["success"] is False
        assert "header" in result["hint"].lower() or "footer" in result["hint"].lower()

    def test_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import add_header_footer
        path = _copy("contract_simple.docx", tmp_path)
        result = add_header_footer(str(path), "Header text")
        assert result["success"] is True
        assert "backup" in result
        assert Path(result["backup"]).exists()

    def test_file_not_found(self, tmp_path):
        from servers.docx_layout.engine import add_header_footer
        result = add_header_footer(str(tmp_path / "ghost.docx"), "text")
        assert result["success"] is False


# ---------------------------------------------------------------------------
# export_pdf
# ---------------------------------------------------------------------------

class TestExportPdf:
    def test_no_converter_returns_graceful_error(self, tmp_path):
        """When no PDF converter is available, return success=False with hint."""
        from unittest.mock import patch
        from servers.docx_layout.engine import export_pdf

        path = _copy("contract_simple.docx", tmp_path)

        # Patch where the engine module looks up get_pdf_converter
        with patch("servers.docx_layout.engine.get_pdf_converter", return_value=None):
            result = export_pdf(str(path))

        assert result["success"] is False
        assert "libreoffice" in result["hint"].lower() or "word" in result["hint"].lower()
        assert result["progress"]

    def test_file_not_found(self, tmp_path):
        from servers.docx_layout.engine import export_pdf
        result = export_pdf(str(tmp_path / "ghost.docx"))
        assert result["success"] is False

    def test_wrong_file_type(self, tmp_path):
        from servers.docx_layout.engine import export_pdf
        f = tmp_path / "test.xlsx"
        f.write_bytes(b"fake")
        result = export_pdf(str(f))
        assert result["success"] is False

    def test_export_with_libreoffice(self, tmp_path):
        """Test LibreOffice export; skip if libreoffice unavailable or conversion fails."""
        import shutil
        if shutil.which("libreoffice") is None and shutil.which("soffice") is None:
            pytest.skip("LibreOffice not installed")

        from servers.docx_layout.engine import export_pdf
        path = _copy("contract_simple.docx", tmp_path)
        out_pdf = tmp_path / "output.pdf"
        result = export_pdf(str(path), str(out_pdf))

        if not result["success"]:
            # LibreOffice may be installed but non-functional in headless CI
            pytest.skip(f"LibreOffice conversion failed: {result.get('error', '')}")

        assert Path(result["output"]).exists()


# ---------------------------------------------------------------------------
# Snapshot creation — all write tools
# ---------------------------------------------------------------------------

class TestAllWriteToolsCreateSnapshot:
    def test_set_heading_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        path = _copy("contract_simple.docx", tmp_path)
        result = set_heading(str(path), 0, 1)
        assert result["success"] is True
        assert Path(result["backup"]).exists()

    def test_set_font_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import set_font
        path = _copy("contract_simple.docx", tmp_path)
        result = set_font(str(path), 0, font_name="Arial")
        assert result["success"] is True
        assert Path(result["backup"]).exists()

    def test_set_paragraph_style_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import set_paragraph_style
        path = _copy("contract_simple.docx", tmp_path)
        result = set_paragraph_style(str(path), 0, "Normal")
        assert result["success"] is True
        assert Path(result["backup"]).exists()

    def test_set_page_margins_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import set_page_margins
        path = _copy("contract_simple.docx", tmp_path)
        result = set_page_margins(str(path), top=3.0)
        assert result["success"] is True
        assert Path(result["backup"]).exists()

    def test_add_header_footer_creates_snapshot(self, tmp_path):
        from servers.docx_layout.engine import add_header_footer
        path = _copy("contract_simple.docx", tmp_path)
        result = add_header_footer(str(path), "Test header")
        assert result["success"] is True
        assert Path(result["backup"]).exists()


# ---------------------------------------------------------------------------
# Progress field present in all responses
# ---------------------------------------------------------------------------

class TestProgressField:
    def test_set_heading_has_progress(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        path = _copy("contract_simple.docx", tmp_path)
        result = set_heading(str(path), 0, 1)
        assert "progress" in result
        assert isinstance(result["progress"], list)

    def test_error_response_has_progress(self, tmp_path):
        from servers.docx_layout.engine import set_heading
        result = set_heading(str(tmp_path / "ghost.docx"), 0, 1)
        assert "progress" in result
        assert isinstance(result["progress"], list)
        assert len(result["progress"]) > 0
