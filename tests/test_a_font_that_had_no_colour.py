"""set_font could set a name, a size, bold and italic. Nothing else.

Between them, colour and paragraph spacing are most of what "format this so it
reads well" means -- a navy heading, body set a little open, air between
sections. None of it could be asked for, so a model rebuilding an unreadable
brief set every one of those properties in python-docx instead.

add_header_footer had the matching gap: text only, which renders as 11pt black
on the left in every document, and no way at all to put a page number in a
footer. Writing the number as literal text is the obvious alternative and is
wrong on every page but the first, which is why the hand-written version
carried a "page number field will be auto, placeholder" comment and shipped
without one.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

import pytest

from docx_layout import engine as docx_layout  # type: ignore[reportMissingImports]

NAVY = "0B1D3A"


@pytest.fixture
def doc(tmp_path, fixtures_dir):
    dest = tmp_path / "contract.docx"
    shutil.copy(fixtures_dir / "contract_simple.docx", dest)
    return dest


def part(path: Path, name: str) -> str:
    zf = zipfile.ZipFile(path)
    return zf.read(name).decode() if name in zf.namelist() else ""


# --- set_font ---------------------------------------------------------------


def test_a_paragraph_can_be_given_a_colour(doc):
    r = docx_layout.set_font(str(doc), 0, color=NAVY)
    assert r["success"] is True
    assert any("color" in c for c in r["changes"])
    assert re.search(rf'<w:color w:val="{NAVY}"', part(doc, "word/document.xml"))


def test_line_spacing_and_space_after_are_applied(doc):
    r = docx_layout.set_font(str(doc), 0, line_spacing=1.5, space_after=12)
    assert r["success"] is True
    assert "line_spacing" in r["changes"]
    assert "space_after" in r["changes"]

    from docx import Document

    para = Document(str(doc)).paragraphs[0]
    assert para.paragraph_format.line_spacing == 1.5
    assert para.paragraph_format.space_after.pt == 12


def test_the_old_arguments_still_behave(doc):
    r = docx_layout.set_font(str(doc), 0, font_name="Calibri", font_size=14, bold="true")
    assert r["success"] is True

    from docx import Document

    run = Document(str(doc)).paragraphs[0].runs[0]
    assert run.font.name == "Calibri"
    assert run.font.size.pt == 14
    assert run.bold is True


def test_a_paragraph_that_is_entirely_a_hyperlink_is_not_a_silent_no_op(tmp_path):
    """`paragraph.runs` returns direct children only.

    A paragraph whose text sits inside a w:hyperlink reports zero runs while
    `paragraph.text` reads fine, so every run-level change applied to nothing
    and still came back successful.
    """
    from docx import Document
    from docx.oxml import OxmlElement

    path = tmp_path / "linked.docx"
    document = Document()
    para = document.add_paragraph("")
    link = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Linked text"
    run.append(text)
    link.append(run)
    para._p.append(link)
    document.save(str(path))

    assert len(Document(str(path)).paragraphs[0].runs) == 0
    assert Document(str(path)).paragraphs[0].text == "Linked text"

    r = docx_layout.set_font(str(path), 0, color=NAVY, bold="true")
    assert r["success"] is True
    assert re.search(rf'<w:color w:val="{NAVY}"', part(path, "word/document.xml"))
    # The text must not be duplicated by a run added alongside the hyperlink.
    assert Document(str(path)).paragraphs[0].text == "Linked text"


def test_a_bad_colour_is_refused_before_the_snapshot(doc):
    r = docx_layout.set_font(str(doc), 0, color="#GGGGGG")
    assert r["success"] is False
    assert "GGGGGG" in r["error"]
    assert r.get("backup") in (None, "")


def test_negative_spacing_is_refused(doc):
    r = docx_layout.set_font(str(doc), 0, line_spacing=-1)
    assert r["success"] is False
    assert "negative" in r["error"]


# --- add_header_footer ------------------------------------------------------


def test_a_footer_can_number_its_own_pages(doc):
    r = docx_layout.add_header_footer(str(doc), "Executive Brief", location="footer", page_numbers=True)
    assert r["success"] is True
    assert r["page_numbers"] is True
    footer = part(doc, "word/footer1.xml")
    assert "PAGE" in footer, "a live field, not a literal number"
    assert 'w:fldCharType="begin"' in footer


def test_a_header_can_be_styled(doc):
    r = docx_layout.add_header_footer(
        str(doc), "BOARD CONFIDENTIAL", location="header", font_size=7, color="595959", align="right"
    )
    assert r["success"] is True
    header = part(doc, "word/header1.xml")
    assert '<w:color w:val="595959"' in header
    assert "<w:sz " in header


def test_plain_text_still_works(doc):
    r = docx_layout.add_header_footer(str(doc), "Just text", location="header")
    assert r["success"] is True
    assert "Just text" in part(doc, "word/header1.xml")


def test_replacing_a_long_header_leaves_nothing_behind(doc):
    docx_layout.add_header_footer(str(doc), "A much longer original header line", location="header")
    docx_layout.add_header_footer(str(doc), "Short", location="header")
    header = part(doc, "word/header1.xml")
    assert "Short" in header
    assert "much longer" not in header


def test_a_bad_alignment_is_refused_before_the_snapshot(doc):
    r = docx_layout.add_header_footer(str(doc), "x", location="header", align="middle")
    assert r["success"] is False
    assert "middle" in r["error"]
    assert r.get("backup") in (None, "")


def test_a_bad_location_is_still_refused(doc):
    r = docx_layout.add_header_footer(str(doc), "x", location="sidebar")
    assert r["success"] is False
    assert "sidebar" in r["error"]
