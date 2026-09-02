"""Nothing in docx_tables could change how a cell looks, only what it says.

A shaded header row and banded body rows are most of what separates a table a
director scans from a grid of numbers, and both need the `w:shd` element that
python-docx does not expose. So the capability was absent, and the model asked
to make an executive brief readable installed python-docx and wrote its own
`set_cell_shading` helper.

`set_cell_style` addresses by range rather than by cell on purpose: one call
per cell would make a 6x4 table 24 writes, 24 snapshots and 24 receipts to
apply two colours. `row=0` shades a header; `band_fill` alone stripes alternate
body rows and deliberately leaves the header out of the banding, because a
striped header reads as data.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

import pytest

from docx_tables import engine as docx_tables  # type: ignore[reportMissingImports]

NAVY = "0B1D3A"
CREAM = "F5F1E8"


@pytest.fixture
def table_doc(tmp_path, fixtures_dir):
    src = fixtures_dir / "report_tables.docx"
    dest = tmp_path / "report.docx"
    shutil.copy(src, dest)
    return dest


def fills(path: Path) -> list[str]:
    xml = zipfile.ZipFile(path).read("word/document.xml").decode()
    return re.findall(r'w:fill="([0-9A-F]{6})"', xml)


def test_a_header_row_can_be_shaded(table_doc):
    before = fills(table_doc).count(NAVY)
    r = docx_tables.set_cell_style(str(table_doc), 0, fill=NAVY, color="FFFFFF", bold="true", row=0)
    assert r["success"] is True
    assert r["cells_styled"] >= 1
    assert fills(table_doc).count(NAVY) > before


def test_banding_stripes_the_body_and_leaves_the_header_alone(table_doc):
    from docx import Document

    rows = len(Document(str(table_doc)).tables[0].rows)
    r = docx_tables.set_cell_style(str(table_doc), 0, band_fill=CREAM)
    assert r["success"] is True
    # Rows 2, 4, 6 ... of a header + body table.
    assert r["rows_banded"] == (rows - 1) // 2

    doc = Document(str(table_doc))
    header_cell = doc.tables[0].rows[0].cells[0]
    header_xml = header_cell._tc.xml
    assert CREAM not in header_xml, "the header must never be banded"


def test_a_single_cell_can_be_targeted(table_doc):
    r = docx_tables.set_cell_style(str(table_doc), 0, fill=NAVY, row=1, col=0)
    assert r["success"] is True
    assert r["cells_styled"] == 1


def test_a_whole_column_can_be_targeted(table_doc):
    from docx import Document

    rows = len(Document(str(table_doc)).tables[0].rows)
    r = docx_tables.set_cell_style(str(table_doc), 0, bold="true", col=0)
    assert r["success"] is True
    assert r["cells_styled"] == rows


def test_a_write_leaves_a_snapshot_and_a_receipt(table_doc):
    r = docx_tables.set_cell_style(str(table_doc), 0, fill=NAVY, row=0)
    assert r["backup"], "every write must snapshot first"
    assert Path(r["backup"]).exists()


# --- refusals ---------------------------------------------------------------


def test_a_call_that_asks_for_no_style_is_refused(table_doc):
    r = docx_tables.set_cell_style(str(table_doc), 0)
    assert r["success"] is False
    assert "fill" in r["error"]
    assert "0B1D3A" in r["hint"], "the hint should show a working call"


def test_a_colour_that_is_not_hex_is_named(table_doc):
    r = docx_tables.set_cell_style(str(table_doc), 0, fill="navy")
    assert r["success"] is False
    assert "navy" in r["error"]
    assert "hex" in r["error"].lower() or "hex" in r["hint"].lower()


def test_a_row_beyond_the_table_is_refused(table_doc):
    r = docx_tables.set_cell_style(str(table_doc), 0, fill=NAVY, row=999)
    assert r["success"] is False
    assert "999" in r["error"]


def test_a_missing_table_is_refused(table_doc):
    r = docx_tables.set_cell_style(str(table_doc), 99, fill=NAVY)
    assert r["success"] is False
    assert "99" in r["error"]


def test_a_missing_file_is_refused(tmp_path):
    r = docx_tables.set_cell_style(str(tmp_path / "nope.docx"), 0, fill=NAVY)
    assert r["success"] is False
    assert r["hint"]
