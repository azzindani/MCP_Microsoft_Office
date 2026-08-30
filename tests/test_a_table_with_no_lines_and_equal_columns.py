"""add_table produced a borderless table with equal columns, and it was unreadable.

Found by rendering a generated report and LOOKING at it, which is the only way
this class of defect is ever found. Every check that had been made of the
document passed: the file opened, the tables were present, `read_table` returned
every cell. The pages were unusable.

Two faults, both in the layout rather than the data.

**No rules.** `doc.add_table(rows, cols)` applies python-docx's default style,
"Normal Table", which draws no borders at all. Word then shows columns of text
floating in whitespace, and a table without rules reads as a table only where
the columns happen to line up. It is now created with "Table Grid".

**Equal columns.** python-docx writes a `tblGrid` splitting the page into equal
fractions, so a column holding
`DecreaseIncreaseInPlacementsWithOtherBanksAndBankIndonesia` got the same third
of the page as one holding `4.019`, and the identifier broke mid-word across
three lines:

    DecreaseIncreaseInPla | CurrentYearDuration | 169.400
    cementsWithOtherBan   |                     |
    ksAndBankIndonesia    |                     |

`table.autofit = True` is the obvious fix and does not work -- it sets
`tblLayout`, and both Word and LibreOffice still lay the table out on the equal
`gridCol` widths. Explicit widths are honoured by both, so columns are now sized
from their content.

**The weighting is damped, and that is the part worth keeping.** Sharing width
out in straight proportion to the longest cell over-feeds the wide column: the
same three columns came out 3.21in / 1.52in / 1.26in, which then broke
`CurrentYearDuration` and `1.640.830.566` across lines. That trade is a loss --
a wrapped identifier is still readable, a wrapped figure invites a misread. On
the square root of the length they come out 2.93in / 1.68in / 1.39in and every
figure fits on one line.

Verified by rendering to PDF and looking at the pixels, not by asserting the
cells exist.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "servers" / "docx_tables"))

from docx_tables import engine  # noqa: E402

# The three columns from the reconciliation this was found on: a long
# identifier, a medium label, and a figure that must never be broken.
ROWS = [
    ["fact", "context", "millions Rp"],
    ["DecreaseIncreaseInPlacementsWithOtherBanksAndBankIndonesia", "CurrentYearDuration", "169.400"],
    ["Assets", "CurrentYearInstant", "1.640.830.566"],
    ["DecreaseIncreaseInLoans", "PriorYearDuration", "-17.906.497"],
]


@pytest.fixture
def document(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("anchor")
    path = tmp_path / "report.docx"
    doc.save(str(path))
    return str(path)


def _table(path):
    from docx import Document

    return Document(path).tables[0]


class TestTheTableHasLines:
    def test_a_created_table_is_bordered(self, document):
        result = engine.add_table(document, 0, len(ROWS), 3, ROWS)
        assert result["success"] is True
        assert _table(document).style.name == "Table Grid"

    def test_the_cells_still_hold_what_was_passed(self, document):
        """Styling must not disturb the data it is applied to."""
        engine.add_table(document, 0, len(ROWS), 3, ROWS)
        table = _table(document)
        assert [c.text for c in table.rows[0].cells] == ROWS[0]
        assert table.rows[2].cells[2].text == "1.640.830.566"


class TestTheColumnsAreNotEqual:
    def test_a_long_column_gets_more_room_than_a_short_one(self, document):
        engine.add_table(document, 0, len(ROWS), 3, ROWS)
        widths = [c.width.inches for c in _table(document).columns]
        assert widths[0] > widths[1] > widths[2], f"columns are not content-sized: {widths}"

    def test_the_widest_column_does_not_starve_the_others(self, document):
        """The damping. Straight proportion breaks the figures instead.

        `1.640.830.566` is 13 characters; a column under about 1.2in breaks it
        across two lines, which is the failure this trade-off exists to avoid.
        """
        engine.add_table(document, 0, len(ROWS), 3, ROWS)
        widths = [c.width.inches for c in _table(document).columns]
        assert widths[2] >= 1.2, f"the figure column is too narrow at {widths[2]:.2f}in"
        assert widths[1] >= 1.5, f"the context column is too narrow at {widths[1]:.2f}in"

    def test_the_row_fits_the_printable_width(self, document):
        """A column set past the margin runs off the page instead of wrapping."""
        from docx import Document
        from docx.shared import Emu

        engine.add_table(document, 0, len(ROWS), 3, ROWS)
        doc = Document(document)
        section = doc.sections[0]
        # Subtracting Length objects yields a plain int of EMU, not a Length.
        usable = Emu(int(section.page_width) - int(section.left_margin) - int(section.right_margin)).inches
        assert sum(c.width.inches for c in doc.tables[0].columns) <= usable + 0.01

    def test_every_cell_carries_the_width_not_only_the_column(self, document):
        """Word reads per-cell `tcW` and treats `gridCol` as a hint.

        Setting the column alone looks right in LibreOffice and changes nothing
        in Word, which is the worse half of the failure to leave in place.
        """
        engine.add_table(document, 0, len(ROWS), 3, ROWS)
        table = _table(document)
        for index, column in enumerate(table.columns):
            for cell in column.cells:
                assert cell.width == column.width, f"cell in column {index} has no width of its own"


class TestATableWithNoDataIsUnchanged:
    def test_an_empty_table_is_still_created_and_bordered(self, document):
        """No data means no content to size from; it must not raise."""
        result = engine.add_table(document, 0, 3, 2, None)
        assert result["success"] is True
        table = _table(document)
        assert table.style.name == "Table Grid"
        assert len(table.rows) == 3
