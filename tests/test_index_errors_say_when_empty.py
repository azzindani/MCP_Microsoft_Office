"""An empty document reported its valid index range as "(0--1)".

Twenty-seven call sites across six engines hand-built the range as
f"(0-{total - 1})". When the collection is empty that renders as

    Table index 0 out of range (0--1)

A coverage sweep hit a freshly created document and got exactly that from
read_table, read_table_row, set_cell and add_row in a row. "0--1" is not a
range a caller can act on: it cannot tell whether the index was wrong or the
document simply has no tables, so it has no way to reach the one move that
helps -- add_table(). Four calls were spent on it.

index_range() now says "— this file has no tables" for an empty collection and
is unchanged, "(0-3)", for a populated one. The docx_tables hints follow: with
no tables, "Use list_tables to see available tables" only sends the caller to
look at an empty list, so it becomes "Add one with add_table() first."
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document
from pptx import Presentation

from shared.progress import index_range


class TestIndexRangeSpeaksPlainly:
    def test_a_populated_collection_reads_as_before(self):
        assert index_range(4, "tables") == "(0-3)"

    def test_a_single_item_is_a_degenerate_but_valid_range(self):
        assert index_range(1, "slides") == "(0-0)"

    def test_an_empty_collection_says_so(self):
        assert index_range(0, "tables") == "— this file has no tables"

    def test_an_empty_collection_never_prints_a_negative_range(self):
        assert "-1" not in index_range(0, "charts")

    def test_the_noun_reaches_the_reader(self):
        assert "slides" in index_range(0, "slides")
        assert "charts" in index_range(0, "charts")


@pytest.fixture()
def empty_docx(tmp_path: Path) -> str:
    p = tmp_path / "empty.docx"
    Document().save(str(p))
    return str(p)


@pytest.fixture()
def empty_pptx(tmp_path: Path) -> str:
    p = tmp_path / "empty.pptx"
    Presentation().save(str(p))
    return str(p)


class TestTheToolsThatBurnedTheSweepsCalls:
    """read_table, read_table_row, set_cell and add_row, on a document with no
    tables -- the four that answered "(0--1)" one after another."""

    def _calls(self, path: str):
        from docx_tables import engine

        return [
            ("read_table", engine.read_table(path, 0)),
            ("read_table_row", engine.read_table_row(path, 0, 0)),
            ("set_cell", engine.set_cell(path, 0, 0, 0, "x")),
            ("add_row", engine.add_row(path, 0, ["x"])),
        ]

    def test_none_of_them_print_a_malformed_range(self, empty_docx: str):
        for name, r in self._calls(empty_docx):
            assert "(0--1)" not in str(r.get("error")), name

    def test_each_says_the_document_has_no_tables(self, empty_docx: str):
        for name, r in self._calls(empty_docx):
            assert "no tables" in str(r.get("error")), f"{name}: {r.get('error')}"

    def test_each_names_the_move_that_would_help(self, empty_docx: str):
        for name, r in self._calls(empty_docx):
            assert "add_table()" in str(r.get("hint")), f"{name}: {r.get('hint')}"

    def test_they_still_fail(self, empty_docx: str):
        """Clearer wording, same verdict."""
        for name, r in self._calls(empty_docx):
            assert r["success"] is False, name


class TestAPopulatedDocumentIsUnaffected:
    def test_a_genuinely_bad_index_still_gets_the_range(self, tmp_path: Path):
        from docx_tables import engine

        p = tmp_path / "one_table.docx"
        doc = Document()
        doc.add_table(rows=2, cols=3)
        doc.save(str(p))

        r = engine.read_table(str(p), 5)
        assert r["success"] is False
        assert "(0-0)" in r["error"]
        assert "list_tables" in r["hint"]

    def test_a_valid_index_still_works(self, tmp_path: Path):
        from docx_tables import engine

        p = tmp_path / "one_table.docx"
        doc = Document()
        doc.add_table(rows=2, cols=3)
        doc.save(str(p))
        assert engine.read_table(str(p), 0)["success"] is True


class TestTheOtherEnginesGotTheSameTreatment:
    def test_a_deck_with_no_slides_says_so(self, empty_pptx: str):
        from pptx_basic import engine

        r = engine.read_slide(empty_pptx, 0)
        assert r["success"] is False
        assert "no slides" in r["error"]
        assert "(0--1)" not in r["error"]

    def test_no_engine_still_builds_the_range_by_hand(self):
        """The point of a shared helper is that nobody keeps a private copy."""
        root = Path(__file__).resolve().parents[1] / "servers"
        offenders = [
            str(p.relative_to(root))
            for p in root.rglob("engine.py")
            if re.search(r"out of range \(0-\{", p.read_text(encoding="utf-8"))
        ]
        assert not offenders, f"these still hand-build the range: {offenders}"
