"""create_from_template meant three different things on three sibling servers.

A coverage sweep called it on docx-new with only template_path -- reasonable,
since copying a template unchanged is a real use -- and got:

    1 validation error for create_from_templateArguments
    substitutions
      Field required [type=missing, ...]

The three tiers disagreed:

    docx-new   create_from_template  required = [template_path, substitutions]
    xlsx-new   create_from_template  required = [template_path, replacements]
    pptx-new   create_from_template  required = [template_path]

pptx already proves that "copy this template" is a complete request, and both
engines iterate the mapping without mutating it, so an empty one is meaningful:
copy the file, substitute nothing. CLAUDE.md agrees -- "For optional parameters,
use a primitive default value" -- so requiring it was the outlier, not pptx.

Follow-up: the mapping was also named differently per tier -- `substitutions` on
docx, `replacements` on xlsx, and absent entirely on pptx. All three now take
`substitutions`, and pptx actually applies them instead of only copying the file.
"""

from __future__ import annotations

import inspect
from pathlib import Path

import pytest

from docx_new import server as docx_server
from pptx_new import server as pptx_server
from xlsx_new import server as xlsx_server


def _tool_fn(tool):
    """The @mcp.tool() decorator returns the plain function on this repo and a
    FunctionTool wrapper on others. Accept either, so this test pins behaviour
    rather than the decorator's return type."""
    return getattr(tool, "fn", tool)


TIERS = [
    pytest.param(docx_server, "substitutions", id="docx"),
    pytest.param(xlsx_server, "substitutions", id="xlsx"),
    pytest.param(pptx_server, "substitutions", id="pptx"),
]


class TestTheMappingIsOptionalEverywhere:
    @pytest.mark.parametrize("module,param", TIERS)
    def test_it_has_a_default(self, module, param: str):
        sig = inspect.signature(_tool_fn(module.create_from_template))
        assert sig.parameters[param].default is not inspect.Parameter.empty, (
            f"{param} is required, so 'just copy the template' is impossible on this tier"
        )

    @pytest.mark.parametrize("module,param", TIERS)
    def test_the_default_is_an_empty_mapping(self, module, param: str):
        sig = inspect.signature(_tool_fn(module.create_from_template))
        assert sig.parameters[param].default == {}

    def test_only_the_template_path_is_ever_required(self):
        """Copying a template unchanged is a complete request on every tier."""
        for module in (docx_server, xlsx_server, pptx_server):
            sig = inspect.signature(_tool_fn(module.create_from_template))
            required = [name for name, p in sig.parameters.items() if p.default is inspect.Parameter.empty]
            assert required == ["template_path"], f"{module.__name__} demands more than a template"

    def test_all_three_call_the_mapping_the_same_thing(self):
        """A caller who learns this tool on one tier must not be wrong on the
        next two. It used to be substitutions / replacements / nothing."""
        for module in (docx_server, xlsx_server, pptx_server):
            sig = inspect.signature(_tool_fn(module.create_from_template))
            assert "substitutions" in sig.parameters, f"{module.__name__} names the mapping something else"
            assert "replacements" not in sig.parameters

    @pytest.mark.parametrize("module,param", TIERS)
    def test_template_path_is_still_required(self, module, param: str):
        """Making the mapping optional must not make everything optional."""
        sig = inspect.signature(_tool_fn(module.create_from_template))
        assert sig.parameters["template_path"].default is inspect.Parameter.empty

    def test_all_three_tiers_require_the_same_thing(self):
        required = {}
        for module in (docx_server, xlsx_server, pptx_server):
            sig = inspect.signature(_tool_fn(module.create_from_template))
            required[module.__name__] = [n for n, p in sig.parameters.items() if p.default is inspect.Parameter.empty]
        assert len({tuple(v) for v in required.values()}) == 1, (
            f"the same tool name still means different things per tier: {required}"
        )


class TestCopyingATemplateUnchangedWorks:
    def test_docx_copies_with_no_substitutions(self, tmp_path: Path):
        src = _blank_docx(tmp_path)
        out = tmp_path / "copy.docx"
        result = _tool_fn(docx_server.create_from_template)(str(src), output_path=str(out))
        assert result["success"] is True, result.get("error")
        assert out.is_file()

    def test_xlsx_copies_with_no_substitutions(self, tmp_path: Path):
        src = _blank_xlsx(tmp_path)
        out = tmp_path / "copy.xlsx"
        result = _tool_fn(xlsx_server.create_from_template)(str(src), output_path=str(out))
        assert result["success"] is True, result.get("error")
        assert out.is_file()

    def test_docx_still_substitutes_when_asked(self, tmp_path: Path):
        from docx import Document

        src = tmp_path / "tpl.docx"
        doc = Document()
        doc.add_paragraph("Hello PLACEHOLDER")
        doc.save(str(src))
        out = tmp_path / "filled.docx"
        result = _tool_fn(docx_server.create_from_template)(str(src), {"PLACEHOLDER": "World"}, output_path=str(out))
        assert result["success"] is True, result.get("error")
        text = "\n".join(p.text for p in Document(str(out)).paragraphs)
        assert "World" in text
        assert "PLACEHOLDER" not in text

    def test_the_default_is_not_shared_between_calls(self, tmp_path: Path):
        """A mutable default is only safe while nothing mutates it."""
        src = _blank_docx(tmp_path)
        _tool_fn(docx_server.create_from_template)(str(src), {"A": "B"}, output_path=str(tmp_path / "one.docx"))
        sig = inspect.signature(_tool_fn(docx_server.create_from_template))
        assert sig.parameters["substitutions"].default == {}


def _blank_docx(tmp_path: Path) -> Path:
    from docx import Document

    p = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_paragraph("unchanged")
    doc.save(str(p))
    return p


def _blank_xlsx(tmp_path: Path) -> Path:
    from openpyxl import Workbook

    p = tmp_path / "tpl.xlsx"
    wb = Workbook()
    wb.active["A1"] = "unchanged"  # type: ignore[index]
    wb.save(str(p))
    return p
