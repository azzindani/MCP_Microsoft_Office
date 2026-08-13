"""Tests for the docx_basic server engine."""

import shutil
from pathlib import Path

import pytest

FIXTURES = Path(__file__).parent / "fixtures"


def _copy(name: str, tmp_path: Path) -> Path:
    src = FIXTURES / name
    if not src.exists():
        pytest.skip(f"Fixture {name} not found. Run create_fixtures.py first.")
    dst = tmp_path / name
    shutil.copy(src, dst)
    return dst


# ─── get_document_outline ─────────────────────────────────────────────────────


class TestGetDocumentOutline:
    def test_returns_headings(self, tmp_path):
        from docx_basic.engine import get_document_outline

        path = _copy("contract_simple.docx", tmp_path)
        result = get_document_outline(str(path))
        assert result["success"] is True
        assert "outline" in result
        assert len(result["outline"]) > 0

    def test_outline_has_required_keys(self, tmp_path):
        from docx_basic.engine import get_document_outline

        path = _copy("contract_simple.docx", tmp_path)
        result = get_document_outline(str(path))
        for entry in result["outline"]:
            assert "index" in entry
            assert "level" in entry
            assert "text" in entry

    def test_file_not_found(self, tmp_path):
        from docx_basic.engine import get_document_outline

        result = get_document_outline(str(tmp_path / "nonexistent.docx"))
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_wrong_file_type(self, tmp_path):
        from docx_basic.engine import get_document_outline

        f = tmp_path / "test.xlsx"
        f.write_bytes(b"fake")
        result = get_document_outline(str(f))
        assert result["success"] is False


# ─── read_document ────────────────────────────────────────────────────────────


class TestReadDocument:
    def test_returns_paragraphs(self, tmp_path):
        from docx_basic.engine import read_document

        path = _copy("contract_simple.docx", tmp_path)
        result = read_document(str(path))
        assert result["success"] is True
        assert result["paragraph_count"] > 0
        assert len(result["paragraphs"]) > 0

    def test_paragraphs_have_required_keys(self, tmp_path):
        from docx_basic.engine import read_document

        path = _copy("contract_simple.docx", tmp_path)
        result = read_document(str(path))
        for p in result["paragraphs"]:
            assert "index" in p
            assert "text" in p
            assert "style" in p

    def test_file_not_found(self, tmp_path):
        from docx_basic.engine import read_document

        result = read_document(str(tmp_path / "missing.docx"))
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_wrong_type(self, tmp_path):
        from docx_basic.engine import read_document

        f = tmp_path / "x.xlsx"
        f.write_bytes(b"fake")
        result = read_document(str(f))
        assert result["success"] is False
        assert ".xlsx" in result["error"]

    def test_has_token_estimate(self, tmp_path):
        from docx_basic.engine import read_document

        path = _copy("contract_simple.docx", tmp_path)
        result = read_document(str(path))
        assert "token_estimate" in result


# ─── read_paragraph ───────────────────────────────────────────────────────────


class TestReadParagraph:
    def test_returns_single_paragraph(self, tmp_path):
        from docx_basic.engine import read_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        result = read_paragraph(str(path), 0)
        assert result["success"] is True
        assert "text" in result
        assert "runs" in result

    def test_returns_run_details(self, tmp_path):
        from docx_basic.engine import read_paragraph

        path = _copy("contract_complex.docx", tmp_path)
        result = read_paragraph(str(path), 0)
        assert result["success"] is True
        # The complex doc has runs with bold
        for run in result["runs"]:
            assert "text" in run
            assert "bold" in run
            assert "italic" in run

    def test_index_out_of_range(self, tmp_path):
        from docx_basic.engine import read_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        result = read_paragraph(str(path), 9999)
        assert result["success"] is False
        assert "out of range" in result["error"]
        assert "hint" in result


# ─── search_paragraphs ────────────────────────────────────────────────────────


class TestSearchParagraphs:
    def test_finds_matching_text(self, tmp_path):
        from docx_basic.engine import search_paragraphs

        path = _copy("contract_simple.docx", tmp_path)
        result = search_paragraphs(str(path), "Payment")
        assert result["success"] is True
        assert len(result["matches"]) > 0

    def test_empty_query_returns_error(self, tmp_path):
        from docx_basic.engine import search_paragraphs

        path = _copy("contract_simple.docx", tmp_path)
        result = search_paragraphs(str(path), "")
        assert result["success"] is False
        assert "empty" in result["error"].lower()

    def test_no_match_returns_empty_list(self, tmp_path):
        from docx_basic.engine import search_paragraphs

        path = _copy("contract_simple.docx", tmp_path)
        result = search_paragraphs(str(path), "XYZZY_NOT_IN_DOC_12345")
        assert result["success"] is True
        assert result["matches"] == []

    def test_respects_max_results(self, tmp_path):
        from docx_basic.engine import search_paragraphs

        path = _copy("contract_simple.docx", tmp_path)
        result = search_paragraphs(str(path), "the", max_results=2)
        assert result["success"] is True
        assert len(result["matches"]) <= 2


# ─── replace_text ─────────────────────────────────────────────────────────────


class TestReplaceText:
    def test_replaces_text(self, tmp_path):
        from docx import Document

        from docx_basic.engine import replace_text

        path = _copy("contract_simple.docx", tmp_path)
        result = replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        assert result["success"] is True
        # Verify change was written
        doc = Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in full_text

    def test_creates_snapshot(self, tmp_path):
        from docx_basic.engine import replace_text

        path = _copy("contract_simple.docx", tmp_path)
        result = replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        assert result["success"] is True
        assert "backup" in result
        assert Path(result["backup"]).exists()

    def test_preserves_bold_formatting(self, tmp_path):
        """Replace text without destroying bold runs in complex doc."""
        from docx import Document

        from docx_basic.engine import replace_text

        path = _copy("contract_complex.docx", tmp_path)

        # The complex doc has PARTY_A_NAME in a bold run
        result = replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        assert result["success"] is True

        doc = Document(str(path))
        # Verify bold formatting is still present somewhere in the doc
        has_bold = any(run.bold for p in doc.paragraphs for run in p.runs if run.text.strip())
        assert has_bold, "Bold formatting was destroyed by replace_text"

    def test_match_not_found(self, tmp_path):
        from docx_basic.engine import replace_text

        path = _copy("contract_simple.docx", tmp_path)
        result = replace_text(str(path), "XYZZY_NOT_HERE", "something")
        assert result["success"] is False
        assert "not found" in result["error"]
        assert "hint" in result

    def test_file_not_found(self, tmp_path):
        from docx_basic.engine import replace_text

        result = replace_text(str(tmp_path / "gone.docx"), "x", "y")
        assert result["success"] is False

    def test_dry_run_makes_no_changes(self, tmp_path):
        from docx import Document

        from docx_basic.engine import replace_text

        path = _copy("contract_simple.docx", tmp_path)
        original_text = Document(str(path)).paragraphs[0].text

        result = replace_text(str(path), "PARTY_A_NAME", "Acme Corp", dry_run=True)
        assert result["success"] is True
        assert result.get("dry_run") is True

        # File should be unchanged
        doc = Document(str(path))
        assert doc.paragraphs[0].text == original_text

    def test_backup_key_in_result(self, tmp_path):
        from docx_basic.engine import replace_text

        path = _copy("contract_simple.docx", tmp_path)
        result = replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        assert "backup" in result


# ─── insert_paragraph ─────────────────────────────────────────────────────────


class TestInsertParagraph:
    def test_inserts_paragraph(self, tmp_path):
        from docx import Document

        from docx_basic.engine import insert_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        original_count = len(Document(str(path)).paragraphs)

        result = insert_paragraph(str(path), 0, "Inserted paragraph")
        assert result["success"] is True

        new_count = len(Document(str(path)).paragraphs)
        assert new_count == original_count + 1

    def test_creates_snapshot(self, tmp_path):
        from docx_basic.engine import insert_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        result = insert_paragraph(str(path), 0, "New text")
        assert result["success"] is True
        assert Path(result["backup"]).exists()

    def test_index_out_of_range(self, tmp_path):
        from docx_basic.engine import insert_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        result = insert_paragraph(str(path), 9999, "text")
        assert result["success"] is False


# ─── delete_paragraph ─────────────────────────────────────────────────────────


class TestDeleteParagraph:
    def test_deletes_by_index(self, tmp_path):
        from docx import Document

        from docx_basic.engine import delete_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        original_count = len(Document(str(path)).paragraphs)

        result = delete_paragraph(str(path), paragraph_index=0)
        assert result["success"] is True

        new_count = len(Document(str(path)).paragraphs)
        assert new_count == original_count - 1

    def test_creates_snapshot(self, tmp_path):
        from docx_basic.engine import delete_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        result = delete_paragraph(str(path), paragraph_index=0)
        assert result["success"] is True
        assert Path(result["backup"]).exists()

    def test_deletes_by_match_text(self, tmp_path):
        from docx_basic.engine import delete_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        result = delete_paragraph(str(path), match_text="Service Agreement")
        # Should succeed or fail gracefully
        assert "success" in result

    def test_match_not_found(self, tmp_path):
        from docx_basic.engine import delete_paragraph

        path = _copy("contract_simple.docx", tmp_path)
        result = delete_paragraph(str(path), match_text="XYZZY_NOT_HERE_99999")
        assert result["success"] is False


# ─── append_text ──────────────────────────────────────────────────────────────


class TestAppendText:
    def test_appends_paragraph(self, tmp_path):
        from docx import Document

        from docx_basic.engine import append_text

        path = _copy("contract_simple.docx", tmp_path)
        original_count = len(Document(str(path)).paragraphs)

        result = append_text(str(path), "Appended paragraph text.")
        assert result["success"] is True

        new_count = len(Document(str(path)).paragraphs)
        assert new_count == original_count + 1

    def test_creates_snapshot(self, tmp_path):
        from docx_basic.engine import append_text

        path = _copy("contract_simple.docx", tmp_path)
        result = append_text(str(path), "Some text")
        assert result["success"] is True
        assert Path(result["backup"]).exists()


# ─── get_history / restore_version ───────────────────────────────────────────


class TestVersionControl:
    def test_get_history_returns_entries(self, tmp_path):
        from docx_basic.engine import get_history_tool, replace_text

        path = _copy("contract_simple.docx", tmp_path)
        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")

        result = get_history_tool(str(path))
        assert result["success"] is True
        assert len(result["history"]) >= 1
        assert result["history"][0]["timestamp"] is not None

    def test_restore_version_reverts_content(self, tmp_path):
        from docx import Document

        from docx_basic.engine import get_history_tool, replace_text, restore_version

        path = _copy("contract_simple.docx", tmp_path)
        original_text = Document(str(path)).paragraphs[0].text

        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        history = get_history_tool(str(path))
        timestamp = history["history"][0]["timestamp"]

        result = restore_version(str(path), timestamp)
        assert result["success"] is True

        restored_text = Document(str(path)).paragraphs[0].text
        assert restored_text == original_text

    def test_restore_nonexistent_snapshot(self, tmp_path):
        from docx_basic.engine import restore_version

        path = _copy("contract_simple.docx", tmp_path)
        result = restore_version(str(path), "2099-01-01T00-00-00Z")
        assert result["success"] is False
        assert "hint" in result


# ─── diff_versions ────────────────────────────────────────────────────────────


class TestDiffVersions:
    def test_diff_detects_change(self, tmp_path):
        from docx_basic.engine import diff_versions, get_history_tool, replace_text

        path = _copy("contract_simple.docx", tmp_path)

        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        history = get_history_tool(str(path))
        timestamp = history["history"][0]["timestamp"]

        result = diff_versions(str(path), timestamp, "current")
        assert result["success"] is True
        assert "change_count" in result
        assert "summary" in result

    def test_diff_returns_summary(self, tmp_path):
        from docx_basic.engine import diff_versions, get_history_tool, replace_text

        path = _copy("contract_simple.docx", tmp_path)
        replace_text(str(path), "PARTY_B_NAME", "Widget Ltd")
        history = get_history_tool(str(path))
        timestamp = history["history"][0]["timestamp"]

        result = diff_versions(str(path), timestamp)
        assert result["success"] is True
        assert isinstance(result.get("summary"), str)


# ─── read_paragraph_range ─────────────────────────────────────────────────────


class TestReadParagraphRange:
    def test_returns_range(self, tmp_path):
        from docx_basic.engine import read_paragraph_range

        path = _copy("contract_simple.docx", tmp_path)
        result = read_paragraph_range(str(path), 0, 3)
        assert result["success"] is True
        assert len(result["paragraphs"]) <= 4

    def test_range_too_large_returns_error(self, tmp_path):
        from docx_basic.engine import read_paragraph_range

        path = _copy("contract_simple.docx", tmp_path)
        result = read_paragraph_range(str(path), 0, 200)
        assert result["success"] is False
        assert "hint" in result


# ─── progress field ───────────────────────────────────────────────────────────


class TestProgressField:
    """Every tool response must include a 'progress' array."""

    def test_read_document_has_progress(self, tmp_path):
        from docx_basic.engine import read_document

        path = _copy("contract_simple.docx", tmp_path)
        result = read_document(str(path))
        assert "progress" in result
        assert isinstance(result["progress"], list)

    def test_replace_text_has_progress(self, tmp_path):
        from docx_basic.engine import replace_text

        path = _copy("contract_simple.docx", tmp_path)
        result = replace_text(str(path), "PARTY_A_NAME", "Acme")
        assert "progress" in result

    def test_error_response_has_progress(self, tmp_path):
        from docx_basic.engine import read_document

        result = read_document("/nonexistent/path/file.docx")
        assert "progress" in result


# ─── get_document_index ────────────────────────────────────────────────────


class TestGetDocumentIndex:
    def test_returns_section_tree(self, tmp_path):
        from docx_basic.engine import get_document_index

        path = _copy("contract_simple.docx", tmp_path)
        result = get_document_index(str(path))
        assert result["success"] is True
        assert result["total_paragraphs"] == 13
        assert len(result["sections"]) > 0
        assert result["sections"][0]["address"] == "§1"

    def test_file_not_found(self, tmp_path):
        from docx_basic.engine import get_document_index

        result = get_document_index(str(tmp_path / "nonexistent.docx"))
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_wrong_file_type(self, tmp_path):
        from docx_basic.engine import get_document_index

        bad = tmp_path / "notes.txt"
        bad.write_text("hello")
        result = get_document_index(str(bad))
        assert result["success"] is False


# ─── fetch_section ──────────────────────────────────────────────────────────


class TestFetchSection:
    def test_returns_paragraphs_for_valid_address(self, tmp_path):
        from docx_basic.engine import fetch_section

        path = _copy("contract_simple.docx", tmp_path)
        result = fetch_section(str(path), "§1")
        assert result["success"] is True
        assert result["heading"] == "Service Agreement"
        assert len(result["paragraphs"]) == 13
        assert result["paragraphs"][0]["addr"] == "§1.p0"

    def test_invalid_address_gives_actionable_error(self, tmp_path):
        from docx_basic.engine import fetch_section

        path = _copy("contract_simple.docx", tmp_path)
        result = fetch_section(str(path), "§99")
        assert result["success"] is False
        assert "get_document_index" in result["hint"]

    def test_file_not_found(self, tmp_path):
        from docx_basic.engine import fetch_section

        result = fetch_section(str(tmp_path / "nonexistent.docx"), "§1")
        assert result["success"] is False


# ─── read_receipt ───────────────────────────────────────────────────────────


class TestReadReceipt:
    def test_empty_log_before_any_operation(self, tmp_path):
        from docx_basic.helpers import read_receipt_tool

        path = _copy("contract_simple.docx", tmp_path)
        result = read_receipt_tool(str(path))
        assert result["success"] is True
        assert result["entries"] == []

    def test_records_entry_after_mutating_op(self, tmp_path):
        from docx_basic.engine import replace_text
        from docx_basic.helpers import read_receipt_tool

        path = _copy("contract_simple.docx", tmp_path)
        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")

        result = read_receipt_tool(str(path))
        assert result["success"] is True
        assert len(result["entries"]) == 1
        assert result["entries"][0]["tool"] == "replace_text"
        assert result["entries"][0]["success"] is True

    def test_last_n_limits_entries(self, tmp_path):
        from docx_basic.engine import replace_text
        from docx_basic.helpers import read_receipt_tool

        path = _copy("contract_simple.docx", tmp_path)
        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        replace_text(str(path), "PARTY_B_NAME", "Widget Ltd")
        replace_text(str(path), "30 days", "45 days")

        result = read_receipt_tool(str(path), last_n=2)
        assert len(result["entries"]) == 2
