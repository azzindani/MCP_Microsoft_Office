"""add_table said "in the empty document" about a document with a table in it.

    add_table(doc.docx, after_paragraph_index=-1, rows=1, cols=1)
      -> "Inserted 1x1 table in the empty document"

The docstring is "Insert table after paragraph N (-1 = first)". The document
held one table and no paragraphs, so -1 had a front to be placed at -- and the
new table went to the back, behind the one already there, described as an empty
document.

The cause is one assumption written into an earlier fix: "With no paragraphs
there is exactly one place a table can go, so there is nothing to
disambiguate." That is true of a brand-new file and false of any document whose
content is tables. total_paras == 0 means "no paragraphs", never "no content",
and the code read it as the second.

A round-16 phase reported this as "-1 insertion-point semantics not honored" on
a "paragraph-only document". That description is wrong -- on a document that
really does have paragraphs, -1 has always worked and still does, and there is
a test below holding that. Reproducing it is what found the case that is
genuinely broken: no paragraphs, but content all the same.

The placement is what matters, so the table-only test reads the first table's
cell text back rather than trusting the message. Before the fix that cell says
EXISTING; after it says NEW.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from docx_tables import engine as te  # type: ignore[reportMissingImports]


def _body_tags(path: Path) -> list[str]:
    return [c.tag.split("}")[-1] for c in docx.Document(str(path)).element.body]


def _placement_message(result: dict) -> str:
    return " | ".join(
        str(p.get("message", "")) for p in result.get("progress", []) if "Inserted" in str(p.get("message", ""))
    )


@pytest.fixture()
def table_only(tmp_path: Path) -> Path:
    """A document whose only content is a table -- no paragraphs at all."""
    p = tmp_path / "tableonly.docx"
    d = docx.Document()
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "EXISTING"
    for para in list(d.paragraphs):
        para._element.getparent().remove(para._element)
    d.save(str(p))
    assert _body_tags(p).count("p") == 0, "fixture must have no paragraphs"
    assert _body_tags(p).count("tbl") == 1
    return p


class TestATableOnlyDocumentIsNotEmpty:
    def test_minus_one_puts_the_new_table_first(self, table_only: Path) -> None:
        """Placement, read back from the file -- not the message."""
        r = te.add_table(str(table_only), after_paragraph_index=-1, rows=1, cols=1, data=[["NEW"]], open_after=False)
        assert r["success"] is True, r.get("error")
        d = docx.Document(str(table_only))
        assert len(d.tables) == 2
        assert d.tables[0].rows[0].cells[0].text == "NEW", "the new table did not go first"
        assert d.tables[1].rows[0].cells[0].text == "EXISTING"

    def test_it_is_not_described_as_empty(self, table_only: Path) -> None:
        r = te.add_table(str(table_only), after_paragraph_index=-1, rows=1, cols=1, data=[["NEW"]], open_after=False)
        msg = _placement_message(r)
        assert "empty document" not in msg, msg

    def test_the_message_says_where_it_went(self, table_only: Path) -> None:
        r = te.add_table(str(table_only), after_paragraph_index=-1, rows=1, cols=1, data=[["NEW"]], open_after=False)
        assert "before the existing content" in _placement_message(r)


class TestATrulyEmptyDocumentStillReadsAsEmpty:
    """ "First" and "last" are the same position there, and saying so is right."""

    def test_message_and_placement(self, tmp_path: Path) -> None:
        p = tmp_path / "empty.docx"
        docx.Document().save(str(p))
        r = te.add_table(str(p), after_paragraph_index=-1, rows=1, cols=1, data=[["x"]], open_after=False)
        assert r["success"] is True, r.get("error")
        assert "in the empty document" in _placement_message(r)
        assert _body_tags(p).count("tbl") == 1


class TestDocumentsWithParagraphsAreUnchanged:
    """The reported symptom, which was never actually broken."""

    @pytest.fixture()
    def two_paragraphs(self, tmp_path: Path) -> Path:
        p = tmp_path / "paras.docx"
        d = docx.Document()
        d.add_paragraph("one")
        d.add_paragraph("two")
        d.save(str(p))
        return p

    def test_minus_one_goes_before_the_first_paragraph(self, two_paragraphs: Path) -> None:
        r = te.add_table(str(two_paragraphs), after_paragraph_index=-1, rows=1, cols=1, data=[["y"]], open_after=False)
        assert r["success"] is True, r.get("error")
        assert _body_tags(two_paragraphs)[:3] == ["tbl", "p", "p"]
        assert "before paragraph 0" in _placement_message(r)

    def test_zero_goes_after_the_first_paragraph(self, two_paragraphs: Path) -> None:
        r = te.add_table(str(two_paragraphs), after_paragraph_index=0, rows=1, cols=1, data=[["z"]], open_after=False)
        assert r["success"] is True, r.get("error")
        assert _body_tags(two_paragraphs)[:3] == ["p", "tbl", "p"]
        assert "after paragraph 0" in _placement_message(r)

    def test_an_out_of_range_index_is_still_refused(self, two_paragraphs: Path) -> None:
        r = te.add_table(str(two_paragraphs), after_paragraph_index=9, rows=1, cols=1, data=[["q"]], open_after=False)
        assert r["success"] is False
        assert "out of range" in r["error"]
