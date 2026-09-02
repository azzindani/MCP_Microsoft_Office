"""create_from_sections can only ever write a heading and a paragraph.

Its `sections` are `{heading, body}` strings, so each one becomes exactly one
Heading 2 and one Normal paragraph. Handed a data summary it produced ten
headings over ten 180-word paragraphs -- 21 paragraphs, 0 tables, 0 bullets --
and the reply that came back was "so many text and numbers in single
paragraphs, not readable for my director".

Nothing was broken. The tool did what it says. It simply has no way to express
a table, a bulleted list, a figure worth pulling out, or the one sentence that
should sit in a box, and the per-paragraph tools that could build those are
addressed by index and shift under every insert -- roughly 150 sequential calls
for a four-page brief. So the work left the tool surface: a model asked to fix
the document installed python-docx and wrote it by hand.

`create_from_blocks` is that document in one call. These tests check the parts
that are invisible in a success response: that the shading is really in the
XML, that a kpi row and a table do not merge into one grid, and that an
unrecognised block is reported rather than silently skipped.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

import pytest

from docx_new import engine as docx_new  # type: ignore[reportMissingImports]

ACCENT = "0B1D3A"

BLOCKS = [
    {"kind": "callout", "title": "Bottom line", "text": "Not a growth story."},
    {"kind": "heading", "text": "At a glance", "level": 2},
    {"kind": "kpi", "items": [{"value": "12.69 M", "label": "Total tons"}, {"value": "43.6k", "label": "Avg"}]},
    {"kind": "table", "header": ["Region", "Share"], "rows": [["US", "42.7%"], ["Asia", "40.8%"], ["Europe", "13%"]]},
    {"kind": "rule"},
    {"kind": "bullets", "items": ["Re-price for reality.", "De-risk carriers."]},
    {"kind": "text", "text": "Snapshot frozen 20 Nov 2023."},
]


@pytest.fixture
def brief(tmp_path):
    out = tmp_path / "brief.docx"
    result = docx_new.create_from_blocks(str(out), "Executive Brief", BLOCKS, accent=ACCENT, open_after=False)
    assert result["success"] is True, result
    return out, result


def document_xml(path: Path) -> str:
    return zipfile.ZipFile(path).read("word/document.xml").decode()


def test_every_block_kind_is_written(brief):
    _, result = brief
    assert result["blocks_by_kind"] == {
        "callout": 1,
        "heading": 1,
        "kpi": 1,
        "table": 1,
        "rule": 1,
        "bullets": 1,
        "text": 1,
    }
    assert result["skipped"] == []


def test_the_document_actually_holds_tables_and_bullets(brief):
    """The complaint was structural. Structure is what is asserted."""
    from docx import Document

    path, _ = brief
    doc = Document(str(path))
    assert len(doc.tables) == 3, "callout, kpi row and the table itself"
    styles = [p.style.name for p in doc.paragraphs]
    assert any(s.startswith("List") for s in styles), "bullets must be a list style"
    assert any(s.startswith("Heading") for s in styles)


def test_the_shading_is_really_in_the_file(brief):
    """A fill that is not in the XML is a success response over a plain table."""
    path, _ = brief
    fills = re.findall(r'w:fill="([0-9A-F]{6})"', document_xml(path))
    assert ACCENT in fills, "the table header must carry the accent"
    # Header + banded body rows + the callout: more than one shaded cell.
    assert len(fills) > 1
    assert re.search(r"<w:pBdr>", document_xml(path)), "the title rule must be present"


def test_a_kpi_row_and_a_table_do_not_merge_into_one_grid(brief):
    """Word joins two tables with nothing between them.

    This is the layout defect that is hardest to see -- the response is
    identical either way, and it only shows when somebody opens the file.
    """
    path, _ = brief
    assert not re.search(r"</w:tbl>\s*<w:tbl>", document_xml(path))


def test_an_unrecognised_block_is_reported_not_dropped(tmp_path):
    out = tmp_path / "x.docx"
    result = docx_new.create_from_blocks(
        str(out),
        "T",
        [{"kind": "text", "text": "kept"}, {"kind": "bar_chart", "text": "no such kind"}, {"nope": 1}],
        open_after=False,
    )
    assert result["success"] is True
    assert result["block_count"] == 1
    assert len(result["skipped"]) == 2
    assert any("bar_chart" in s for s in result["skipped"])
    assert any(p["status"] == "warn" for p in result["progress"])


def test_a_table_block_with_no_rows_is_not_silently_empty(tmp_path):
    result = docx_new.create_from_blocks(str(tmp_path / "x.docx"), "T", [{"kind": "table"}], open_after=False)
    assert result["success"] is True
    assert result["skipped"] and "table" in result["skipped"][0]


def test_the_kind_key_may_be_spelled_type(tmp_path):
    """Alias spellings elsewhere in this repo exist because callers use them."""
    result = docx_new.create_from_blocks(
        str(tmp_path / "x.docx"), "T", [{"type": "text", "content": "hello"}], open_after=False
    )
    assert result["blocks_by_kind"] == {"text": 1}


def test_a_bad_accent_is_refused_before_anything_is_written(tmp_path):
    out = tmp_path / "x.docx"
    result = docx_new.create_from_blocks(
        str(out), "T", [{"kind": "text", "text": "a"}], accent="navy", open_after=False
    )
    assert result["success"] is False
    assert "navy" in result["error"]
    assert not out.exists()


def test_blocks_must_be_a_list(tmp_path):
    result = docx_new.create_from_blocks(str(tmp_path / "x.docx"), "T", "not a list", open_after=False)
    assert result["success"] is False
    assert "list" in result["error"]


def test_the_default_accent_still_produces_a_shaded_header(tmp_path):
    out = tmp_path / "x.docx"
    docx_new.create_from_blocks(str(out), "T", [{"kind": "table", "header": ["A"], "rows": [["1"]]}], open_after=False)
    assert re.search(r'w:fill="[0-9A-F]{6}"', document_xml(out))
