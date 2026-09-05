"""A board paper referencing charts it could not contain.

    3.7 `create_from_sections` (docx) cannot embed images/charts
    The board paper references charts that live in separate HTML files.
    Acceptable, but image-in-block support would close the loop.
    Room for improvement: an `image` block type accepting a `data/` path or URL.

The other half of that item was already here: `create_from_blocks` has `kpi`,
`table` and `callout`, so the KPI row and the findings table the review also
asked for exist. The gap was the picture.

**Three refusals matter more than the feature.** The file a caller has in hand
is almost always a `.html` chart, because that is what every generator in this
fleet writes and it is self-contained by design. python-docx cannot place one.
A generic "unsupported format" sends an agent to guess; naming the format and
the way out does not. The same is true of a path that does not exist -- usually
a path from the wrong side of the container boundary -- and of a URL that
returned a status, because "the document has no image" and "the file server
refused" are different problems with different fixes.

A refused block is reported in `skipped` and the rest of the document is still
written. One unreachable chart should not cost the paper.
"""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest

from servers.docx_new.docx_new.engine import BLOCK_KINDS, create_from_blocks


def _png(tmp_path: Path, name: str = "chart.png") -> Path:
    """A real 1x1 PNG, built here so the test needs no fixture file."""

    def chunk(kind: bytes, data: bytes) -> bytes:
        body = kind + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body))

    raw = zlib.compress(b"\x00\xff\x00\x00")
    png = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", raw)
        + chunk(b"IEND", b"")
    )
    p = tmp_path / name
    p.write_bytes(png)
    return p


def test_image_is_a_block_kind():
    assert "image" in BLOCK_KINDS


def test_a_png_is_embedded(tmp_path):
    img = _png(tmp_path)
    out = tmp_path / "paper.docx"
    r = create_from_blocks(
        str(out),
        "Board paper",
        [{"kind": "heading", "text": "Findings"}, {"kind": "image", "path": str(img)}],
        open_after=False,
    )
    assert r["success"] is True, r
    assert r["images_embedded"] == 1
    assert r["skipped"] == []
    assert out.exists()


def test_the_image_is_really_in_the_file(tmp_path):
    """A .docx is a zip; an embedded picture is a part inside it."""
    import zipfile

    img = _png(tmp_path)
    out = tmp_path / "paper.docx"
    create_from_blocks(str(out), "T", [{"kind": "image", "path": str(img)}], open_after=False)
    with zipfile.ZipFile(out) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    assert media, "no media part -- the block reported success and embedded nothing"


def test_a_caption_is_written_beneath_it(tmp_path):
    from docx import Document

    img = _png(tmp_path)
    out = tmp_path / "paper.docx"
    create_from_blocks(
        str(out),
        "T",
        [{"kind": "image", "path": str(img), "caption": "Figure 1 — charged-off by grade"}],
        open_after=False,
    )
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("Figure 1" in t for t in texts)


# ---------------------------------------------------------------------------
# the refusals, which are the point
# ---------------------------------------------------------------------------


def test_an_html_chart_is_refused_by_name(tmp_path):
    """The file a caller actually has, and the one they will try first."""
    chart = tmp_path / "Credit_Risk_bar.html"
    chart.write_text("<html>a self-contained plotly page</html>", encoding="utf-8")
    out = tmp_path / "paper.docx"
    r = create_from_blocks(str(out), "T", [{"kind": "image", "path": str(chart)}], open_after=False)
    assert r["success"] is True, "one bad block must not cost the document"
    assert r["images_embedded"] == 0
    assert len(r["skipped"]) == 1
    note = r["skipped"][0]
    assert "HTML page" in note
    assert "render one to PNG" in note, "a refusal without a way out sends an agent guessing"


def test_a_missing_file_says_where_paths_resolve(tmp_path):
    out = tmp_path / "paper.docx"
    r = create_from_blocks(str(out), "T", [{"kind": "image", "path": "/nope/chart.png"}], open_after=False)
    assert r["images_embedded"] == 0
    assert "does not exist" in r["skipped"][0]
    assert "resolved on the server" in r["skipped"][0]


def test_a_block_with_no_path_says_so(tmp_path):
    out = tmp_path / "paper.docx"
    r = create_from_blocks(str(out), "T", [{"kind": "image"}], open_after=False)
    assert "no path or url" in r["skipped"][0]


def test_a_non_image_file_names_what_can_be_placed(tmp_path):
    csv = tmp_path / "data.csv"
    csv.write_text("a,b\n1,2\n", encoding="utf-8")
    out = tmp_path / "paper.docx"
    r = create_from_blocks(str(out), "T", [{"kind": "image", "path": str(csv)}], open_after=False)
    assert ".png" in r["skipped"][0]


def test_a_file_that_lies_about_being_a_png_is_caught(tmp_path):
    """The extension is a claim; python-docx is the one that has to read it."""
    fake = tmp_path / "not_really.png"
    fake.write_text("this is not a png", encoding="utf-8")
    out = tmp_path / "paper.docx"
    r = create_from_blocks(str(out), "T", [{"kind": "image", "path": str(fake)}], open_after=False)
    assert r["images_embedded"] == 0
    assert "not a readable image" in r["skipped"][0]


def test_one_bad_image_does_not_cost_the_rest_of_the_paper(tmp_path):
    from docx import Document

    img = _png(tmp_path)
    out = tmp_path / "paper.docx"
    r = create_from_blocks(
        str(out),
        "Board paper",
        [
            {"kind": "heading", "text": "Findings"},
            {"kind": "image", "path": "/nope.png"},
            {"kind": "image", "path": str(img)},
            {"kind": "text", "text": "Charged-off rate was 13.82%."},
        ],
        open_after=False,
    )
    assert r["success"] is True
    assert r["images_embedded"] == 1
    assert len(r["skipped"]) == 1
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("13.82%" in t for t in texts)


def test_the_width_is_bounded(tmp_path):
    """A caller asking for 40 inches gets a page-width image, not a broken file."""
    img = _png(tmp_path)
    out = tmp_path / "paper.docx"
    r = create_from_blocks(str(out), "T", [{"kind": "image", "path": str(img), "width_in": 40}], open_after=False)
    assert r["images_embedded"] == 1


@pytest.mark.parametrize("key", ["path", "src", "url", "file_path", "image"])
def test_the_path_is_read_from_any_of_the_names_a_caller_uses(tmp_path, key):
    img = _png(tmp_path, f"c_{key}.png")
    out = tmp_path / f"p_{key}.docx"
    r = create_from_blocks(str(out), "T", [{"kind": "image", key: str(img)}], open_after=False)
    assert r["images_embedded"] == 1, r["skipped"]
