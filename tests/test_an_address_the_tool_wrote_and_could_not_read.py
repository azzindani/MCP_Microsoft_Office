"""It handed back addresses it could not accept.

    fetch_section(doc, "§1")
      -> paragraphs: [{"addr": "§1.p0", ...}, {"addr": "§1.p1", ...}, ...]

    fetch_section(doc, "§1.p1")
      -> Invalid section address: '§1.p1'. Use §N for a section, e.g. '§1',
         or §N.pM for a paragraph inside one, e.g. '§1.p3'.

Three statements, no two of which agree. The tool labels every paragraph it
returns with a `§N.pM` address; its own error message documents `§N.pM` and
offers '§1.p3' as the example; and the branch that reads the address matches
`^§(\\d+)$`, which admits neither. The documented example fails exactly as the
emitted address does.

An address a tool prints is an identifier a caller will send back -- that is
what makes it an address rather than a label. This one was a dead end in both
directions.

The cause is the one round 14 kept finding: a second copy. `resolve_docx_address`,
forty lines above in the same file, has parsed `§N.pM` correctly all along --
same section lookup, same `para_start + rel_idx` arithmetic, and a message that
names every accepted form. `fetch_section_content` did not call it; it grew its
own narrower matcher. So one branch of one file understood the notation and the
other did not, and the tests only ever exercised the branch that did.

Fixed by giving both the same compiled pattern, so a form the sibling accepts
cannot be a form this one rejects.

Found in round 15, phase 24: "Tool emits §N.pM labels it cannot parse back.
Only bare §N works."
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document

from docx_basic.engine import fetch_section  # type: ignore[reportMissingImports]

RESOLVER = Path(__file__).parent.parent / "shared" / "shared" / "address_resolver.py"


@pytest.fixture()
def doc(tmp_path: Path) -> str:
    path = tmp_path / "doc.docx"
    d = Document()
    d.add_heading("Ad campaign spend review", level=1)
    for line in ("Total spend was 2,503,118.77.", "Google Ads carried 1,939,003.26.", "Facebook Ads 564,115.51."):
        d.add_paragraph(line)
    d.add_heading("Channel notes", level=1)
    d.add_paragraph("Facebook link clicks are missing for 546 rows.")
    d.save(str(path))
    return str(path)


def addrs_of(response: dict) -> list[str]:
    return [p["addr"] for p in response["paragraphs"]]


class TestEveryAddressItPrintsIsOneItAccepts:
    def test_the_section_read_succeeds(self, doc: str) -> None:
        r = fetch_section(doc, "§1")
        assert r["success"] is True, r.get("error")
        assert len(addrs_of(r)) >= 3, addrs_of(r)

    def test_each_emitted_address_can_be_fetched(self, doc: str) -> None:
        """The whole defect, stated once: send back what it just gave you."""
        emitted = addrs_of(fetch_section(doc, "§1"))
        refused = [a for a in emitted if not fetch_section(doc, a)["success"]]
        assert not refused, f"emitted but not accepted: {refused}"

    def test_an_emitted_address_returns_that_paragraph(self, doc: str) -> None:
        section = fetch_section(doc, "§1")
        wanted = section["paragraphs"][2]
        one = fetch_section(doc, wanted["addr"])
        assert one["success"] is True, one.get("error")
        assert [p["text"] for p in one["paragraphs"]] == [wanted["text"]]

    def test_it_keeps_the_section_heading(self, doc: str) -> None:
        section = fetch_section(doc, "§1")
        one = fetch_section(doc, section["paragraphs"][1]["addr"])
        assert one["heading"] == section["heading"], (one["heading"], section["heading"])

    def test_the_second_section_addresses_work_too(self, doc: str) -> None:
        """Relative indexing must be relative to the right section."""
        second = fetch_section(doc, "§2")
        assert second["success"] is True, second.get("error")
        for a in addrs_of(second):
            assert a.startswith("§2."), a
            assert fetch_section(doc, a)["success"] is True, a


class TestTheDocumentedExampleWorks:
    def test_the_form_the_error_message_offers(self, doc: str) -> None:
        """'§1.p3' is the example the refusal prints. It used to fail too."""
        bad = fetch_section(doc, "1")
        assert bad["success"] is False
        offered = re.findall(r"§\d+\.p\d+", f"{bad.get('error', '')} {bad.get('hint', '')}")
        assert offered, "the refusal no longer shows a §N.pM example"

        r = fetch_section(doc, "§1.p3")
        assert r["success"] is True, r.get("error")


class TestItStillRefusesWhatIsActuallyWrong:
    def test_a_paragraph_past_the_end_of_the_section(self, doc: str) -> None:
        r = fetch_section(doc, "§1.p99")
        assert r["success"] is False
        assert "§1" in f"{r.get('error', '')}{r.get('hint', '')}"

    def test_a_section_that_does_not_exist(self, doc: str) -> None:
        r = fetch_section(doc, "§99")
        assert r["success"] is False

    @pytest.mark.parametrize("guess", ["1", "p1x", "section:1", "§1.q2", "", "§"])
    def test_a_malformed_address_is_still_refused(self, doc: str, guess: str) -> None:
        assert fetch_section(doc, guess)["success"] is False, guess


class TestTheIndexSaysWhereItCutSections:
    """Phase 24's other note: one section where the outline showed five headings.

    Both numbers are right -- a section is cut at Heading 1 and the outline
    lists every heading -- but nothing said so, and read together they look
    like one of the two tools is broken. The scheme is not changed; changing
    it would move every existing §N address.
    """

    @pytest.fixture()
    def nested(self, tmp_path: Path) -> str:
        path = tmp_path / "nested.docx"
        d = Document()
        d.add_heading("Ad campaign spend review", level=1)
        for h in ("Spend", "Impressions", "Clicks", "Notes"):
            d.add_heading(h, level=2)
            d.add_paragraph(f"Body for {h}.")
        d.save(str(path))
        return str(path)

    def test_one_section_for_one_top_level_heading(self, nested: str) -> None:
        from docx_basic.engine import get_document_index  # type: ignore[reportMissingImports]

        r = get_document_index(nested)
        assert r["success"] is True, r.get("error")
        assert len(r["sections"]) == 1, r["sections"]

    def test_it_reports_the_level_it_cut_on(self, nested: str) -> None:
        from docx_basic.engine import get_document_index  # type: ignore[reportMissingImports]

        assert get_document_index(nested)["sections_cut_at_heading_level"] == 1

    def test_it_counts_the_headings_it_passed_over(self, nested: str) -> None:
        from docx_basic.engine import get_document_index  # type: ignore[reportMissingImports]

        assert get_document_index(nested)["heading_counts"] == {"1": 1, "2": 4}

    def test_the_hint_accounts_for_the_difference(self, nested: str) -> None:
        from docx_basic.engine import get_document_index  # type: ignore[reportMissingImports]

        hint = get_document_index(nested)["hint"]
        assert "Heading 1" in hint and "4 deeper heading" in hint, hint
        assert "get_document_outline" in hint, hint

    def test_a_flat_hint_when_every_heading_is_top_level(self, tmp_path: Path) -> None:
        """The note must fire on a difference, not on every document."""
        from docx_basic.engine import get_document_index  # type: ignore[reportMissingImports]

        path = tmp_path / "flat.docx"
        d = Document()
        for h in ("One", "Two"):
            d.add_heading(h, level=1)
            d.add_paragraph("Body.")
        d.save(str(path))
        r = get_document_index(str(path))
        assert len(r["sections"]) == 2
        assert "deeper heading" not in r["hint"], r["hint"]

    def test_the_deeper_headings_are_reachable(self, nested: str) -> None:
        """The disclosure is only fair if the addresses it points at work."""
        from docx_basic.engine import get_document_index  # type: ignore[reportMissingImports]

        r = get_document_index(nested)
        addrs = [p["addr"] for p in fetch_section(nested, r["sections"][0]["address"])["paragraphs"]]
        assert len(addrs) == 9, addrs
        assert all(fetch_section(nested, a)["success"] for a in addrs)


class TestThereIsOnlyOnePatternForThisNotation:
    """A second matcher is what let the two branches disagree."""

    def test_no_branch_compiles_its_own(self) -> None:
        """The shape of the defect: two matchers for one notation.

        The canonical definition is excluded by name -- everything else that
        builds a regex mentioning the section sign is a second one.
        """
        src = RESOLVER.read_text(encoding="utf-8")
        elsewhere = "\n".join(
            ln for ln in src.splitlines() if not ln.lstrip().startswith("_SECTION_ADDRESS = re.compile")
        )
        inline = re.findall(r"re\.(?:match|compile|fullmatch|search)\(\s*r?[\"'][^\"']*§", elsewhere)
        assert not inline, f"a section-address regex is still written inline: {inline}"

    def test_the_pattern_is_defined_once_and_shared(self) -> None:
        src = RESOLVER.read_text(encoding="utf-8")
        assert src.count("_SECTION_ADDRESS = re.compile") == 1, "defined more than once"
        assert src.count("_SECTION_ADDRESS.match") >= 2, "both branches must read it from the same place"
