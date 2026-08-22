"""A freshly created Word document could not be given a table, an image, or a
first paragraph.

create_document() writes a .docx with zero paragraphs. Three tools then anchor
their work to a paragraph index, and every index into an empty collection is out
of range, so all three refused:

    add_table(after_paragraph_index=0)   Paragraph index 0 out of range
                                         — this file has no paragraphs
    add_image(paragraph_index=0)         same
    insert_paragraph(after_index=0)      paragraph_index 0 out of range
                                         — this file has no paragraphs

For add_table and add_image there was no way round it: they are the only tools
that create a table or place an image, so the obvious two-step -- create the
document, then put something in it -- simply did not work. The hints made it
worse by pointing at read_document "to see paragraph indices", which on an empty
document returns an empty list; the caller is sent to look at nothing, the same
dead end the docx_tables hints were fixed for earlier.

insert_paragraph had a working alternative in append_text, but failed harder.
after_index=-1 is its documented "insert at the beginning" value and the one a
caller reaches for on an empty file. -1 passed the guard, reached
doc.paragraphs[0], and returned the raw IndexError to the caller:

    error: "list index out of range"
    hint:  "Use restore_version to undo if a snapshot was taken."

-- a message naming nothing, and a hint offering to undo a change that had not
happened.

None of this is ambiguous to resolve. With no paragraphs there is exactly one
place the content can go, so any index means that place. The guards now apply
only when there is something to index, and the empty document is filled.
"""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from docx_basic.engine import append_text, insert_paragraph  # type: ignore[reportMissingImports]
from docx_layout.engine import add_image  # type: ignore[reportMissingImports]
from docx_tables.engine import add_table, read_table  # type: ignore[reportMissingImports]


def _png(path: Path) -> str:
    """Smallest valid 1x1 PNG, built here so the test needs no fixture file."""

    def chunk(tag: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload)) + tag + payload + struct.pack(">I", zlib.crc32(tag + payload) & 0xFFFFFFFF)
        )

    raw = zlib.compress(b"\x00\xff\xff\xff")
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", raw)
        + chunk(b"IEND", b"")
    )
    return str(path)


@pytest.fixture()
def empty_doc(tmp_path: Path) -> str:
    """What create_document() produces: a .docx with no paragraphs at all."""
    path = tmp_path / "new.docx"
    Document().save(str(path))
    assert len(Document(str(path)).paragraphs) == 0, "fixture must start empty"
    return str(path)


class TestTheDocumentReallyStartsEmpty:
    def test_zero_paragraphs(self, empty_doc: str):
        assert len(Document(empty_doc).paragraphs) == 0


class TestATableCanBeAdded:
    def test_add_table_succeeds_on_an_empty_document(self, empty_doc: str):
        r = add_table(empty_doc, 0, 2, 2)
        assert r["success"] is True, r.get("error")

    def test_the_table_is_really_in_the_file(self, empty_doc: str):
        add_table(empty_doc, 0, 3, 3)
        assert len(Document(empty_doc).tables) == 1

    def test_the_data_lands_in_the_cells(self, empty_doc: str):
        add_table(empty_doc, 0, 2, 2, data=[["Platform", "Spend"], ["Google Ads", "1939000"]])
        table = Document(empty_doc).tables[0]
        assert table.cell(0, 0).text == "Platform"
        assert table.cell(1, 1).text == "1939000"

    def test_it_can_then_be_read_back(self, empty_doc: str):
        """The five reader tools were unreachable while add_table was."""
        add_table(empty_doc, 0, 2, 2, data=[["a", "b"], ["c", "d"]])
        r = read_table(empty_doc, 0)
        assert r["success"] is True, r.get("error")

    def test_the_progress_says_where_it_went(self, empty_doc: str):
        """ "after paragraph 0" would be a lie -- there is no paragraph 0."""
        r = add_table(empty_doc, 0, 2, 2)
        blob = str(r["progress"])
        assert "empty document" in blob, blob


class TestAnImageCanBeAdded:
    def test_add_image_succeeds_on_an_empty_document(self, empty_doc: str, tmp_path: Path):
        r = add_image(empty_doc, 0, _png(tmp_path / "dot.png"))
        assert r["success"] is True, r.get("error")

    def test_the_image_is_really_embedded(self, empty_doc: str, tmp_path: Path):
        add_image(empty_doc, 0, _png(tmp_path / "dot.png"))
        doc = Document(empty_doc)
        assert any(rel.reltype.endswith("/image") for rel in doc.part.rels.values())


class TestTheFirstParagraphCanBeInserted:
    @pytest.mark.parametrize("after_index", [0, -1])
    def test_insert_paragraph_succeeds_on_an_empty_document(self, empty_doc: str, after_index: int):
        r = insert_paragraph(empty_doc, after_index, "First paragraph")
        assert r["success"] is True, r.get("error")

    @pytest.mark.parametrize("after_index", [0, -1])
    def test_the_text_lands(self, empty_doc: str, after_index: int):
        insert_paragraph(empty_doc, after_index, "First paragraph")
        assert [p.text for p in Document(empty_doc).paragraphs] == ["First paragraph"]

    @pytest.mark.parametrize("after_index", [0, -1])
    def test_it_reports_index_zero_not_one(self, empty_doc: str, after_index: int):
        """after_index + 1 would claim the paragraph landed at 1 or 0 depending
        on which value was passed, for the same single paragraph."""
        r = insert_paragraph(empty_doc, after_index, "First paragraph")
        assert r["inserted_at_index"] == 0

    def test_no_raw_index_error_survives(self, empty_doc: str):
        """-1 used to return IndexError's own words to the caller."""
        r = insert_paragraph(empty_doc, -1, "First paragraph")
        assert "list index out of range" not in str(r)

    def test_append_text_still_works(self, empty_doc: str):
        """The alternative path that always worked must keep working."""
        assert append_text(empty_doc, "Appended")["success"] is True


class TestAPopulatedDocumentIsUnchanged:
    @pytest.fixture()
    def two_paragraphs(self, tmp_path: Path) -> str:
        path = tmp_path / "two.docx"
        doc = Document()
        doc.add_paragraph("one")
        doc.add_paragraph("two")
        doc.save(str(path))
        return str(path)

    def test_a_genuinely_bad_table_anchor_is_still_refused(self, two_paragraphs: str):
        r = add_table(two_paragraphs, 9, 2, 2)
        assert r["success"] is False
        assert "(0-1)" in r["error"], r["error"]

    def test_a_genuinely_bad_image_anchor_is_still_refused(self, two_paragraphs: str, tmp_path: Path):
        r = add_image(two_paragraphs, 9, _png(tmp_path / "dot.png"))
        assert r["success"] is False

    def test_a_genuinely_bad_insert_anchor_is_still_refused(self, two_paragraphs: str):
        r = insert_paragraph(two_paragraphs, 9, "x")
        assert r["success"] is False
        assert "(0-1)" in r["error"], r["error"]

    def test_the_table_still_lands_after_the_named_paragraph(self, two_paragraphs: str):
        add_table(two_paragraphs, 0, 2, 2)
        body = list(Document(two_paragraphs).element.body)
        tags = [el.tag.rsplit("}", 1)[-1] for el in body]
        assert tags[:3] == ["p", "tbl", "p"], tags

    def test_insert_after_the_first_paragraph_still_lands_second(self, two_paragraphs: str):
        insert_paragraph(two_paragraphs, 0, "middle")
        assert [p.text for p in Document(two_paragraphs).paragraphs] == ["one", "middle", "two"]

    def test_insert_at_minus_one_still_lands_first(self, two_paragraphs: str):
        insert_paragraph(two_paragraphs, -1, "zeroth")
        assert [p.text for p in Document(two_paragraphs).paragraphs] == ["zeroth", "one", "two"]
