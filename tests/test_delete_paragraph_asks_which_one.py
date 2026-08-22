"""delete_paragraph's documented minimal call failed on its own default.

The schema marks only file_path required; paragraph_index and match_text are
both optional, defaulting to -1 and "". Calling with just the required argument
-- the call an LLM reading tools/list makes -- carried the -1 into the range
check and came back with:

    error: paragraph_index -1 out of range (0-6)
    hint:  Use read_document to get current paragraph count.

The caller never wrote -1, so the error names a value it cannot map back to
anything it sent, and the hint offers to count the paragraphs when the actual
problem is that no paragraph was chosen. Counting them would not have helped:
every index in 0-6 was already valid.

The tool now says which of the two selectors to supply, and shows both calls.
A real out-of-range index is still reported as before.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_basic.engine import delete_paragraph  # type: ignore[reportMissingImports]

PARAS = [
    "Ad campaign spend review",
    "Google Ads carried 1,939,000 of the 2,500,000 total spend.",
    "205 rows are exact duplicates of another row.",
]


@pytest.fixture()
def doc(tmp_path: Path) -> str:
    path = tmp_path / "report.docx"
    d = Document()
    for text in PARAS:
        d.add_paragraph(text)
    d.save(str(path))
    return str(path)


def texts(path: str) -> list[str]:
    return [p.text for p in Document(path).paragraphs]


class TestTheSchemaMinimalCall:
    def test_it_fails(self, doc: str):
        assert delete_paragraph(doc)["success"] is False

    def test_the_error_does_not_quote_an_index_the_caller_never_sent(self, doc: str):
        error = delete_paragraph(doc)["error"]
        assert "-1" not in error, error

    def test_the_error_says_a_paragraph_must_be_chosen(self, doc: str):
        error = delete_paragraph(doc)["error"]
        assert "paragraph_index" in error and "match_text" in error, error

    def test_the_hint_shows_both_ways_to_choose_one(self, doc: str):
        hint = delete_paragraph(doc)["hint"]
        assert "paragraph_index=" in hint and "match_text=" in hint, hint

    def test_the_hint_names_a_tool_that_lists_the_paragraphs(self, doc: str):
        assert "get_document_outline" in delete_paragraph(doc)["hint"]

    def test_nothing_is_deleted(self, doc: str):
        delete_paragraph(doc)
        assert texts(doc) == PARAS


class TestChoosingOneStillWorks:
    def test_by_index(self, doc: str):
        r = delete_paragraph(doc, paragraph_index=1)
        assert r["success"] is True, r.get("error")
        assert texts(doc) == [PARAS[0], PARAS[2]]

    def test_by_text(self, doc: str):
        r = delete_paragraph(doc, match_text="205 rows")
        assert r["success"] is True, r.get("error")
        assert texts(doc) == PARAS[:2]

    def test_index_zero_is_not_mistaken_for_unset(self, doc: str):
        r = delete_paragraph(doc, paragraph_index=0)
        assert r["success"] is True, r.get("error")
        assert texts(doc) == PARAS[1:]


class TestRealFailuresAreUnchanged:
    def test_an_out_of_range_index_still_says_so(self, doc: str):
        r = delete_paragraph(doc, paragraph_index=99)
        assert r["success"] is False
        assert "99" in r["error"] and "out of range" in r["error"], r["error"]

    def test_text_that_is_not_there_still_says_so(self, doc: str):
        r = delete_paragraph(doc, match_text="Facebook Ads")
        assert r["success"] is False
        assert "match_text" in r["error"], r["error"]

    def test_a_missing_file_still_says_so(self, tmp_path: Path):
        r = delete_paragraph(str(tmp_path / "ghost.docx"), paragraph_index=0)
        assert r["success"] is False
        assert "not found" in r["error"].lower(), r["error"]
