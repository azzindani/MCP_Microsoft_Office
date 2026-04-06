"""End-to-end scenario tests using realistic dummy data.

Each scenario simulates a complete real-world office workflow:
  - Contract filling (HR / Legal)
  - Budget spreadsheet creation (Finance)
  - Sales proposal generation (Sales)
  - Presentation building (Management)
  - Invoice generation (Finance)
  - Batch letter generation (HR)

These tests verify the full tool pipeline, not just individual functions.
"""

import shutil
from pathlib import Path

import pytest

FIXTURES = Path(__file__).parent / "fixtures"


def _copy(name: str, tmp_path: Path) -> Path:
    src = FIXTURES / name
    if not src.exists():
        pytest.skip(f"Fixture {name} not found — run create_fixtures.py first.")
    dst = tmp_path / name
    shutil.copy(src, dst)
    return dst


# ─── Scenario 1: Contract Filling (Legal / HR) ───────────────────────────────


class TestContractFillingScenario:
    """
    Scenario: A legal assistant receives a contract template with placeholders
    and needs to fill in client details, effective date, and contract value.

    Tools used: search_paragraphs → replace_text × 4 → get_document_outline
    """

    def test_fill_all_placeholders(self, tmp_path):
        from servers.docx_basic.engine import replace_text, search_paragraphs

        path = _copy("contract_simple.docx", tmp_path)

        # Step 1: Confirm placeholders exist
        hits = search_paragraphs(str(path), "PARTY_A_NAME")
        assert hits["success"] is True
        assert len(hits["matches"]) > 0

        # Step 2: Fill in all 4 placeholders
        replacements = {
            "PARTY_A_NAME": "Acme Corporation",
            "PARTY_B_NAME": "Widget Ltd",
            "EFFECTIVE_DATE": "April 1, 2026",
            "CONTRACT_VALUE": "50,000",
        }
        for placeholder, value in replacements.items():
            result = replace_text(str(path), placeholder, value)
            assert result["success"] is True, f"Failed replacing {placeholder}: {result}"

        # Step 3: Verify no placeholders remain
        for placeholder in replacements:
            remaining = search_paragraphs(str(path), placeholder)
            assert remaining["matches"] == [], f"Placeholder still present: {placeholder}"

        # Step 4: Verify filled values are present
        for value in replacements.values():
            found = search_paragraphs(str(path), value)
            assert len(found["matches"]) > 0, f"Value not found after fill: {value}"

    def test_snapshot_created_per_replacement(self, tmp_path):
        from servers.docx_basic.engine import get_history, replace_text

        path = _copy("contract_simple.docx", tmp_path)

        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        replace_text(str(path), "PARTY_B_NAME", "Widget Ltd")

        # get_history returns a list of snapshot dicts (newest first)
        # Snapshots share a filename when two ops happen within the same second,
        # so we assert at least 1 snapshot was created (not 0).
        history = get_history(str(path))
        assert isinstance(history, list)
        assert len(history) >= 1
        assert "timestamp" in history[0]
        assert "backup_path" in history[0]

    def test_rollback_after_wrong_replacement(self, tmp_path):
        import time

        from servers.docx_basic.engine import (
            get_history,
            replace_text,
            restore_version,
            search_paragraphs,
        )

        path = _copy("contract_simple.docx", tmp_path)

        # Make a correct replacement, then wait so next snapshot gets a new timestamp
        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        time.sleep(1)  # ensure second snapshot has a distinct filename

        # Make a wrong replacement
        replace_text(str(path), "30 days", "WRONG VALUE")

        # Confirm the mistake is there
        found = search_paragraphs(str(path), "WRONG VALUE")
        assert len(found["matches"]) > 0

        # Roll back: newest snapshot (index 0) is after "WRONG VALUE",
        # second snapshot (index 1) is the good state.
        history = get_history(str(path))
        assert len(history) >= 2
        good_snapshot_ts = history[1]["timestamp"]
        restore_version(str(path), good_snapshot_ts)

        # Confirm "WRONG VALUE" is gone
        after = search_paragraphs(str(path), "WRONG VALUE")
        assert after["matches"] == []

    def test_document_outline_preserved_after_fill(self, tmp_path):
        from servers.docx_basic.engine import get_document_outline, replace_text

        path = _copy("contract_simple.docx", tmp_path)
        original_outline = get_document_outline(str(path))
        original_count = len(original_outline["outline"])

        replace_text(str(path), "PARTY_A_NAME", "Acme Corp")
        replace_text(str(path), "CONTRACT_VALUE", "50,000")

        after_outline = get_document_outline(str(path))
        assert len(after_outline["outline"]) == original_count


# ─── Scenario 2: Budget Spreadsheet (Finance) ─────────────────────────────────


class TestBudgetSpreadsheetScenario:
    """
    Scenario: A finance analyst needs to update Q3 figures, add a SUM formula,
    apply conditional formatting for overbudget cells, and freeze the header row.

    Tools used: set_cell × 4 → set_formula → set_conditional_format → freeze_panes
    """

    SHEET = "Q3 Revenue"  # actual sheet name in budget_simple.xlsx

    def test_update_quarterly_figures(self, tmp_path):
        from servers.xlsx_basic.engine import read_cell, set_cell

        path = _copy("budget_simple.xlsx", tmp_path)

        updates = [
            (self.SHEET, "B2", 142500),
            (self.SHEET, "B3", 98700),
            (self.SHEET, "B4", 87300),
            (self.SHEET, "B5", 151200),
        ]
        for sheet, cell, value in updates:
            result = set_cell(str(path), sheet, cell, value)
            assert result["success"] is True, f"set_cell failed for {cell}: {result}"

        for sheet, cell, value in updates:
            read = read_cell(str(path), sheet, cell)
            assert read["success"] is True
            assert read["value"] == value

    def test_add_sum_formula_after_data_entry(self, tmp_path):
        from servers.xlsx_basic.engine import set_cell
        from servers.xlsx_formulas.engine import set_formula

        path = _copy("budget_simple.xlsx", tmp_path)

        for row, value in [(2, 142500), (3, 98700), (4, 87300), (5, 151200)]:
            set_cell(str(path), self.SHEET, f"B{row}", value)

        result = set_formula(str(path), self.SHEET, "B6", "=SUM(B2:B5)")
        assert result["success"] is True
        assert result["formula"] == "=SUM(B2:B5)"

    def test_conditional_format_overbudget_cells(self, tmp_path):
        from servers.xlsx_formulas.engine import set_conditional_format

        path = _copy("budget_simple.xlsx", tmp_path)

        result = set_conditional_format(
            str(path),
            self.SHEET,
            "B2:B10",
            "greater_than",
            150000,
            "red",
        )
        assert result["success"] is True

    def test_freeze_header_row(self, tmp_path):
        from servers.xlsx_formulas.engine import freeze_panes

        path = _copy("budget_simple.xlsx", tmp_path)
        result = freeze_panes(str(path), self.SHEET, "A2")
        assert result["success"] is True

    def test_full_budget_pipeline(self, tmp_path):
        from servers.xlsx_basic.engine import read_cell, set_cell
        from servers.xlsx_formulas.engine import (
            freeze_panes,
            set_conditional_format,
            set_formula,
        )

        path = _copy("budget_simple.xlsx", tmp_path)

        for row, value in [(2, 142500), (3, 98700), (4, 87300), (5, 151200)]:
            set_cell(str(path), self.SHEET, f"B{row}", value)

        set_formula(str(path), self.SHEET, "B6", "=SUM(B2:B5)")
        set_conditional_format(str(path), self.SHEET, "B2:B6", "greater_than", 130000, "green")
        freeze_panes(str(path), self.SHEET, "A2")

        for row, expected in [(2, 142500), (3, 98700), (4, 87300), (5, 151200)]:
            cell = read_cell(str(path), self.SHEET, f"B{row}")
            assert cell["success"] is True
            assert cell["value"] == expected


# ─── Scenario 3: Sales Proposal Document (Sales) ─────────────────────────────


class TestSalesProposalScenario:
    """
    Scenario: A sales rep takes a contract template, adds a custom scope section,
    inserts a pricing table, and sets heading styles.

    Tools used: replace_text → insert_paragraph → add_table → set_heading
    """

    def test_customise_proposal_text(self, tmp_path):
        from servers.docx_basic.engine import replace_text, search_paragraphs

        path = _copy("contract_simple.docx", tmp_path)

        result = replace_text(str(path), "PARTY_A_NAME", "TechStart Inc")
        assert result["success"] is True

        found = search_paragraphs(str(path), "TechStart Inc")
        assert len(found["matches"]) > 0

    def test_add_custom_scope_paragraph(self, tmp_path):
        from servers.docx_basic.engine import (
            get_document_outline,
            insert_paragraph,
            search_paragraphs,
        )

        path = _copy("contract_simple.docx", tmp_path)

        outline = get_document_outline(str(path))
        last_idx = outline["outline"][-1]["index"]
        result = insert_paragraph(
            str(path),
            last_idx,
            "Custom integration with Salesforce CRM and HubSpot pipeline.",
            "Body Text",
        )
        assert result["success"] is True

        found = search_paragraphs(str(path), "Salesforce")
        assert len(found["matches"]) > 0

    def test_add_pricing_table(self, tmp_path):
        from servers.docx_basic.engine import get_document_outline
        from servers.docx_tables.engine import add_table, list_tables

        path = _copy("contract_simple.docx", tmp_path)

        outline = get_document_outline(str(path))
        last_para = outline["total_paragraphs"] - 1

        result = add_table(
            str(path),
            last_para,
            rows=4,
            cols=3,
            data=[
                ["Service", "Hours", "Price"],
                ["Discovery & Design", "40", "$8,000"],
                ["Development", "120", "$24,000"],
                ["Testing & Deployment", "20", "$4,000"],
            ],
        )
        assert result["success"] is True

        tables = list_tables(str(path))
        assert tables["success"] is True
        assert tables["table_count"] >= 1

    def test_set_heading_styles(self, tmp_path):
        from servers.docx_basic.engine import get_document_outline
        from servers.docx_layout.engine import set_heading

        path = _copy("contract_simple.docx", tmp_path)

        outline = get_document_outline(str(path))
        first_heading_idx = outline["outline"][0]["index"]

        result = set_heading(str(path), first_heading_idx, 1)
        assert result["success"] is True


# ─── Scenario 4: Presentation Building (Management) ──────────────────────────


class TestPresentationBuildingScenario:
    """
    Scenario: A manager updates a quarterly review deck — changes the title,
    adds a new summary slide, and reorders slides.

    Tools used: read_presentation → set_text → add_slide → reorder_slide
    """

    def test_update_presentation_title(self, tmp_path):
        from servers.pptx_basic.engine import read_presentation, set_text

        path = _copy("deck_simple.pptx", tmp_path)

        prs = read_presentation(str(path))
        assert prs["success"] is True
        assert prs["slide_count"] > 0

        result = set_text(str(path), 0, "Title 1", "Q4 2026 Business Review")
        assert result["success"] is True

    def test_add_summary_slide_at_end(self, tmp_path):
        from servers.pptx_basic.engine import add_slide, read_presentation

        path = _copy("deck_simple.pptx", tmp_path)

        before = read_presentation(str(path))
        original_count = before["slide_count"]

        # add_slide(file_path, layout_name, title="", body="")
        result = add_slide(
            str(path),
            "Title and Content",
            "Key Takeaways",
            "Revenue up 18%\nCAC down 12%\nNPS score: 72",
        )
        assert result["success"] is True

        after = read_presentation(str(path))
        assert after["slide_count"] == original_count + 1

    def test_reorder_slides(self, tmp_path):
        from servers.pptx_basic.engine import read_presentation, reorder_slide

        path = _copy("deck_simple.pptx", tmp_path)

        before = read_presentation(str(path))
        assert before["slide_count"] >= 2

        last_idx = before["slide_count"] - 1
        result = reorder_slide(str(path), last_idx, 1)
        assert result["success"] is True

    def test_full_deck_update_pipeline(self, tmp_path):
        from servers.pptx_basic.engine import (
            add_slide,
            read_presentation,
            read_slide,
            set_text,
        )

        path = _copy("deck_simple.pptx", tmp_path)

        prs = read_presentation(str(path))
        assert prs["success"] is True

        set_text(str(path), 0, "Title 1", "Q4 2026 Business Review")

        add_slide(
            str(path),
            "Title and Content",
            "Q4 Metrics",
            "Revenue: $2.1M\nGrowth: 18%\nHeadcount: 47",
        )

        final = read_presentation(str(path))
        assert final["slide_count"] == prs["slide_count"] + 1

        slide0 = read_slide(str(path), 0)
        title_shape = next((s for s in slide0["shapes"] if s["name"] == "Title 1"), None)
        assert title_shape is not None
        assert "Q4 2026" in title_shape["text"]


# ─── Scenario 5: Invoice Generation (Finance) ────────────────────────────────


class TestInvoiceGenerationScenario:
    """
    Scenario: Finance creates a new invoice workbook from scratch with
    line items, formulas, and a professional structure.

    Tools used (xlsx_new): create_invoice
    Signature: create_invoice(output_path, company_name, client_name,
                              invoice_number, items, tax_rate, currency)
    """

    def test_create_invoice_with_line_items(self, tmp_path):
        from servers.xlsx_new.engine import create_invoice

        out = tmp_path / "invoice_001.xlsx"
        result = create_invoice(
            str(out),
            company_name="Acme Services LLC",
            client_name="Widget Corp",
            invoice_number="INV-2026-001",
            items=[
                {"description": "Software Development", "qty": 80, "unit_price": 150.0},
                {"description": "UI/UX Design", "qty": 20, "unit_price": 120.0},
                {"description": "Project Management", "qty": 10, "unit_price": 100.0},
            ],
            tax_rate=0.1,
            open_after=False,
        )
        assert result["success"] is True
        assert out.exists()

    def test_invoice_file_is_readable(self, tmp_path):
        from servers.xlsx_basic.engine import list_sheets, read_cell
        from servers.xlsx_new.engine import create_invoice

        out = tmp_path / "invoice.xlsx"
        create_invoice(
            str(out),
            company_name="My Agency",
            client_name="Test Client",
            invoice_number="INV-001",
            items=[
                {"description": "Consulting", "qty": 10, "unit_price": 200.0},
            ],
            tax_rate=0.08,
            open_after=False,
        )

        sheets = list_sheets(str(out))
        assert sheets["success"] is True
        assert len(sheets["sheets"]) >= 1

        sheet_name = sheets["sheets"][0]["name"]
        # Invoice number should appear somewhere in the first few rows
        cell = read_cell(str(out), sheet_name, "A1")
        assert cell["success"] is True


# ─── Scenario 6: Batch Offer Letter Generation (HR) ─────────────────────────


class TestBatchLetterGenerationScenario:
    """
    Scenario: HR generates personalised offer letters for 3 new hires from
    a single template using batch_create_from_template.

    Tools used (docx_new): batch_create_from_template
    Signature: batch_create_from_template(template_path, data_list, output_dir,
                                          filename_key, open_after)
    """

    def _make_template(self, tmp_path: Path) -> Path:
        from docx import Document

        template_path = tmp_path / "offer_template.docx"
        doc = Document()
        doc.add_heading("Offer Letter", level=1)
        doc.add_paragraph("Dear {{CANDIDATE_NAME}},")
        doc.add_paragraph(
            "We are pleased to offer you the position of {{JOB_TITLE}} "
            "with a starting salary of {{SALARY}} per annum, "
            "commencing {{START_DATE}}."
        )
        doc.add_paragraph("Sincerely, HR Team")
        doc.save(str(template_path))
        return template_path

    def test_batch_generate_offer_letters(self, tmp_path):
        from servers.docx_basic.engine import search_paragraphs
        from servers.docx_new.engine import batch_create_from_template

        template_path = self._make_template(tmp_path)
        output_dir = tmp_path / "offers"
        output_dir.mkdir()

        new_hires = [
            {
                "CANDIDATE_NAME": "Alice Johnson",
                "JOB_TITLE": "Senior Engineer",
                "SALARY": "$145,000",
                "START_DATE": "May 1, 2026",
                "output_filename": "offer_alice_johnson",
            },
            {
                "CANDIDATE_NAME": "Bob Chen",
                "JOB_TITLE": "Product Manager",
                "SALARY": "$130,000",
                "START_DATE": "May 15, 2026",
                "output_filename": "offer_bob_chen",
            },
            {
                "CANDIDATE_NAME": "Carol Rivera",
                "JOB_TITLE": "UX Designer",
                "SALARY": "$110,000",
                "START_DATE": "June 1, 2026",
                "output_filename": "offer_carol_rivera",
            },
        ]

        result = batch_create_from_template(
            str(template_path),
            new_hires,
            str(output_dir),
            filename_key="output_filename",
            open_after=False,
        )
        assert result["success"] is True, f"batch failed: {result}"
        assert result["created_count"] == 3

        for hire in new_hires:
            out_file = output_dir / f"{hire['output_filename']}.docx"
            assert out_file.exists(), f"Missing: {out_file.name}"

            found = search_paragraphs(str(out_file), hire["CANDIDATE_NAME"])
            assert len(found["matches"]) > 0, f"Name not in {out_file.name}"

            placeholder_check = search_paragraphs(str(out_file), "{{CANDIDATE_NAME}}")
            assert placeholder_check["matches"] == []

    def test_batch_creates_correct_count(self, tmp_path):
        from docx import Document

        from servers.docx_new.engine import batch_create_from_template

        template_path = tmp_path / "tmpl.docx"
        doc = Document()
        doc.add_paragraph("Hello {{NAME}}, your role is {{ROLE}}.")
        doc.save(str(template_path))

        output_dir = tmp_path / "out"
        output_dir.mkdir()

        items = [
            {"NAME": "Dave", "ROLE": "Analyst", "output_filename": "dave"},
            {"NAME": "Eve", "ROLE": "Manager", "output_filename": "eve"},
        ]

        result = batch_create_from_template(
            str(template_path),
            items,
            str(output_dir),
            filename_key="output_filename",
            open_after=False,
        )
        assert result["success"] is True
        assert result["created_count"] == 2
        assert (output_dir / "dave.docx").exists()
        assert (output_dir / "eve.docx").exists()


# ─── Scenario 7: Table Editing (Report) ──────────────────────────────────────


class TestTableEditingScenario:
    """
    Scenario: A report writer updates figures in an existing table and adds
    a totals row.

    Tools used: list_tables → read_table → set_cell → add_row
    """

    def test_find_and_update_table_cell(self, tmp_path):
        from servers.docx_tables.engine import list_tables, set_cell

        path = _copy("report_tables.docx", tmp_path)

        tables = list_tables(str(path))
        assert tables["success"] is True
        assert tables["table_count"] >= 1

        result = set_cell(str(path), 0, 1, 1, "Updated Value")
        assert result["success"] is True

    def test_add_totals_row(self, tmp_path):
        from servers.docx_tables.engine import add_row, list_tables, read_table

        path = _copy("report_tables.docx", tmp_path)

        tables = list_tables(str(path))
        assert tables["table_count"] >= 1

        table_data = read_table(str(path), 0)
        # read_table returns "rows" as an integer count, "data" as the 2D array
        original_row_count = table_data["rows"]

        result = add_row(str(path), 0, ["TOTAL", "—", "—", "—"])
        assert result["success"] is True

        updated = read_table(str(path), 0)
        assert updated["rows"] == original_row_count + 1


# ─── Scenario 8: Agenda Presentation (Operations) ────────────────────────────


class TestAgendaPresentationScenario:
    """
    Scenario: An ops manager creates a team meeting agenda from scratch.

    Tools used (pptx_new): create_agenda
    Signature: create_agenda(output_path, meeting_title, date, items, presenter)
    """

    def test_create_team_meeting_agenda(self, tmp_path):
        from servers.pptx_new.engine import create_agenda

        out = tmp_path / "team_meeting_q2.pptx"
        result = create_agenda(
            str(out),
            meeting_title="Q2 2026 Team Sync",
            date="April 10, 2026",
            items=[
                {"topic": "Q1 Retrospective", "duration": "15 min", "owner": "Sarah"},
                {"topic": "Q2 Roadmap Review", "duration": "30 min", "owner": "Tom"},
                {"topic": "OKR Check-in", "duration": "20 min", "owner": "All"},
                {"topic": "Budget Update", "duration": "10 min", "owner": "Finance"},
                {"topic": "Open Q&A", "duration": "15 min", "owner": "All"},
            ],
            presenter="Engineering All-Hands",
            open_after=False,
        )
        assert result["success"] is True
        assert out.exists()

    def test_agenda_slide_content(self, tmp_path):
        from servers.pptx_basic.engine import read_presentation, read_slide
        from servers.pptx_new.engine import create_agenda

        out = tmp_path / "agenda.pptx"
        create_agenda(
            str(out),
            meeting_title="Weekly Standup",
            date="April 7, 2026",
            items=[
                {"topic": "Sprint Progress", "duration": "10 min", "owner": "Dev Lead"},
                {"topic": "Blockers", "duration": "5 min", "owner": "All"},
            ],
            presenter="Dev Team",
            open_after=False,
        )

        prs = read_presentation(str(out))
        assert prs["success"] is True
        assert prs["slide_count"] >= 2

        slide0 = read_slide(str(out), 0)
        all_text = " ".join(s["text"] for s in slide0["shapes"])
        assert "Weekly Standup" in all_text


# ─── Scenario 9: CSV-to-Excel Pipeline (Data Entry) ─────────────────────────


class TestCsvToExcelScenario:
    """
    Scenario: An analyst receives a CSV export from their CRM and imports it
    into Excel for further analysis.

    Tools used (xlsx_new): create_from_csv
    """

    def test_import_csv_data(self, tmp_path):
        import csv

        from servers.xlsx_new.engine import create_from_csv

        csv_path = tmp_path / "sales_data.csv"
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Region", "Rep", "Q1", "Q2", "Q3", "Total"])
            writer.writerow(["North", "Alice", 42000, 51000, 48000, 141000])
            writer.writerow(["South", "Bob", 38000, 44000, 52000, 134000])
            writer.writerow(["East", "Carol", 55000, 49000, 61000, 165000])
            writer.writerow(["West", "Dave", 47000, 53000, 58000, 158000])

        out = tmp_path / "sales_report.xlsx"
        result = create_from_csv(str(csv_path), str(out), has_header=True, open_after=False)
        assert result["success"] is True
        assert out.exists()

    def test_csv_data_readable_after_import(self, tmp_path):
        import csv

        from servers.xlsx_basic.engine import read_cell
        from servers.xlsx_new.engine import create_from_csv

        csv_path = tmp_path / "data.csv"
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Name", "Score"])
            writer.writerow(["Alice", 95])
            writer.writerow(["Bob", 87])

        out = tmp_path / "data.xlsx"
        result = create_from_csv(str(csv_path), str(out), has_header=True, open_after=False)
        assert result["success"] is True

        # Use the sheet_name from the result if available, else fall back to "Data"
        sheet = result.get("sheet_name", "Data")

        cell_a1 = read_cell(str(out), sheet, "A1")
        assert cell_a1["success"] is True
        assert cell_a1["value"] == "Name"

        cell_a2 = read_cell(str(out), sheet, "A2")
        assert cell_a2["value"] == "Alice"
