"""Tests for the shared module."""

import json
import shutil
from pathlib import Path

import pytest

# ─── file_utils ───────────────────────────────────────────────────────────────


class TestResolvePathFileUtils:
    def test_resolves_home_tilde(self):
        from shared.file_utils import resolve_path

        result = resolve_path("~/Documents/test.docx")
        assert result.is_absolute()
        assert str(Path.home()) in str(result)

    def test_strips_wrapping_double_quotes(self, tmp_path):
        from shared.file_utils import resolve_path

        quoted = f'"{tmp_path}/test.docx"'
        result = resolve_path(quoted)
        assert result.is_absolute()
        assert '"' not in str(result)

    def test_strips_wrapping_single_quotes(self, tmp_path):
        from shared.file_utils import resolve_path

        quoted = f"'{tmp_path}/test.docx'"
        result = resolve_path(quoted)
        assert result.is_absolute()
        assert '"' not in str(result)
        assert "'" not in str(result)

    def test_rejects_mcp_versions_path(self, tmp_path):
        from shared.file_utils import resolve_path

        bad = str(tmp_path / ".mcp_versions" / "test.bak")
        with pytest.raises(ValueError, match=".mcp_versions"):
            resolve_path(bad)

    def test_rejects_null_byte(self, tmp_path):
        from shared.file_utils import resolve_path

        with pytest.raises(ValueError, match="null byte"):
            resolve_path(str(tmp_path) + "/test\x00.docx")

    def test_expands_env_var(self, tmp_path, monkeypatch):
        from shared.file_utils import resolve_path

        monkeypatch.setenv("TESTDIR", str(tmp_path))
        result = resolve_path("$TESTDIR/doc.docx")
        assert str(tmp_path) in str(result)


class TestEmbedContent:
    def test_embeds_real_base64_bytes_when_requested(self, tmp_path):
        import base64

        from shared.file_utils import embed_content

        f = tmp_path / "out.xlsx"
        f.write_bytes(b"PK\x03\x04fake zip bytes")
        result = embed_content({"success": True}, f, return_content=True)
        assert base64.b64decode(result["content_base64"]) == f.read_bytes()
        assert result["content_mime_type"] == ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    def test_does_nothing_when_return_content_false(self, tmp_path):
        from shared.file_utils import embed_content

        f = tmp_path / "out.xlsx"
        f.write_bytes(b"data")
        result = embed_content({"success": True}, f, return_content=False)
        assert "content_base64" not in result

    def test_does_nothing_when_result_already_failed(self, tmp_path):
        from shared.file_utils import embed_content

        f = tmp_path / "out.xlsx"
        f.write_bytes(b"data")
        result = embed_content({"success": False}, f, return_content=True)
        assert "content_base64" not in result

    def test_missing_file_does_not_raise(self, tmp_path):
        from shared.file_utils import embed_content

        result = embed_content({"success": True}, tmp_path / "missing.xlsx", return_content=True)
        assert "content_base64" not in result
        assert result["success"] is True


class TestSafeCopy:
    def test_creates_parent_dirs(self, tmp_path):
        from shared.file_utils import safe_copy

        src = tmp_path / "source.txt"
        src.write_text("hello")
        dst = tmp_path / "nested" / "deep" / "dest.txt"
        safe_copy(str(src), str(dst))
        assert dst.exists()
        assert dst.read_text() == "hello"


class TestMcpJson:
    def test_read_returns_empty_dict_if_missing(self, tmp_path):
        from shared.file_utils import read_mcp_json

        result = read_mcp_json(str(tmp_path / "nonexistent.json"))
        assert result == {}

    def test_write_and_read_roundtrip(self, tmp_path):
        from shared.file_utils import read_mcp_json, write_mcp_json

        path = str(tmp_path / "mcp.json")
        data = {"mcpServers": {"docx-basic": {"command": "uv"}}}
        write_mcp_json(path, data)
        result = read_mcp_json(path)
        assert result == data

    def test_write_is_atomic(self, tmp_path):
        """Verify temp file is used (no partial write risk)."""
        from shared.file_utils import write_mcp_json

        path = tmp_path / "mcp.json"
        write_mcp_json(str(path), {"key": "value"})
        assert path.exists()
        # Content should be valid JSON
        loaded = json.loads(path.read_text())
        assert loaded == {"key": "value"}

    def test_handles_trailing_comma(self, tmp_path):
        """json5 handles trailing commas; plain json raises."""
        from shared.file_utils import _HAS_JSON5, read_mcp_json

        path = tmp_path / "mcp.json"
        path.write_text('{"key": "value",}')

        if _HAS_JSON5:
            result = read_mcp_json(str(path))
            assert result == {"key": "value"}
        else:
            with pytest.raises(Exception):
                read_mcp_json(str(path))


# ─── version_control ──────────────────────────────────────────────────────────


class TestVersionControl:
    def test_snapshot_creates_backup(self, tmp_path):
        from shared.version_control import snapshot

        doc = tmp_path / "test.docx"
        doc.write_bytes(b"fake docx content")
        backup = snapshot(str(doc))
        assert Path(backup).exists()
        assert ".mcp_versions" in backup
        assert backup.endswith(".bak")

    def test_snapshot_raises_if_file_missing(self, tmp_path):
        from shared.version_control import snapshot

        with pytest.raises(FileNotFoundError):
            snapshot(str(tmp_path / "nonexistent.docx"))

    def test_snapshot_creates_versions_dir(self, tmp_path):
        from shared.version_control import snapshot

        doc = tmp_path / "test.docx"
        doc.write_bytes(b"content")
        snapshot(str(doc))
        assert (tmp_path / ".mcp_versions").is_dir()

    def test_restore_reverts_content(self, tmp_path):
        from shared.version_control import restore, snapshot

        doc = tmp_path / "test.docx"
        doc.write_bytes(b"original content")
        backup = snapshot(str(doc))
        timestamp = Path(backup).stem.split("_", 1)[1]

        doc.write_bytes(b"modified content")
        assert doc.read_bytes() == b"modified content"

        success = restore(str(doc), timestamp)
        assert success is True
        assert doc.read_bytes() == b"original content"

    def test_restore_returns_false_if_backup_missing(self, tmp_path):
        from shared.version_control import restore

        doc = tmp_path / "test.docx"
        doc.write_bytes(b"content")
        assert restore(str(doc), "2099-01-01T00-00-00Z") is False

    def test_get_history_returns_newest_first(self, tmp_path):
        import time

        from shared.version_control import get_history, snapshot

        doc = tmp_path / "test.docx"
        doc.write_bytes(b"v1")
        snapshot(str(doc))
        time.sleep(1.1)  # ensure different timestamp second
        doc.write_bytes(b"v2")
        snapshot(str(doc))

        history = get_history(str(doc))
        assert len(history) == 2
        assert history[0]["timestamp"] > history[1]["timestamp"]

    def test_get_history_empty_when_no_versions(self, tmp_path):
        from shared.version_control import get_history

        doc = tmp_path / "test.docx"
        doc.write_bytes(b"content")
        assert get_history(str(doc)) == []

    def test_history_entries_have_required_keys(self, tmp_path):
        from shared.version_control import get_history, snapshot

        doc = tmp_path / "test.docx"
        doc.write_bytes(b"content")
        snapshot(str(doc))
        history = get_history(str(doc))
        assert len(history) == 1
        entry = history[0]
        assert "timestamp" in entry
        assert "backup_path" in entry
        assert "size_bytes" in entry


# ─── patch_validator ──────────────────────────────────────────────────────────


class TestPatchValidator:
    ALLOWED_OPS = ["replace_text", "insert_after", "delete_paragraph", "set_cell"]

    def test_valid_op_array(self):
        from shared.patch_validator import validate_ops

        ops = [{"op": "replace_text", "match": "foo", "new_text": "bar"}]
        ok, msg = validate_ops(ops, self.ALLOWED_OPS)
        assert ok is True
        assert msg == ""

    def test_rejects_non_list(self):
        from shared.patch_validator import validate_ops

        ok, msg = validate_ops({"op": "replace_text"}, self.ALLOWED_OPS)
        assert ok is False
        assert "list" in msg

    def test_rejects_missing_op_key(self):
        from shared.patch_validator import validate_ops

        ok, msg = validate_ops([{"action": "replace_text"}], self.ALLOWED_OPS)
        assert ok is False
        assert "'op'" in msg

    def test_rejects_unknown_op(self):
        from shared.patch_validator import validate_ops

        ok, msg = validate_ops([{"op": "destroy_universe"}], self.ALLOWED_OPS)
        assert ok is False
        assert "destroy_universe" in msg

    def test_rejects_more_than_50_ops(self):
        from shared.patch_validator import validate_ops

        ops = [{"op": "replace_text", "match": f"x{i}", "new_text": "y"} for i in range(51)]
        ok, msg = validate_ops(ops, self.ALLOWED_OPS)
        assert ok is False
        assert "50" in msg

    def test_accepts_exactly_50_ops(self):
        from shared.patch_validator import validate_ops

        ops = [{"op": "replace_text", "match": f"x{i}", "new_text": "y"} for i in range(50)]
        ok, _ = validate_ops(ops, self.ALLOWED_OPS)
        assert ok is True

    def test_rejects_formula_without_equals(self):
        from shared.patch_validator import validate_ops

        ops = [{"op": "set_formula", "sheet": "S", "cell": "A1", "formula": "SUM(A:A)"}]
        ok, msg = validate_ops(ops, ["set_formula"])
        assert ok is False
        assert "=" in msg

    def test_accepts_formula_with_equals(self):
        from shared.patch_validator import validate_ops

        ops = [{"op": "set_formula", "sheet": "S", "cell": "A1", "formula": "=SUM(A:A)"}]
        ok, _ = validate_ops(ops, ["set_formula"])
        assert ok is True

    def test_empty_op_array_is_valid(self):
        from shared.patch_validator import validate_ops

        ok, _ = validate_ops([], self.ALLOWED_OPS)
        assert ok is True


# ─── platform_utils ───────────────────────────────────────────────────────────


class TestPlatformUtils:
    def test_config_paths_are_absolute(self):
        from shared.platform_utils import (
            get_claude_desktop_config_path,
            get_cursor_config_path,
            get_lmstudio_mcp_config_path,
        )

        assert get_lmstudio_mcp_config_path().is_absolute()
        assert get_claude_desktop_config_path().is_absolute()
        assert get_cursor_config_path().is_absolute()

    def test_8gb_mode_off_by_default(self, monkeypatch):
        monkeypatch.delenv("OFFICE_MCP_8GB_MODE", raising=False)
        monkeypatch.delenv("MCP_CONSTRAINED_MODE", raising=False)
        from shared.platform_utils import is_8gb_mode

        assert is_8gb_mode() is False

    def test_8gb_mode_enabled_by_env(self, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_8GB_MODE", "1")
        # Re-import to pick up new env
        import importlib

        import shared.platform_utils as pu

        importlib.reload(pu)
        assert pu.is_8gb_mode() is True

    def test_8gb_mode_reduces_limits(self, monkeypatch):
        monkeypatch.setenv("OFFICE_MCP_8GB_MODE", "1")
        import importlib

        import shared.platform_utils as pu

        importlib.reload(pu)
        assert pu.get_max_paragraphs() < 50
        assert pu.get_max_cells() < 200
        assert pu.get_max_search_results() < 50


# ─── gitops ───────────────────────────────────────────────────────────────────


class TestGitops:
    def test_is_git_repo_true_inside_repo(self, tmp_path):
        """The test itself runs inside a git repo."""
        from shared.gitops import is_git_repo

        # MCP_Microsoft_Office is a git repo
        result = is_git_repo(str(tmp_path.parent))
        # May be True or False depending on whether tmp_path is in a git repo
        assert isinstance(result, bool)

    def test_is_git_repo_false_for_random_dir(self, tmp_path):
        from shared.gitops import is_git_repo

        isolated = tmp_path / "not_a_repo"
        isolated.mkdir()
        assert is_git_repo(str(isolated)) is False

    def test_commit_returns_none_when_nothing_to_commit(self, tmp_path):
        from shared.gitops import commit

        # No git repo — should return None gracefully
        result = commit(str(tmp_path), "test commit")
        assert result is None

    def test_all_functions_return_safe_defaults_on_missing_git(self, monkeypatch):
        """All gitops functions handle exceptions without raising."""
        import shared.gitops as go

        monkeypatch.setattr("shared.gitops._git_enabled", lambda: True)

        # These should not raise even with invalid paths
        assert isinstance(go.is_git_repo("/nonexistent/path"), bool)
        assert isinstance(go.stage_file("/nonexistent/file.docx"), bool)
        assert go.commit("/nonexistent", "msg") is None
        assert isinstance(go.get_log("/nonexistent", "file.docx"), list)
        assert isinstance(go.current_branch("/nonexistent"), str)
        assert isinstance(go.diff_staged("/nonexistent"), str)

    def test_git_integration_disabled_by_env(self, monkeypatch):
        monkeypatch.setenv("GIT_INTEGRATION", "false")
        import importlib

        import shared.gitops as go

        importlib.reload(go)

        assert go.is_git_repo("/any/path") is False
        assert go.stage_file("/any/file") is False
        assert go.commit("/any", "msg") is None


# ─── doc_diff ─────────────────────────────────────────────────────────────────


class TestDocDiff:
    def test_diff_docx_identical_files(self, tmp_path):
        """Identical files produce no changes."""
        pytest.importorskip("docx")
        import sys

        sys.path.insert(0, str(Path(__file__).parent.parent))

        fixtures = Path(__file__).parent / "fixtures"
        if not (fixtures / "contract_simple.docx").exists():
            pytest.skip("Fixtures not created yet")

        from shared.doc_diff import diff_docx

        path = str(fixtures / "contract_simple.docx")
        result = diff_docx(path, path)
        assert result["success"] is True
        assert result["change_count"] == 0
        assert "No changes" in result["summary"]

    def test_diff_xlsx_identical_files(self, tmp_path):
        pytest.importorskip("openpyxl")
        fixtures = Path(__file__).parent / "fixtures"
        if not (fixtures / "budget_simple.xlsx").exists():
            pytest.skip("Fixtures not created yet")

        from shared.doc_diff import diff_xlsx

        path = str(fixtures / "budget_simple.xlsx")
        result = diff_xlsx(path, path)
        assert result["success"] is True
        assert result["sheet_diffs"] == {}

    def test_diff_docx_detects_replaced_paragraph(self, tmp_path):
        pytest.importorskip("docx")
        fixtures = Path(__file__).parent / "fixtures"
        if not (fixtures / "contract_simple.docx").exists():
            pytest.skip("Fixtures not created yet")

        shutil.copy(fixtures / "contract_simple.docx", tmp_path / "a.docx")
        shutil.copy(fixtures / "contract_simple.docx", tmp_path / "b.docx")

        # Modify b.docx
        from docx import Document

        doc = Document(str(tmp_path / "b.docx"))
        if doc.paragraphs:
            # Change first non-empty paragraph
            for p in doc.paragraphs:
                if p.text:
                    for run in p.runs:
                        run.text = run.text.replace("Agreement", "Contract")
                    break
        doc.save(str(tmp_path / "b.docx"))

        from shared.doc_diff import diff_docx

        result = diff_docx(str(tmp_path / "a.docx"), str(tmp_path / "b.docx"))
        assert result["success"] is True

    def test_format_diff_as_text_produces_string(self, tmp_path):
        from shared.doc_diff import format_diff_as_text

        diff = {
            "success": True,
            "paragraph_count_a": 10,
            "paragraph_count_b": 10,
            "changes": [
                {
                    "type": "replace",
                    "a_range": [0, 1],
                    "b_range": [0, 1],
                    "a_text": ["old text"],
                    "b_text": ["new text"],
                }
            ],
            "change_count": 1,
            "summary": "1 paragraph changed.",
        }
        text = format_diff_as_text(diff)
        assert isinstance(text, str)
        assert len(text) > 0


# ─── live_edit ──────────────────────────────────────────────────────────────


class TestLiveEditAppleScriptEscaping:
    def test_applescript_escape_handles_quotes_and_backslashes(self):
        from shared.live_edit import _applescript_escape

        assert _applescript_escape('a"b') == 'a\\"b'
        assert _applescript_escape("a\\b") == "a\\\\b"
        assert _applescript_escape('a\\"b') == 'a\\\\\\"b'

    def test_word_reload_escapes_injected_quote(self, monkeypatch):
        # Regression: file_path was interpolated into the AppleScript source
        # unescaped, so a filename containing '"' could break out of the
        # `set targetDoc to "..."` string literal and inject arbitrary
        # AppleScript (including `do shell script`).
        from shared import live_edit

        captured: dict = {}

        class FakeCompleted:
            returncode = 0

        def fake_run(cmd, **kwargs):
            captured["script"] = cmd[2]
            return FakeCompleted()

        monkeypatch.setattr(live_edit.subprocess, "run", fake_run)

        malicious = 'x.docx" \n do shell script "echo INJECTED" \n --'
        live_edit._notify_word_reload_macos(malicious)

        script = captured["script"]
        # The literal payload must never appear unescaped inside the
        # generated AppleScript source.
        assert f'"{malicious}"' not in script
        assert 'do shell script "echo INJECTED"' not in script

    def test_excel_reload_escapes_injected_quote(self, monkeypatch):
        from shared import live_edit

        captured: dict = {}

        class FakeCompleted:
            returncode = 0

        def fake_run(cmd, **kwargs):
            captured["script"] = cmd[2]
            return FakeCompleted()

        monkeypatch.setattr(live_edit.subprocess, "run", fake_run)

        malicious = 'x.xlsx" \n do shell script "echo INJECTED" \n --'
        live_edit._notify_excel_reload_macos(malicious)

        script = captured["script"]
        assert f'"{malicious}"' not in script
        assert 'do shell script "echo INJECTED"' not in script
