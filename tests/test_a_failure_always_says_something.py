"""add_image answered a bad image with an empty error string.

A targeted re-sweep wrote base64 text into /workspace/data/q2_dot.png without
decoding it, then called add_image on it. The extension check passes -- it only
reads the suffix -- so python-docx got the file and raised UnrecognizedImageError,
which it raises bare, with no message. str(exc) is therefore "", and the generic
handler put that straight into the response:

    success: false
    error:   ""
    hint:    "Use restore_version to undo if a snapshot was taken."

An empty `error` reads as a forgotten field rather than a reason, and the hint
offered to undo a change that had not happened. The sweep model's own summary
called it "failed silently", which is exactly right.

This is the second time an exception with nothing to say has reached a caller:
str(KeyError("")) is the two characters '', which is how a chart tool reported a
missing column earlier the same day. So there are two guards here. add_image
names the actual problem, and describe_error() backstops every _error()
constructor in the Office servers so no tool can report a blank failure however
the exception was raised.
"""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from docx_layout.engine import add_image  # type: ignore[reportMissingImports]
from shared.progress import describe_error


def _real_png(path: Path) -> str:
    def chunk(tag: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload)) + tag + payload + struct.pack(">I", zlib.crc32(tag + payload) & 0xFFFFFFFF)
        )

    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(b"\x00\xff\xff\xff"))
        + chunk(b"IEND", b"")
    )
    return str(path)


@pytest.fixture()
def doc(tmp_path: Path) -> str:
    path = tmp_path / "d.docx"
    d = Document()
    d.add_paragraph("Campaign summary")
    d.save(str(path))
    return str(path)


class TestDescribeError:
    @pytest.mark.parametrize("blank", ["", "   ", "''", '""'])
    def test_a_blank_message_is_replaced(self, blank: str):
        assert describe_error(blank) != blank
        assert describe_error(blank).strip()

    def test_a_real_message_is_untouched(self):
        assert describe_error("Table index 0 out of range") == "Table index 0 out of range"

    def test_the_caller_can_choose_the_fallback(self):
        assert describe_error("", "nothing to report") == "nothing to report"

    def test_a_message_that_merely_contains_quotes_survives(self):
        assert describe_error("column '' not found") == "column '' not found"


class TestABadImageSaysWhy:
    @pytest.fixture()
    def not_an_image(self, tmp_path: Path) -> str:
        """What the sweep produced: base64 text saved under a .png name."""
        path = tmp_path / "q2_dot.png"
        path.write_text("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ", encoding="utf-8")
        return str(path)

    def test_it_fails(self, doc: str, not_an_image: str):
        assert add_image(doc, 0, not_an_image)["success"] is False

    def test_the_error_is_not_empty(self, doc: str, not_an_image: str):
        error = add_image(doc, 0, not_an_image)["error"]
        assert error.strip() not in ("", "''", '""'), repr(error)

    def test_the_error_names_the_file(self, doc: str, not_an_image: str):
        assert "q2_dot.png" in add_image(doc, 0, not_an_image)["error"]

    def test_the_error_says_it_is_not_a_readable_image(self, doc: str, not_an_image: str):
        assert "readable image" in add_image(doc, 0, not_an_image)["error"]

    def test_the_hint_explains_the_extension_is_not_enough(self, doc: str, not_an_image: str):
        hint = add_image(doc, 0, not_an_image)["hint"]
        assert ".png" in hint and "contents" in hint

    def test_the_hint_no_longer_offers_to_undo_nothing(self, doc: str, not_an_image: str):
        assert "restore_version" not in add_image(doc, 0, not_an_image)["hint"]

    def test_the_document_is_unharmed(self, doc: str, not_an_image: str):
        add_image(doc, 0, not_an_image)
        assert [p.text for p in Document(doc).paragraphs] == ["Campaign summary"]


class TestTheGoodPathIsUnchanged:
    def test_a_real_png_still_embeds(self, doc: str, tmp_path: Path):
        r = add_image(doc, 0, _real_png(tmp_path / "dot.png"))
        assert r["success"] is True, r.get("error")
        rels = Document(doc).part.rels.values()
        assert any(rel.reltype.endswith("/image") for rel in rels)

    def test_a_missing_file_still_says_so(self, doc: str, tmp_path: Path):
        r = add_image(doc, 0, str(tmp_path / "ghost.png"))
        assert r["success"] is False
        assert "not found" in r["error"].lower()

    def test_an_unsupported_extension_still_says_so(self, doc: str, tmp_path: Path):
        bad = tmp_path / "notes.txt"
        bad.write_text("hello")
        r = add_image(doc, 0, str(bad))
        assert r["success"] is False
        assert "format" in r["error"].lower()


class TestNoOfficeErrorConstructorCanReturnBlank:
    def test_every_error_helper_routes_through_describe_error(self):
        """The guard belongs in the constructor, not at each raise site."""
        import re

        root = Path(__file__).resolve().parents[1] / "servers"
        offenders = []
        for path in list(root.rglob("engine.py")) + list(root.rglob("helpers.py")):
            text = path.read_text(encoding="utf-8")
            for match in re.finditer(r"def _error\(", text):
                body = text[match.start() : match.start() + 700]
                if '"error":' in body and "describe_error" not in body:
                    offenders.append(str(path.relative_to(root)))
        assert not offenders, f"these can still emit a blank error: {offenders}"
