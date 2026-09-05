"""A board paper that had to reference its charts instead of showing them.

    `create_from_sections` (docx) cannot embed images/charts. The board paper
    references charts that live in separate HTML files. Acceptable, but
    image-in-block support would close the loop.
    Room for improvement: an `image` block type accepting a `data/` path or URL.

That was answered by giving `create_from_blocks` an `image` block kind, and
`create_from_sections` was left as {heading, body}. Defensible on paper -- the
richer tool exists -- but a caller reaching for the simpler one and attaching a
chart to a section is making an ordinary request, and the key was dropped in
silence. A picture missing from a document with nothing in the response saying
so is the same failure this fleet has spent the round closing.

Both tools now share `_add_image`, so they cannot disagree about what counts as
a usable image or about how a refusal is worded. And a refusal *is* worded: an
HTML chart -- which is what every chart tool in this fleet produces, and so the
first thing a caller will pass -- comes back as a warning naming the conversion,
not as a document quietly missing a figure.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
for _p in (str(ROOT), str(ROOT / "servers" / "docx_new")):
    if _p not in sys.path:
        sys.path.insert(0, _p)

from docx_new.engine import IMAGE_SOURCE_KEYS, create_from_sections  # noqa: E402


def _png(path: Path) -> str:
    """A real 1x1 PNG, so python-docx is exercised rather than mocked."""
    import base64

    path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
        )
    )
    return str(path)


def _doc(path: Path):
    from docx import Document

    return Document(str(path))


class TestASectionCanCarryAPicture:
    def test_the_image_is_placed(self, tmp_path):
        out = tmp_path / "paper.docx"
        result = create_from_sections(
            str(out),
            "Board Paper",
            [{"heading": "Volume", "body": "Cargo grew 7%.", "image": _png(tmp_path / "c.png")}],
            open_after=False,
        )
        assert result["success"] is True, result.get("error")
        assert result["images_placed"] == 1
        assert len(_doc(out).inline_shapes) == 1

    def test_the_text_around_it_still_arrives(self, tmp_path):
        out = tmp_path / "paper2.docx"
        create_from_sections(
            str(out),
            "Board Paper",
            [{"heading": "Volume", "body": "Cargo grew 7%.", "image": _png(tmp_path / "c.png")}],
            open_after=False,
        )
        text = "\n".join(p.text for p in _doc(out).paragraphs)
        assert "Volume" in text
        assert "Cargo grew 7%." in text

    @pytest.mark.parametrize("key", IMAGE_SOURCE_KEYS)
    def test_every_documented_key_works(self, tmp_path, key):
        """One tuple shared with the block kind, so both accept the same keys."""
        out = tmp_path / f"k_{key}.docx"
        result = create_from_sections(
            str(out), "T", [{"heading": "H", key: _png(tmp_path / f"{key}.png")}], open_after=False
        )
        assert result["success"] is True, result.get("error")
        assert result["images_placed"] == 1

    def test_a_section_of_only_a_picture_is_a_figure_not_an_empty_section(self, tmp_path):
        out = tmp_path / "fig.docx"
        result = create_from_sections(str(out), "T", [{"image": _png(tmp_path / "f.png")}], open_after=False)
        assert result["success"] is True, result.get("error")
        assert result["images_placed"] == 1
        assert not [p for p in result["progress"] if "written empty" in str(p)]


class TestARefusalIsSpokenAloud:
    def test_an_html_chart_is_refused_with_the_way_out(self, tmp_path):
        """The file a caller actually holds, since every chart here is HTML."""
        chart = tmp_path / "chart.html"
        chart.write_text("<html></html>", encoding="utf-8")
        out = tmp_path / "h.docx"
        result = create_from_sections(
            str(out), "T", [{"heading": "H", "body": "B", "image": str(chart)}], open_after=False
        )
        assert result["success"] is True, result.get("error")
        assert result["images_placed"] == 0
        assert any("image not placed" in str(p) for p in result["progress"])

    def test_the_document_is_still_written(self, tmp_path):
        """A bad picture must not cost the caller the paper."""
        chart = tmp_path / "chart.html"
        chart.write_text("<html></html>", encoding="utf-8")
        out = tmp_path / "h2.docx"
        create_from_sections(str(out), "T", [{"heading": "H", "body": "B", "image": str(chart)}], open_after=False)
        assert out.is_file()
        assert "B" in "\n".join(p.text for p in _doc(out).paragraphs)

    def test_a_missing_file_is_reported_not_swallowed(self, tmp_path):
        out = tmp_path / "m.docx"
        result = create_from_sections(
            str(out), "T", [{"heading": "H", "image": str(tmp_path / "nope.png")}], open_after=False
        )
        assert result["images_placed"] == 0
        assert any("image not placed" in str(p) for p in result["progress"])


class TestTheOldShapeIsUntouched:
    def test_sections_without_images_behave_exactly_as_before(self, tmp_path):
        out = tmp_path / "plain.docx"
        result = create_from_sections(
            str(out), "T", [{"heading": "A", "body": "one"}, {"heading": "B", "body": "two"}], open_after=False
        )
        assert result["success"] is True, result.get("error")
        assert result["section_count"] == 2
        assert result["images_placed"] == 0
        assert not _doc(out).inline_shapes

    def test_a_genuinely_empty_section_is_still_flagged(self, tmp_path):
        out = tmp_path / "empty.docx"
        result = create_from_sections(str(out), "T", [{}], open_after=False)
        assert any("written empty" in str(p) for p in result["progress"])
