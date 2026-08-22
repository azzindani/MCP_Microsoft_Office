"""add_table rejected the anchor value insert_paragraph documents.

docx-basic's insert_paragraph takes `after_index` and treats -1 as "insert
before the first paragraph" -- that is its documented edge value, added when an
empty document could not be given a first paragraph. docx-tables' add_table
takes `after_paragraph_index`, the same kind of anchor into the same document,
and rejected -1:

    Paragraph index -1 out of range (0-6)

A coverage sweep hit it exactly that way: it had just used -1 successfully on
one tool and reached for it on the next. The two tools sit one server apart in
the same family, so a caller has no reason to expect the convention to stop at
the boundary.

The value now means the same thing in both: the table goes ahead of paragraph 0.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_basic.engine import insert_paragraph  # type: ignore[reportMissingImports]
from docx_tables.engine import add_table  # type: ignore[reportMissingImports]

FIRST = "Ad campaign spend review"
SECOND = "Google Ads carried 1,939,000 of the 2,500,000 total spend."


@pytest.fixture()
def doc(tmp_path: Path) -> str:
    path = tmp_path / "report.docx"
    d = Document()
    d.add_paragraph(FIRST)
    d.add_paragraph(SECOND)
    d.save(str(path))
    return str(path)


def body_order(path: str) -> list[str]:
    """Tables and paragraphs in the order they appear in the document body."""
    doc = Document(path)
    order = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            order.append("p")
        elif tag == "tbl":
            order.append("tbl")
    return order


class TestAddTableAcceptsMinusOne:
    def test_it_succeeds(self, doc: str):
        r = add_table(doc, -1, 2, 2)
        assert r["success"] is True, f"{r.get('error')} / {r.get('hint')}"

    def test_the_table_lands_before_the_first_paragraph(self, doc: str):
        add_table(doc, -1, 2, 2)
        assert body_order(doc)[:3] == ["tbl", "p", "p"], body_order(doc)

    def test_the_paragraphs_are_untouched(self, doc: str):
        add_table(doc, -1, 2, 2)
        assert [p.text for p in Document(doc).paragraphs] == [FIRST, SECOND]

    def test_the_progress_says_where_it_went(self, doc: str):
        r = add_table(doc, -1, 2, 2)
        said = " ".join(str(s.get("msg", "")) for s in r["progress"])
        assert "before paragraph 0" in said, said

    def test_data_still_lands_in_the_cells(self, doc: str):
        add_table(doc, -1, 2, 2, [["Platform", "Spend"], ["Google Ads", "1,939,000"]])
        table = Document(doc).tables[0]
        assert table.cell(0, 0).text == "Platform"
        assert table.cell(1, 1).text == "1,939,000"


class TestTheTwoToolsAgree:
    """The same anchor value must place content in the same relative position."""

    def test_insert_paragraph_still_puts_text_first(self, doc: str):
        r = insert_paragraph(doc, -1, "Inserted first")
        assert r["success"] is True, r.get("error")
        assert [p.text for p in Document(doc).paragraphs][0] == "Inserted first"

    def test_neither_tool_rejects_minus_one(self, doc: str, tmp_path: Path):
        second = tmp_path / "second.docx"
        Document(doc).save(str(second))
        assert insert_paragraph(doc, -1, "x")["success"] is True
        assert add_table(str(second), -1, 2, 2)["success"] is True


class TestOutOfRangeIsStillRejected:
    @pytest.mark.parametrize("index", [-2, -7, 2, 99])
    def test_a_real_out_of_range_index_still_fails(self, doc: str, index: int):
        r = add_table(doc, index, 2, 2)
        assert r["success"] is False, index
        assert str(index) in r["error"], r["error"]

    def test_the_document_is_unchanged_after_a_rejection(self, doc: str):
        add_table(doc, 99, 2, 2)
        assert Document(doc).tables == []


class TestAppendingIsUnchanged:
    def test_after_the_last_paragraph_still_appends(self, doc: str):
        add_table(doc, 1, 2, 2)
        assert body_order(doc)[:3] == ["p", "p", "tbl"], body_order(doc)

    def test_after_the_first_paragraph_still_lands_between(self, doc: str):
        add_table(doc, 0, 2, 2)
        assert body_order(doc)[:3] == ["p", "tbl", "p"], body_order(doc)

    def test_an_empty_document_still_works(self, tmp_path: Path):
        path = tmp_path / "empty.docx"
        Document().save(str(path))
        r = add_table(str(path), -1, 2, 2)
        assert r["success"] is True, r.get("error")
        assert len(Document(str(path)).tables) == 1
