"""A required argument had an unguessable format and nothing said so.

fetch_section takes `address`, and its whole documented contract was:

    "Fetch content of addressed section or paragraph only."

The live schema carries no parameter descriptions, so that sentence is
everything a caller sees. A section address is written with a section sign --
'§1', or '§1.p3' for a paragraph inside it -- which is not a character anyone
types by accident and not one anybody guesses. Every reasonable attempt:

    '1'  'p1'  'para:1'  '#1'  'section:1'  'Introduction'

produced the same sentence, "Invalid section address: '1'", with a hint that
sent the caller to another tool to find out the format.

The sibling resolver a few lines above already answers this properly -- "Use
§N, §N.pM, pN, or slide[N]/shape[name] notation" -- so one branch explained
itself and the other did not.

Found by calling every tool with only its schema-required arguments.
"""

from __future__ import annotations

import ast
from pathlib import Path

import pytest
from docx import Document

from docx_basic.engine import fetch_section, get_document_index  # type: ignore[reportMissingImports]

SERVER = Path(__file__).parent.parent / "servers" / "docx_basic" / "docx_basic" / "server.py"

GUESSES = ["1", "p1", "para:1", "#1", "section:1", "Introduction", ""]


@pytest.fixture()
def sectioned_doc(tmp_path: Path) -> str:
    path = tmp_path / "doc.docx"
    d = Document()
    d.add_heading("Ad campaign spend review", level=1)
    for line in ("Total spend was 2,503,118.77.", "Google Ads carried 1,939,003.26.", "Facebook Ads 564,115.51."):
        d.add_paragraph(line)
    d.save(str(path))
    return str(path)


class TestABadAddressShowsTheNotation:
    @pytest.mark.parametrize("guess", GUESSES)
    def test_it_refuses(self, sectioned_doc: str, guess: str):
        r = fetch_section(sectioned_doc, guess)
        assert r["success"] is False

    @pytest.mark.parametrize("guess", GUESSES)
    def test_the_section_sign_appears_somewhere(self, sectioned_doc: str, guess: str):
        r = fetch_section(sectioned_doc, guess)
        blob = f"{r.get('error', '')} {r.get('hint', '')}"
        assert "§" in blob, blob

    @pytest.mark.parametrize("guess", GUESSES)
    def test_a_concrete_example_is_given(self, sectioned_doc: str, guess: str):
        r = fetch_section(sectioned_doc, guess)
        blob = f"{r.get('error', '')} {r.get('hint', '')}"
        assert "§1" in blob, blob

    def test_the_paragraph_form_is_shown_too(self, sectioned_doc: str):
        r = fetch_section(sectioned_doc, "1")
        blob = f"{r.get('error', '')} {r.get('hint', '')}"
        assert "p3" in blob or "pM" in blob, blob

    def test_the_address_the_caller_sent_is_quoted_back(self, sectioned_doc: str):
        r = fetch_section(sectioned_doc, "section:1")
        assert "section:1" in r["error"], r["error"]


class TestTheDocstringCarriesTheFormat:
    def docstring(self) -> str:
        tree = ast.parse(SERVER.read_text(encoding="utf-8"))
        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef) and node.name == "fetch_section":
                return ast.get_docstring(node) or ""
        raise AssertionError("fetch_section not found in server.py")

    def test_it_shows_an_address(self):
        assert "§1" in self.docstring(), self.docstring()

    def test_it_is_within_the_eighty_char_limit(self):
        assert len(self.docstring()) <= 80, len(self.docstring())

    def test_it_still_says_what_the_tool_does(self):
        doc = self.docstring().lower()
        assert "section" in doc and "paragraph" in doc


class TestAValidAddressStillWorks:
    def test_the_section_form_reads(self, sectioned_doc: str):
        r = fetch_section(sectioned_doc, "§1")
        assert r["success"] is True, r.get("error")

    def test_the_address_the_index_advertises_is_accepted(self, sectioned_doc: str):
        idx = get_document_index(sectioned_doc)
        assert idx["success"] is True, idx.get("error")
        for section in idx["sections"]:
            r = fetch_section(sectioned_doc, section["address"])
            assert r["success"] is True, (section["address"], r.get("error"))

    def test_a_section_number_that_does_not_exist_says_how_many_there_are(self, sectioned_doc: str):
        r = fetch_section(sectioned_doc, "§9")
        assert r["success"] is False
        assert "9" in r["error"], r["error"]
