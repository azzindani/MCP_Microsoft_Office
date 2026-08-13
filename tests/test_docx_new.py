"""Tests for docx_new engine functions."""

from pathlib import Path

from docx import Document

from docx_new.engine import (
    create_document,
    create_from_sections,
    create_from_template,
    create_from_text,
    create_letter,
    merge_documents,
)

FIXTURES = Path(__file__).parent / "fixtures"
CONTRACT_SIMPLE = FIXTURES / "contract_simple.docx"
REPORT_TABLES = FIXTURES / "report_tables.docx"


# ---------------------------------------------------------------------------
# create_document
# ---------------------------------------------------------------------------


def test_create_document_creates_blank_docx(tmp_path: Path) -> None:
    out = tmp_path / "blank.docx"
    result = create_document(str(out), open_after=False)
    assert result["success"] is True
    assert out.exists()

    doc = Document(str(out))
    assert len(doc.paragraphs) == 0


def test_create_document_return_content_embeds_real_bytes(tmp_path: Path) -> None:
    import base64

    out = tmp_path / "blank.docx"
    result = create_document(str(out), open_after=False, return_content=True)
    assert result["success"] is True
    decoded = base64.b64decode(result["content_base64"])
    assert decoded == out.read_bytes()
    assert decoded[:2] == b"PK"  # a real .docx is a zip archive
    assert result["content_mime_type"] == ("application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def test_create_document_no_return_content_by_default(tmp_path: Path) -> None:
    out = tmp_path / "blank.docx"
    result = create_document(str(out), open_after=False)
    assert "content_base64" not in result


def test_create_document_invalid_path_gives_error(tmp_path: Path) -> None:
    # "blocker" is a file, not a directory — mkdir(parents=True) on a path
    # beneath it must fail.
    blocker = tmp_path / "blocker"
    blocker.write_text("not a directory")
    result = create_document(str(blocker / "blank.docx"), open_after=False)
    assert result["success"] is False
    assert "hint" in result


# ---------------------------------------------------------------------------
# create_from_text
# ---------------------------------------------------------------------------


def test_create_from_text_writes_paragraphs(tmp_path: Path) -> None:
    out = tmp_path / "text.docx"
    paragraphs = [
        {"text": "Title", "style": "Heading 1"},
        {"text": "Body line one", "style": "Normal"},
    ]
    result = create_from_text(str(out), paragraphs, open_after=False)
    assert result["success"] is True
    assert result["paragraph_count"] == 2

    doc = Document(str(out))
    assert doc.paragraphs[0].text == "Title"
    assert doc.paragraphs[1].text == "Body line one"


def test_create_from_text_unknown_style_falls_back_to_normal(tmp_path: Path) -> None:
    out = tmp_path / "text.docx"
    result = create_from_text(str(out), [{"text": "X", "style": "NotAStyle"}], open_after=False)
    assert result["success"] is True
    doc = Document(str(out))
    assert doc.paragraphs[0].style.name == "Normal"


def test_create_from_text_rejects_non_list(tmp_path: Path) -> None:
    out = tmp_path / "text.docx"
    result = create_from_text(str(out), {"text": "not a list"}, open_after=False)
    assert result["success"] is False
    assert "list" in result["error"].lower()


# ---------------------------------------------------------------------------
# create_from_sections
# ---------------------------------------------------------------------------


def test_create_from_sections_writes_title_and_sections(tmp_path: Path) -> None:
    out = tmp_path / "sections.docx"
    sections = [
        {"heading": "Intro", "body": "First section body"},
        {"heading": "Conclusion", "body": "Last section body"},
    ]
    result = create_from_sections(str(out), "Report Title", sections, open_after=False)
    assert result["success"] is True
    assert result["section_count"] == 2

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Report Title" in texts
    assert "Intro" in texts
    assert "First section body" in texts


def test_create_from_sections_rejects_non_list(tmp_path: Path) -> None:
    out = tmp_path / "sections.docx"
    result = create_from_sections(str(out), "Title", "not a list", open_after=False)
    assert result["success"] is False


# ---------------------------------------------------------------------------
# create_from_template
# ---------------------------------------------------------------------------


def test_create_from_template_applies_substitutions(tmp_path: Path) -> None:
    out = tmp_path / "filled.docx"
    result = create_from_template(
        str(CONTRACT_SIMPLE),
        str(out),
        {"PARTY_A_NAME": "Acme Corporation", "PARTY_B_NAME": "Widget Ltd"},
        open_after=False,
    )
    assert result["success"] is True
    assert result["substitutions_applied"] > 0

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Acme Corporation" in full_text
    assert "PARTY_A_NAME" not in full_text


def test_create_from_template_missing_template(tmp_path: Path) -> None:
    out = tmp_path / "filled.docx"
    result = create_from_template(str(tmp_path / "ghost.docx"), str(out), {"X": "Y"}, open_after=False)
    assert result["success"] is False
    assert "not found" in result["error"].lower()


def test_create_from_template_rejects_non_dict_substitutions(tmp_path: Path) -> None:
    out = tmp_path / "filled.docx"
    result = create_from_template(str(CONTRACT_SIMPLE), str(out), ["not", "a", "dict"], open_after=False)
    assert result["success"] is False


# ---------------------------------------------------------------------------
# create_letter
# ---------------------------------------------------------------------------


def test_create_letter_contains_all_fields(tmp_path: Path) -> None:
    out = tmp_path / "letter.docx"
    result = create_letter(
        str(out),
        from_name="Jane Doe",
        to_name="John Smith",
        subject="Contract Renewal",
        body="Please find attached the renewal terms.\nLet us know if you have questions.",
        open_after=False,
    )
    assert result["success"] is True

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane Doe" in full_text
    assert "To: John Smith" in full_text
    assert "Subject: Contract Renewal" in full_text
    assert "Please find attached the renewal terms." in full_text
    assert "Sincerely," in full_text


# ---------------------------------------------------------------------------
# merge_documents
# ---------------------------------------------------------------------------


def test_merge_documents_combines_paragraphs(tmp_path: Path) -> None:
    out = tmp_path / "merged.docx"
    result = merge_documents(
        [str(CONTRACT_SIMPLE), str(REPORT_TABLES)],
        str(out),
        add_page_break=False,
        open_after=False,
    )
    assert result["success"] is True
    assert result["merged_count"] == 2

    merged_doc = Document(str(out))
    doc_a = Document(str(CONTRACT_SIMPLE))
    doc_b = Document(str(REPORT_TABLES))
    assert len(merged_doc.paragraphs) == len(doc_a.paragraphs) + len(doc_b.paragraphs)


def test_merge_documents_empty_list_rejected(tmp_path: Path) -> None:
    out = tmp_path / "merged.docx"
    result = merge_documents([], str(out), open_after=False)
    assert result["success"] is False
    assert "non-empty" in result["error"]


def test_merge_documents_missing_file(tmp_path: Path) -> None:
    out = tmp_path / "merged.docx"
    result = merge_documents([str(tmp_path / "ghost.docx")], str(out), open_after=False)
    assert result["success"] is False
    assert "not found" in result["error"].lower()


# ---------------------------------------------------------------------------
# Global contract: every tool response has a progress field
# ---------------------------------------------------------------------------


def test_all_responses_have_progress_field(tmp_path: Path) -> None:
    results = [
        create_document(str(tmp_path / "a.docx"), open_after=False),
        create_from_text(str(tmp_path / "b.docx"), [{"text": "x"}], open_after=False),
        create_from_sections(str(tmp_path / "c.docx"), "T", [{"heading": "H", "body": "B"}], open_after=False),
        create_from_template(str(CONTRACT_SIMPLE), str(tmp_path / "d.docx"), {"PARTY_A_NAME": "X"}, open_after=False),
        create_letter(str(tmp_path / "e.docx"), "A", "B", "S", "Body", open_after=False),
        merge_documents([str(CONTRACT_SIMPLE)], str(tmp_path / "f.docx"), open_after=False),
    ]
    for r in results:
        assert "progress" in r, f"Missing 'progress' in: {r}"
        assert isinstance(r["progress"], list)
